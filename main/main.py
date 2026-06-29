import re
import os
import io
import tempfile
from pathlib import Path
from utils import *
from gui import *
from statistics_gui import *
from add_test import *
import psutil
from tkinter import *
from collections import namedtuple
import time
import atexit
from tkinter import Toplevel, Label, Button, StringVar
from tkinter import ttk, messagebox, font
import tkinter.font as tkFont
from tkinter.font import Font
from tkcalendar import Calendar
import tkinter.filedialog as filedialog
import sqlite3
from datetime import datetime, timedelta
from dateutil import parser as date_parser
from PIL import Image, ImageTk
from docx import Document
import subprocess
if os.name == 'nt':
    from ctypes import windll
from PIL import Image, ImageTk, ImageGrab
from reportlab.pdfgen import canvas as pdf_canvas
import tempfile
import platform
import subprocess
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import fitz  # PyMuPDF
import math
from pypdf import PdfReader, PdfWriter
from pypdf.generic import (
    NameObject,
    NumberObject,
    TextStringObject,
    BooleanObject,
)

# TO-DO: 
# - Add_Student_Demographics: Dropdown menus for country and site
# - Add_Student_ILP: Resize checkboxes in panedwindows; delete trailing space in left-hand pane
# - Student Window Edit: Entry fields are a little short
# - Unable to filter Attendance by Weekly or Monthly total
# - Remove line underneath Add Attendance (Student) combobox
# - Monthly total isn't producing correct results - needs to be changed to calendar month, not rolling 30-day period (?)
# - Filter AttendanceAdd (tried ChatGPT's solution "Modified AttendanceAddTab Class with Filtering Functionality" but it broke Save Attendance)
# - Resize create lesson plan tab windows
# - Eliminate popup during create lesson: probably need to completely reorganize how the window is being drawn (more akin to the other tabs) / LessonPlanner -> create_lesson() -> self.root (recreating a root window?)
# - Reorder widgets on create_lesson() program info screen
# - Batch edit records
# - Default sort class records by most recent date.
# - Clicking on sort fields in Attendence All is altering the student count
# - Phantom "None" in some Attendance All

PROJECT_ROOT = Path(__file__).resolve().parent.parent
    
DB_FILE_PATH = PROJECT_ROOT / "database" / "master.db"
LP_DIR_PATH = PROJECT_ROOT / "lesson_plans"

MAIN_DIR = Path(__file__).resolve().parent

class AutocompleteCombobox(ttk.Combobox):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._completion_list = [item for item in kwargs.get("values", []) if isinstance(item, str)]
        self._hits = []
        self._typed_string = ""
        self.bind("<KeyRelease>", self._on_key_release)
        self.bind("<Tab>", self._accept_suggestion)

    def _on_key_release(self, event):
        """Triggered on each key release for autocomplete."""
        if event.keysym in ("BackSpace", "Left", "Right", "Up", "Down", "Tab"):
            return

        # Ignore Shift key and other modifiers
        if event.keysym in ("Shift_L", "Shift_R", "Control_L", "Control_R", "Alt_L", "Alt_R"):
            return

        typed_string = self.get()
        if typed_string == "":
            self._reset()
            return

        # Match typed string with available options
        self._hits = [item for item in self._completion_list if item.lower().startswith(typed_string.lower())]

        # Display the closest match visually
        if self._hits:
            self._typed_string = typed_string
            self._show_suggestions()

    def _accept_suggestion(self, event):
        """Accept the suggestion when the Tab key is pressed."""
        if self._hits:
            self.set(self._hits[0])  # Set the combobox value to the first suggestion
            self.icursor(tk.END)  # Place cursor at the end
        return "break"  # Prevent default Tab key behavior

    def _show_suggestions(self):
        """Show the matched suggestions."""
        self.delete(0, tk.END)
        self.insert(0, self._hits[0])
        self.select_range(len(self._typed_string), tk.END)

    def _reset(self):
        """Reset the combobox suggestions."""
        self._hits = []
        self._typed_string = ""

class SRS_Master:
    def __init__(self, master):
        self.master = master
        self.master.title("Student Reporting System")

        # Shared font for the application
        if os.name == 'nt':
            self.default_font = font.Font(family="Segoe UI", size=11)
            self.bold_font = tkFont.Font(family=self.default_font.cget("family"), 
                            size=self.default_font.cget("size"), 
                            weight="bold")
        else:
            self.default_font = font.Font(family="Segoe UI", size=16)

        # Edit Notebook style
        style=ttk.Style()
        style.layout("TNotebook", [])
        style.configure("TNotebook.Tab", font=self.default_font)
        style.configure("TNotebook", highlightbackground="#848a98",tabmargins=0, font=self.default_font)
        
        # Create a Notebook
        self.notebook = ttk.Notebook(master)

        # Create and add tabs
        self.students_tab = StudentsTab(self.notebook, self.default_font)
        self.attendance_tab = AttendanceTab(self.notebook, self.default_font)
        self.class_tab = ClassTab(self.notebook, self.default_font)
        self.testing_tab = TestingTab(self.notebook, self.default_font)
        self.statistics_tab = StatisticsTab(self.notebook, self.default_font, self.bold_font)
        self.notifications_tab = ttk.Frame(self.notebook)

        self.notebook.add(self.students_tab, text='Students')
        self.notebook.add(self.attendance_tab, text='Attendance')
        self.notebook.add(self.class_tab, text='Class')
        self.notebook.add(self.testing_tab, text='Testing')
        self.notebook.add(self.statistics_tab, text='Statistics')
        self.notebook.add(self.notifications_tab, text='Notifications')

        # Pack the Notebook widget to make it visible
        self.notebook.grid(row=0, column=0, padx=10, pady=10, sticky='nsew')

        # Create increase/decrease font size buttons
        self.increase_button = Button(root, text="A", command=self.increase_font_size)
        self.increase_button.grid(row=0, column=1, padx=10, pady=10, sticky='n')

        self.decrease_button = Button(root, text="a", command=self.decrease_font_size)
        self.decrease_button.grid(row=0, column=2, padx=10, pady=10, sticky='n')

        # Configure the master window to expand properly
        self.master.grid_rowconfigure(0, weight=2)
        self.master.grid_columnconfigure(0, weight=2)
        self.master.grid_columnconfigure(1, weight=0)
        self.master.grid_columnconfigure(2, weight=0)

    def cleanup_temp_files(self):
        if self.testing_tab:
            self.testing_tab.cleanup_temp_files()

    def increase_font_size(self):
        # Increase font size
        current_size = self.default_font.cget("size")
        self.default_font.configure(size=current_size + 2)
        self.bold_font.configure(size=current_size + 2)
        # Increase window sizes proportionally
        increase_window_sizes(self.master)  # Assuming `self.master` is the root window
        update_all_treeviews(self.master, self.default_font)
        # update_all_treeviews(self.master, self.bold_font)

    def decrease_font_size(self):
        # Decrease font size
        current_size = self.default_font.cget("size")
        if current_size > 4:  # Prevent font size from becoming too small
            self.default_font.configure(size=current_size - 2)
        # Decrease window sizes proportionally
        decrease_window_sizes(self.master)  # Assuming `self.master` is the root window
        update_all_treeviews(self.master, self.default_font)

class StudentsTab(ttk.Frame):
    def __init__(self, parent, default_font):
        super().__init__(parent)
        self.default_font = default_font
        self._sort_column = None
        self._sort_reverse = False
        self.columns = ['ID', 'Last Name', 'First Name', 'Site', 'Last Class', 'Mobile', 'Email', 'Orientation Date']
        self.visible_columns = self.columns[:]
        self.filter_text = StringVar()  # Add this line
        self.filter_column = StringVar()  # Add this line
        self.student_counter_var = IntVar()
        self.create_widgets()
        self.populate_treeview()

    def on_treeview_click(self, event):
        region = self.tree.identify('region', event.x, event.y)
        if region != 'heading':
            self.on_double_click(event)

    def create_widgets(self):
        # Configure row height of Treeview
        style = ttk.Style(self)
        style.configure("Treeview", rowheight=40, font=self.default_font)  # Set the row height to 40 pixels
        style.configure("Treeview.Heading", font=self.default_font)  # Set the font for Treeview headings

        # Create Treeview widget
        self.tree = ttk.Treeview(self, columns=('ID', 'Last Name', 'First Name', 'Site', 'Last Class', 'Mobile', 'Email', 'Orientation Date'), show='headings')
        self.tree.heading('ID', text='ID')
        self.tree.heading('Last Name', text='Last Name')
        self.tree.heading('First Name', text='First Name')
        self.tree.heading('Site', text='Site')
        self.tree.heading('Last Class', text='Last Class')
        self.tree.heading('Mobile', text='Mobile')
        self.tree.heading('Email', text='Email')
        self.tree.heading('Orientation Date', text='Orientation Date')
        #self.tree.heading('Residency Document', text='Residency Document')

        # Define headings and bind click event
        for col in self.columns:
            self.tree.heading(col, text=col, command=lambda c=col: self.sort_by_column(c, False))
            self.tree.column(col, anchor='center')

        # Bind double click to treeview items (to view student's file)
        self.tree.bind("<Double-1>", self.on_treeview_click)

        # Pack Treeview widget
        self.tree.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky='nsew')

        # Add vertical and horizontal scrollbars
        self.vscrollbar = ttk.Scrollbar(self, orient=VERTICAL, command=self.tree.yview) 
        self.vscrollbar.grid(column=2, row=1, sticky=(N,S)) 
        self.tree['yscrollcommand'] = self.vscrollbar.set

        self.xscrollbar = ttk.Scrollbar(self, orient=HORIZONTAL, command=self.tree.xview) 
        self.xscrollbar.grid(column=0, row=2, columnspan=2, sticky=(E,W))
        self.tree['xscrollcommand'] = self.xscrollbar.set

        # Add Filter button
        self.filter_button = Button(self, text="Filter", command=self.open_filter_window, font=self.default_font)
        self.filter_button.grid(row=0, column=0, padx=5, pady=5, sticky='w')

        # Student Counter label frame
        student_counter_frame = ttk.Frame(self)
        student_counter_frame.grid(row=3, column=0, columnspan=2, padx=5, pady=(2, 0), sticky='w')

        # Student Counter label
        self.student_counter_prefix = Label(student_counter_frame, text="Total Students:", font=self.default_font)
        self.student_counter_prefix.grid(row=3, column=0, padx=2, pady=1, sticky='w')

        # Student Counter label
        self.student_counter = Label(student_counter_frame, textvariable=self.student_counter_var, font=self.default_font)
        self.student_counter.grid(row=3, column=1, padx=2, pady=1, sticky='w')

        # Add Add Student button
        self.add_student_button = Button(self, text="Add Student", command=self.add_student, font=self.default_font)
        self.add_student_button.grid(row=3, column=1, padx=5, pady=5, sticky='w')

        # Configure grid to expand properly
        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure(2, weight=0) 
        
    def populate_treeview(self, filtered_rows=None):
        # Clear existing rows
        for row in self.tree.get_children():
            self.tree.delete(row)

        # Connect to the SQLite database
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        if filtered_rows is None:
            # Fetch student data along with the most recent class_date
            query = '''
                SELECT 
                    s.student_id, 
                    s.last_name, 
                    s.first_name, 
                    GROUP_CONCAT(st.site_name, ", ") AS site_names,  -- This aggregates multiple site_pks into a comma-separated string
                    a.class_date, 
                    s.mobile, 
                    s.email, 
                    s.orientation_date
                FROM 
                    Student s
                LEFT JOIN 
                    (
                        SELECT 
                            student_id, 
                            MAX(class_date) AS class_date
                        FROM 
                            Attendance
                        GROUP BY 
                            student_id
                    ) a 
                    ON s.student_id = a.student_id
                LEFT JOIN
                    StudentSites ss  -- Remove the GROUP BY and join StudentSites directly
                    ON s.student_id = ss.student_id
                LEFT JOIN
                    Sites st
                    ON ss.site_pk = st.site_pk  -- Join Sites to get the site_pk or site_name
                GROUP BY 
                    s.student_id, 
                    s.last_name, 
                    s.first_name, 
                    a.class_date, 
                    s.mobile, 
                    s.email, 
                    s.orientation_date
                ORDER BY 
                    s.student_id DESC;

            '''
            cursor.execute(query)
            rows = cursor.fetchall()
        else:
            rows = filtered_rows

        # Update student counter intvar
        c = len(rows)
        self.student_counter_var.set(c)

        # Insert data into the Treeview
        for row in rows:
            formatted_row = list(row)
            formatted_row[4] = format_date(formatted_row[4]) if formatted_row[4] else ''  # Format last_class date
            formatted_row[5] = format_mobile(formatted_row[5])
            formatted_row[7] = format_date(formatted_row[7]) if formatted_row[7] else ''  # Format dob date
            self.tree.insert('', 'end', values=formatted_row)

        # Adjust column widths based on content
        for col in self.tree['columns']:
            self.tree.column(col, width=tkFont.Font().measure(col))  # Start with the header width
            for item in self.tree.get_children():
                cell_value = self.tree.item(item, 'values')[self.tree['columns'].index(col)]
                cell_width = tkFont.Font().measure(cell_value)
                if self.tree.column(col, width=None) < cell_width:
                    self.tree.column(col, width=cell_width)

        # Close the database connection
        conn.close()

    def sort_by_column(self, col, reverse):
        # Fetch data from Treeview
        data = [(self.tree.set(child, col), child) for child in self.tree.get_children('')]
        
        # Convert data to appropriate types for sorting
        if col in ['ID', 'Mobile:']:
            data = [(int(d.replace('-', '').replace(' ','').replace('(','').replace(')','')) if d else float('-inf'), child) for d, child in data]
        elif col in ['Orientation Date', 'Last Class', 'Class Date']:
            data = [(datetime.strptime(d, '%m-%d-%Y') if d else datetime.min, child) for d, child in data]
        else:
            data = [(d, child) for d, child in data]

        # Sort data
        data.sort(reverse=reverse)
        
        # Rearrange items in sorted positions
        for index, (val, child) in enumerate(data):
            self.tree.move(child, '', index)
        
        # Reverse sort next time
        self._sort_column = col
        self._sort_reverse = not reverse

        # Update heading with new sort direction
        self.tree.heading(col, command=lambda: self.sort_by_column(col, not reverse))

    def apply_filter(self):
        column = self.filter_column.get()
        filter_text = self.filter_text.get().lower().strip()
        filtered_rows = []

        if column and filter_text:
            # Connect to the SQLite database
            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()
            query = '''
                SELECT 
                    s.student_id, 
                    s.last_name, 
                    s.first_name, 
                    GROUP_CONCAT(st.site_name, ", ") AS site_names,  -- This aggregates multiple site_pks into a comma-separated string
                    a.class_date, 
                    s.mobile, 
                    s.email, 
                    s.orientation_date
                FROM 
                    Student s
                LEFT JOIN 
                    (
                        SELECT 
                            student_id, 
                            MAX(class_date) AS class_date
                        FROM 
                            Attendance
                        GROUP BY 
                            student_id
                    ) a 
                    ON s.student_id = a.student_id
                LEFT JOIN
                    StudentSites ss  -- Remove the GROUP BY and join StudentSites directly
                    ON s.student_id = ss.student_id
                LEFT JOIN
                    Sites st
                    ON ss.site_pk = st.site_pk  -- Join Sites to get the site_pk or site_name
                GROUP BY 
                    s.student_id, 
                    s.last_name, 
                    s.first_name, 
                    a.class_date, 
                    s.mobile, 
                    s.email, 
                    s.orientation_date
            '''
            #     ORDER BY 
            #         s.last_name ASC;
            # '''

            cursor.execute(query)
            rows = cursor.fetchall()

            # Apply the filter to the rows
            for row in rows:
                formatted_row = list(row)

                # Convert date to different searchable formats
                def date_matches(date_str, search_text):
                    try:
                        date_obj = datetime.strptime(date_str, '%m-%d-%Y')
                        formatted_date = [
                            date_obj.strftime('%Y-%m-%d'),  # Full date
                            date_obj.strftime('%m-%d-%Y'),  # Alternative full date format
                            date_obj.strftime('%y-%m-%d'),  # Short year with full date
                            date_obj.strftime('%m-%d'),     # Month-day only
                            date_obj.strftime('%Y-%m'),     # Month
                            date_obj.strftime('%m'),        # Year and month
                            date_obj.strftime('%Y'),        # Year only
                            date_obj.strftime('%y'),        # Short year only
                            date_obj.strftime('%B'),        # full month name
                            date_obj.strftime('%b'),        # short month name
                            date_obj.strftime('%B %Y'),     # full month name with year
                            date_obj.strftime('%b %Y')      # short month name with yar                        
                        ]

                        # Normalize month and day by removing leading zeros
                        normalized_date = [
                            re.sub(r'\b0+(\d)', r'\1', d) for d in formatted_date
                        ]

                        # Check if any normalized version matches the search text
                        return any(search_text in nd.lower() for nd in normalized_date)
                    except ValueError:
                        return False

                formatted_row[4] = format_date(formatted_row[4])
                formatted_row[5] = format_mobile(formatted_row[5])
                formatted_row[7] = format_date(formatted_row[7])
                
                # Check if the filter text matches in any format
                if column in ['Last Class', 'Orientation Date']:
                    date_str = formatted_row[self.columns.index(column)]
                    if date_matches(date_str, filter_text):
                        filtered_rows.append(formatted_row)
                else:
                    if filter_text in str(formatted_row[self.columns.index(column)]).lower():
                        filtered_rows.append(formatted_row)

        self.filter_window.destroy()
        self.populate_treeview(filtered_rows if filtered_rows else None)

    def reset_filters(self):
        self.filter_text.set('')
        self.filter_column.set('')
        self.populate_treeview()

    def open_filter_window(self):
        self.filter_window = Toplevel(self)
        self.filter_window.title("Filter Records")

        Label(self.filter_window, text="Select Column:", font=self.default_font).grid(row=0, column=0, padx=10, pady=5, sticky='w')
        column_menu = AutocompleteCombobox(self.filter_window, textvariable=self.filter_column, font=self.default_font)
        column_menu['values'] = self.columns  # Populate with column names
        column_menu.grid(row=0, column=1, padx=10, pady=5, sticky='w')

        Label(self.filter_window, text="Enter Filter Text:", font=self.default_font).grid(row=1, column=0, padx=10, pady=5, sticky='w')
        filter_entry = Entry(self.filter_window, textvariable=self.filter_text)
        filter_entry.grid(row=1, column=1, padx=10, pady=5, sticky='w')

        Button(self.filter_window, text="Apply Filter", font=self.default_font, command=self.apply_filter).grid(row=2, column=0, columnspan=2, pady=5)
        Button(self.filter_window, text="Reset", font=self.default_font, command=self.reset_filters).grid(row=3, column=0, columnspan=2, pady=5)

    def add_student(self):
        def on_tab_changed(*args):
            selected_tab = self.add_student_notebook.tab(self.add_student_notebook.select(), "text")
            if selected_tab == "Demographics":
                center_window(self.add_student_window, width=600, height=800)
            else:
                center_window(self.add_student_window, width=screen_width, height=screen_height)
                self.add_student_window.state("zoomed")

        def update_database(demographics_tab, ilp_tab, db_path=DB_FILE_PATH):
            # Connect to the database
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()

            # Extract demographics data
            demographics_data = {label: var.get() for label, var in demographics_tab.fields}

            # Extract selected sites (from checkbuttons)
            selected_sites = [site_name for site_name, var in demographics_tab.site_vars.items() if var.get()]

            # Extract text widget data
            notes_text_content = ilp_tab.notes_text.get("1.0", "end-1c").strip() or ''
            other_text_content = ilp_tab.other_problems_var.get().strip()

            # Check if student already exists
            query = '''
                SELECT * 
                FROM 
                    Student 
                WHERE 
                    first_name = ? AND last_name = ?
            '''
            cursor.execute(query, (
                demographics_data['First Name:'].strip(), 
                demographics_data['Last Name:'].strip()
            ))
            existing_student = cursor.fetchone()

            if existing_student:
                msg = f"Student {demographics_data['First Name:'].strip()} {demographics_data['Last Name:'].strip()} already exists."
                choice = ask_student_action(demographics_tab, msg)

                if choice == 'modify':
                    student_id = existing_student[0]
                    student_name = f"{demographics_data['Last Name:'].strip()}, {demographics_data['First Name:'].strip()}"
                    self.open_student_window(student_id, student_name)
                elif choice == 'add':
                    add_new_student(demographics_data, ilp_tab, cursor, selected_sites, notes_text_content, other_text_content)
                else:
                    messagebox.showinfo("Operation Cancelled", "No changes were made.")
            else:
                add_new_student(demographics_data, ilp_tab, cursor, selected_sites, notes_text_content, other_text_content)

            # Commit changes and close the connection
            conn.commit()
            conn.close()

        def add_new_student(demographics_data, ilp_tab, cursor, selected_sites, notes_text, other_text):
            # Insert into Student table
            query = '''
                INSERT INTO Student (last_name, first_name, gender, mobile, landline, email, dob, country, first_language, other_languages, job, home_address, nrs_level, orientation_date, student_notes, residency_document, laces_id)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            '''
            cursor.execute(query, (
                demographics_data.get('Last Name:', '').strip(), 
                demographics_data.get('First Name:', '').strip(),
                demographics_data.get('Gender:', '').strip(),
                format_mobile(demographics_data.get('Mobile:', ''), out=True), 
                format_mobile(demographics_data.get('Landline:', ''), out=True), 
                demographics_data.get('Email:', '').strip(), 
                format_date(demographics_data.get('Date of Birth:'), out=True),
                demographics_data.get('Country:', '').strip(), 
                demographics_data.get('Native Language:', '').strip(), 
                demographics_data.get('Other Languages:', '').strip(), 
                demographics_data.get('Job:', '').strip(), 
                demographics_data.get('Home Address:', '').strip(), 
                demographics_data.get('NRS Level:', '').strip(), 
                format_date(demographics_data.get('Orientation Date:'), out=True),
                demographics_data.get('Student Notes:', '').strip(),
                demographics_data.get('Residency Document:', '').strip(),
                demographics_data.get('LACES ID:', '').strip()
            ))

            student_id = cursor.lastrowid

            # Insert selected sites into StudentSites table
            for site_name in selected_sites:
                # Lookup site_pk by site_name
                cursor.execute('SELECT site_pk FROM Sites WHERE site_name = ?', (site_name,))
                site_pk = cursor.fetchone()
                if site_pk:
                    cursor.execute('''
                        INSERT INTO StudentSites (student_id, site_pk)
                        VALUES (?, ?)
                    ''', (student_id, site_pk[0]))

            # Insert ILP data (as previously implemented)
            ilp_data = {
                'goal_job': ilp_tab.long_goal_job.get() or 0,
                'goal_college': ilp_tab.long_goal_college.get() or 0,
                'goal_citizenship': ilp_tab.long_goal_citizenship.get() or 0,
                'goal_speaking': ilp_tab.short_goal_speaking.get() or 0,
                'goal_listening': ilp_tab.short_goal_listening.get() or 0,
                'goal_reading': ilp_tab.short_goal_reading.get() or 0,
                'goal_writing': ilp_tab.short_goal_writing.get() or 0,
                'progress_meet_teacher': ilp_tab.progress_meet.get() or 0,
                'progress_increase_one_level': ilp_tab.progress_increase_one_level.get() or 0,
                'progress_citizenship': ilp_tab.progress_pass_citizenship.get() or 0,
                'progress_get_job': ilp_tab.progress_job.get() or 0,
                'how_attend': ilp_tab.how_attend.get() or 0,
                'how_homework': ilp_tab.how_homework.get() or 0,
                'how_practice_home': ilp_tab.how_practice_at_home.get() or 0,
                'how_meet_job_specialist': ilp_tab.how_job_specialist.get() or 0,
                'problem_childcare': ilp_tab.problems_childcare.get() or 0,
                'problem_transport': ilp_tab.problems_transportation.get() or 0,
                'problem_time': ilp_tab.problems_time.get() or 0,
                'problem_difficult': ilp_tab.problems_difficult.get() or 0,
                'problem_location': ilp_tab.problems_location.get() or 0,
                'problem_moving': ilp_tab.problems_moving.get() or 0,
                'problem_other': other_text,
                'time_less_six': ilp_tab.time_less_six.get() or 0,
                'time_six_to_twelve': ilp_tab.time_six_to_twelve.get() or 0,
                'time_more_twelve': ilp_tab.time_more_twelve.get() or 0,
                'modality_visual': ilp_tab.modality_visual.get() or 0,
                'modality_verbal': ilp_tab.modality_verbal.get() or 0,
                'modality_written': ilp_tab.modality_written.get() or 0,
                'modality_auditory': ilp_tab.modality_auditory.get() or 0,
                'modality_kinesthetic': ilp_tab.modality_kinesthetic.get() or 0,
                'notes': notes_text
            }

            # Insert ILP data (as previously implemented)
            query = '''
                INSERT INTO ILP (student_id, goal_job, goal_college, goal_citizenship, goal_speaking, goal_listening, goal_reading, goal_writing, progress_meet_teacher, progress_increase_one_level, progress_citizenship, progress_get_job, how_attend, how_homework, how_practice_home, how_meet_job_specialist, problem_childcare, problem_transport, problem_time, problem_difficult, problem_location, problem_moving, problem_other, time_less_six, time_six_to_twelve, time_more_twelve, modality_visual, modality_verbal, modality_written, modality_auditory, modality_kinesthetic, notes)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            '''
            cursor.execute(query, (
                student_id,
                ilp_data['goal_job'],
                ilp_data['goal_college'],
                ilp_data['goal_citizenship'],
                ilp_data['goal_speaking'],
                ilp_data['goal_listening'],
                ilp_data['goal_reading'],
                ilp_data['goal_writing'],
                ilp_data['progress_meet_teacher'],
                ilp_data['progress_increase_one_level'],
                ilp_data['progress_citizenship'],
                ilp_data['progress_get_job'],
                ilp_data['how_attend'],
                ilp_data['how_homework'],
                ilp_data['how_practice_home'],
                ilp_data['how_meet_job_specialist'],
                ilp_data['problem_childcare'],
                ilp_data['problem_transport'],
                ilp_data['problem_time'],
                ilp_data['problem_difficult'],
                ilp_data['problem_location'],
                ilp_data['problem_moving'],
                ilp_data['problem_other'],
                ilp_data['time_less_six'],
                ilp_data['time_six_to_twelve'],
                ilp_data['time_more_twelve'],
                ilp_data['modality_visual'],
                ilp_data['modality_verbal'],
                ilp_data['modality_written'],
                ilp_data['modality_auditory'],
                ilp_data['modality_kinesthetic'],
                ilp_data['notes'].strip()
            ))

            messagebox.showinfo("Success", "New student has been added successfully.")

        def populate_database(*args):
            # Check required fields before allowing add
            if not validate_required_demographics(self.demographics_tab):
                return

            # Verify add before populating database
            confirm_add = messagebox.askyesno(
                message='Are you sure you want to add this student?',
                icon='question',
                title='Add Student'
            )

            if confirm_add:
                update_database(self.demographics_tab, self.ilp_tab)
                self.populate_treeview()
                self.add_student_window.destroy()

        def validate_required_demographics(demographics_tab):
            demographics_data = {
                label: var.get().strip()
                for label, var in demographics_tab.fields
            }

            required_fields = [
                "Last Name:",
                "First Name:",
                "Gender:",
                "Date of Birth:",
                "LACES ID:"
            ]

            missing_fields = [
                field.replace(":", "")
                for field in required_fields
                if not demographics_data.get(field)
            ]

            if missing_fields:
                messagebox.showwarning(
                    "Missing Required Field",
                    "Please fill in the following required field(s):\n\n"
                    + "\n".join(f"• {field}" for field in missing_fields)
                )
                return False

            return True

        self.add_student_window = Toplevel(self)
        self.add_student_window.title("Add Student")

        self.add_student_notebook = ttk.Notebook(self.add_student_window)

        self.add_student_window.grid_rowconfigure(0, weight=1)
        self.add_student_window.grid_columnconfigure(0, weight=1)

        # Create and add tabs
        self.demographics_tab = EnterDemographicsTab(self.add_student_notebook, self.default_font)
        self.ilp_tab = EnterILPTab(self.add_student_notebook, self.default_font)

        self.add_student_notebook.add(self.demographics_tab, text='Demographics')
        self.add_student_notebook.add(self.ilp_tab, text='ILP')

        # Pack the Notebook widget to make it visible
        self.add_student_notebook.grid(row=0, column=0, padx=10, pady=10, sticky='nsew')

        # Bind the event to handle tab changes
        self.add_student_notebook.bind("<<NotebookTabChanged>>", on_tab_changed)

        # Create Add button
        add_button = Button(self.add_student_window, text="Add", command=populate_database, font=self.default_font).grid(row=0, column=1, padx=10, pady=14, sticky='n')


    # If student entry is clicked in Students notebook, open student's file
    def open_student_window(self, student_id, student_name):
        """
        Opens the student window with the provided student ID and name.
        """
        def on_tab_changed(*args):
            selected_tab = self.student_notebook.tab(self.student_notebook.select(), "text")
            if selected_tab == "Demographics":
                self.student_window.state("normal")
                center_window(self.student_window, width=650, height=1000)
            elif selected_tab == "Attendance":
                self.student_window.state("normal")
                center_window(self.student_window, width=950, height=700)
            elif selected_tab == "ILP":
                # center_window(self.student_window, width=900, height=900)
                # center_window(self.student_window, width=screen_width, height=screen_height)
                self.student_window.state("zoomed")
            elif selected_tab == "Meetings":
                # center_window(self.student_window, width=675, height=750)
                # center_window(self.student_window, width=screen_width, height=screen_height)
                self.student_window.state("zoomed")

        self.student_window = Toplevel(self)
        self.student_window.title(student_name)
        center_window(self.student_window, width=900, height=1000)

        self.student_notebook = ttk.Notebook(self.student_window)

        self.student_window.grid_rowconfigure(0, weight=1)
        self.student_window.grid_columnconfigure(0, weight=1)

        # Create and add tabs
        self.student_demographics_tab = StudentDemographicsTab(self.student_notebook, student_id, self, self.default_font)
        self.student_attendance_tab = StudentAttendanceTab(self.student_notebook, self, student_id, self.default_font)
        self.student_ilp_tab = StudentILPTab(self.student_notebook, student_id, self.default_font)
        self.student_meetings_tab = StudentMeetingTab(self.student_notebook, self, student_id, self.default_font)
        self.student_testing_tab = StudentTestingTab(self.student_notebook, self.default_font)

        self.student_notebook.add(self.student_demographics_tab, text='Demographics')
        self.student_notebook.add(self.student_attendance_tab, text='Attendance')
        self.student_notebook.add(self.student_ilp_tab, text='ILP')
        self.student_notebook.add(self.student_meetings_tab, text='Meetings')
        self.student_notebook.add(self.student_testing_tab, text='Testing')

        # Pack the Notebook widget to make it visible
        self.student_notebook.grid(row=0, column=0, padx=10, pady=10, sticky='nsew')

        # Bind the event to handle tab changes
        self.student_notebook.bind("<<NotebookTabChanged>>", on_tab_changed)

    def on_double_click(self, event):
        """
        Handles the double-click event to open the student's window.
        """
        curItem = self.tree.focus()
        student_id = self.tree.item(curItem)['values'][0]
        student_name = self.tree.item(curItem)['values'][1] + ', ' + self.tree.item(curItem)['values'][2]

        self.open_student_window(student_id, student_name)

class EnterDemographicsTab(ttk.Frame):
    def __init__(self, parent, default_font):
        super().__init__(parent)
        self.default_font = default_font
        self.fields = [
            ("Last Name:", StringVar()),
            ("First Name:", StringVar()),
            ("Gender:", StringVar()),
            ("Mobile:", StringVar()),
            ("Landline:", StringVar()),
            ("Email:", StringVar()),
            ("Date of Birth:", StringVar()),
            ("Country:", StringVar()),
            ("Native Language:", StringVar()),
            ("Other Languages:", StringVar()),
            ("Job:", StringVar()),
            ("Home Address:", StringVar()),
            ("NRS Level:", StringVar()),
            ("Orientation Date:", StringVar()),
            ("Student Notes:", StringVar()),
            ("Residency Document:", StringVar()),
            ("LACES ID:", StringVar())
        ]

        self.site_vars = {}  # Store BooleanVar for each site
        self.create_widgets()

    def query_sites(self):
        """Query the database to get unique site names"""
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT DISTINCT site_name FROM Sites ORDER BY site_name ASC")
        sites = cursor.fetchall()
        conn.close()
        return [site[0] for site in sites]

    # ---------- Scroll helpers ----------
    def _bind_mousewheel(self, widget, target):
        # Windows/Mac
        widget.bind("<Enter>", lambda e: widget.bind_all("<MouseWheel>", lambda ev: target.yview_scroll(int(-1*(ev.delta/120)), "units")))
        widget.bind("<Leave>", lambda e: widget.unbind_all("<MouseWheel>"))
        # Linux (Button-4/5)
        widget.bind("<Enter>", lambda e: (widget.bind_all("<Button-4>", lambda ev: target.yview_scroll(-1, "units")),
                                          widget.bind_all("<Button-5>", lambda ev: target.yview_scroll( 1, "units"))))
        widget.bind("<Leave>", lambda e: (widget.unbind_all("<Button-4>"),
                                          widget.unbind_all("<Button-5>")))

    def create_widgets(self):
        """Create demographic fields in a scrollable area and add a scrollable site selection frame."""
        # === Title ===
        title = ttk.Label(self, text="Demographics", font=self.default_font)
        title.grid(row=0, column=0, columnspan=2, padx=10, pady=(10, 4), sticky='w')

        # === Scrollable Demographics Fields ===
        # Tunables: approximate per-row height and visible row count
        ROW_HEIGHT = 34     # adjust if your font/padding is taller
        VISIBLE_ROWS = 10   # show ~10 rows then scroll
        canvas_height = ROW_HEIGHT * VISIBLE_ROWS

        demo_container = ttk.Frame(self)
        demo_container.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=10, pady=(0, 8))

        self.demo_canvas = tk.Canvas(demo_container, highlightthickness=0, height=canvas_height)
        demo_scroll = ttk.Scrollbar(demo_container, orient="vertical", command=self.demo_canvas.yview)
        self.demo_canvas.configure(yscrollcommand=demo_scroll.set)

        self.demo_canvas.grid(row=0, column=0, sticky="nsew")
        demo_scroll.grid(row=0, column=1, sticky="ns")

        demo_container.grid_rowconfigure(0, weight=1)
        demo_container.grid_columnconfigure(0, weight=1)

        # Inner frame that actually holds the label/entry widgets
        self.fields_frame = ttk.Frame(self.demo_canvas)
        self.demo_canvas.create_window((0, 0), window=self.fields_frame, anchor="nw")

        # Build the label/entry grid inside fields_frame
        for i, (label_text, string_var) in enumerate(self.fields):
            label = ttk.Label(self.fields_frame, text=label_text, font=self.default_font)
            entry = Entry(self.fields_frame, textvariable=string_var, font=self.default_font)
            label.grid(row=i, column=0, padx=(6, 10), pady=5, sticky='w')
            entry.grid(row=i, column=1, padx=(0, 10), pady=5, sticky='ew')
            self.fields_frame.grid_columnconfigure(1, weight=1)

        # Update scrollregion when inner frame changes size
        def _on_configure(_event=None):
            self.demo_canvas.configure(scrollregion=self.demo_canvas.bbox("all"))
            # Keep the inner frame width synced to the canvas width
            self.demo_canvas.itemconfig(self._fields_window_id, width=self.demo_canvas.winfo_width())

        # Recreate window id so width sync works even after geometry changes
        self.demo_canvas.delete("all")
        self._fields_window_id = self.demo_canvas.create_window((0, 0), window=self.fields_frame, anchor="nw")

        self.fields_frame.bind("<Configure>", _on_configure)
        self.demo_canvas.bind("<Configure>", _on_configure)

        # Mouse wheel scrolling when cursor is over the canvas/inner frame
        self._bind_mousewheel(self.fields_frame, self.demo_canvas)
        self._bind_mousewheel(self.demo_canvas, self.demo_canvas)

        # === Sites (label + scrollable checklist) ===
        site_label = ttk.Label(self, text="Site:", font=self.default_font)
        site_label.grid(row=2, column=0, padx=10, pady=5, sticky='w')

        self.site_container = ttk.Frame(self)
        self.site_container.grid(row=2, column=1, padx=10, pady=5, sticky='nsew')

        self.canvas = tk.Canvas(self.site_container, height=300)  # fixed visible height for sites
        self.site_frame = ttk.Frame(self.canvas)
        self.scrollbar = ttk.Scrollbar(self.site_container, orient="vertical", command=self.canvas.yview)

        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        self.canvas.create_window((0, 0), window=self.site_frame, anchor="nw")
        self.site_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        # Mouse wheel for sites too
        self._bind_mousewheel(self.site_frame, self.canvas)
        self._bind_mousewheel(self.canvas, self.canvas)

        self.populate_site_checkbuttons()

        # Stretch main grid
        self.grid_rowconfigure(1, weight=1)  # the demographics scroll area grows
        self.grid_columnconfigure(1, weight=1)

    def populate_site_checkbuttons(self):
        """Fetch sites and add them as checkbuttons inside the scrollable frame."""
        sites = self.query_sites()

        # Clear existing checkbuttons if they exist
        for widget in self.site_frame.winfo_children():
            widget.destroy()

        self.site_vars = {}  # Reset site vars
        for site in sites:
            var = BooleanVar(value=False)
            checkbutton = tk.Checkbutton(self.site_frame, text=site, variable=var, font=self.default_font, anchor="w")
            checkbutton.pack(anchor="w", padx=5, pady=2, fill="x")
            self.site_vars[site] = var  # Store for reference

    def get_selected_sites(self):
        """Returns a list of selected sites"""
        return [site for site, var in self.site_vars.items() if var.get()]

    def submit_data(self):
        """Placeholder for form submission"""
        print("Selected Sites:", self.get_selected_sites())

class EnterILPTab(ttk.Frame):
    def __init__(self, parent, default_font):
        super().__init__(parent)
        self.default_font = default_font
        self.fields = [
            "long_goal_job",
            "long_goal_college",
            "long_goal_citizenship",
            "short_goal_speaking",
            "short_goal_listening",
            "short_goal_reading",
            "short_goal_writing",
            "progress_meet",
            "progress_increase_one_level",
            "progress_pass_citizenship",
            "progress_job",
            "how_attend",
            "how_homework",
            "how_practice_at_home",
            "how_job_specialist",
            "problems_childcare",
            "problems_transportation",
            "problems_time",
            "problems_difficult",
            "problems_location",
            "problems_moving",
            "problems_other",
            "time_less_six",
            "time_six_to_twelve",
            "time_more_twelve",
            "modality_visual",
            "modality_verbal",
            "modality_written",
            "modality_auditory",
            "modality_kinesthetic"
        ]
        self.string_vars = {field: StringVar() for field in self.fields}
        self.other_problems_var = StringVar()

        for field in self.fields:
            setattr(self, field, self.string_vars[field])
        self.text_widgets = {}
        self.row_states = {}
        self.create_widgets()

    def create_widgets(self):

        # Create Panedwindows to separate the checkboxes 
        self.ilp_goal_pane_left = ttk.Panedwindow(self, orient=VERTICAL)
        self.ilp_goal_pane_right = ttk.Panedwindow(self, orient=VERTICAL)

        # Create a bold font
        bold_font = tkFont.Font(weight="bold")

        # Define the style for the bold Labelframes
        style = ttk.Style()
        style.configure("Bold.TLabelframe.Label", font=bold_font)

        # Helper function to create a scrollable Labelframe
        def create_scrollable_labelframe(parent, text):
            container = ttk.Labelframe(parent, text=text, style="Bold.TLabelframe", width=100, height=100)
            
            # Create a Canvas widget inside the Labelframe
            canvas = Canvas(container)
            scrollbar = ttk.Scrollbar(container, orient=VERTICAL, command=canvas.yview)
            
            # Create a frame inside the Canvas to hold the widgets
            frame = ttk.Frame(canvas)
            
            # Configure the canvas
            canvas.configure(yscrollcommand=scrollbar.set)
            
            # Place the canvas and scrollbar in the Labelframe
            canvas.grid(row=0, column=0, sticky="nsew")
            scrollbar.grid(row=0, column=1, sticky="ns")
            
            # Add the frame to the canvas
            canvas.create_window((0, 0), window=frame, anchor="nw")
            
            # Make the frame resize with the canvas
            frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
            
            # Configure grid in Labelframe to make the canvas expandable
            container.grid_rowconfigure(0, weight=1)
            container.grid_columnconfigure(0, weight=1)
            
            return container, frame

        # Create scrollable Labelframes
        long_term_goals, long_term_goals_frame = create_scrollable_labelframe(self.ilp_goal_pane_left, 'Long-Term Goals')
        short_term_goals, short_term_goals_frame = create_scrollable_labelframe(self.ilp_goal_pane_left, 'Short-Term Goals')
        progress, progress_frame = create_scrollable_labelframe(self.ilp_goal_pane_left, 'Progress')
        how, how_frame = create_scrollable_labelframe(self.ilp_goal_pane_left, "How I'll Succeed")
        problems, problems_frame = create_scrollable_labelframe(self.ilp_goal_pane_right, 'Problems')
        time, time_frame = create_scrollable_labelframe(self.ilp_goal_pane_right, 'Time')
        modality, modality_frame = create_scrollable_labelframe(self.ilp_goal_pane_right, 'Learning Modalities')
        notes, notes_frame = create_scrollable_labelframe(self.ilp_goal_pane_right, 'Notes')

        # Add Labelframes to Panedwindows
        for labelframe in (long_term_goals, short_term_goals, progress, how):
            self.ilp_goal_pane_left.add(labelframe, weight=1)
        for labelframe in (problems, time, modality, notes):
            self.ilp_goal_pane_right.add(labelframe, weight=1)

        # Create Checkbuttons with the given variables and text
        checkbuttons = [
            (long_term_goals_frame, "Get a job or better job", self.long_goal_job),
            (long_term_goals_frame, "Enter college/training program", self.long_goal_college),
            (long_term_goals_frame, "Get my USA citizenship", self.long_goal_citizenship),
            (short_term_goals_frame, "Increase my speaking skills", self.short_goal_speaking),
            (short_term_goals_frame, "Increase my listening skills", self.short_goal_listening),
            (short_term_goals_frame, "Increase my reading skills", self.short_goal_reading),
            (short_term_goals_frame, "Increase my writing skills", self.short_goal_writing),
            (progress_frame, "Meet with my teacher every 90 days", self.progress_meet),
            (progress_frame, "Increase my English by 1 level with a progress test", self.progress_increase_one_level),
            (progress_frame, "Pass my USA citizenship test", self.progress_pass_citizenship),
            (progress_frame, "Get a job or enter college/training program", self.progress_job),
            (how_frame, "Attend every class", self.how_attend),
            (how_frame, "Complete my homework", self.how_homework),
            (how_frame, "Practice English at home or line for 20-30 minutes 3x/week", self.how_practice_at_home),
            (how_frame, "Meet with a job/career specialist", self.how_job_specialist),
            (problems_frame, "Childcare", self.problems_childcare),
            (problems_frame, "Transportation", self.problems_transportation),
            (problems_frame, "Not enough time", self.problems_time),
            (problems_frame, "Too difficult", self.problems_difficult),
            (problems_frame, "A different location", self.problems_location),
            (problems_frame, "Moving", self.problems_moving),
            (time_frame, "Fewer than 6 months", self.time_less_six),
            (time_frame, "6-12 months", self.time_six_to_twelve),
            (time_frame, "More than 1 year", self.time_more_twelve),
            (modality_frame, "Visual", self.modality_visual),
            (modality_frame, "Verbal", self.modality_verbal),
            (modality_frame, "Written", self.modality_written),
            (modality_frame, "Auditory", self.modality_auditory),
            (modality_frame, "Kinesthetic (Hands-On)", self.modality_kinesthetic)
        ]

        # Add text field for "Other" problems
        other_label = Label(problems_frame, text="Other problems:", font=self.default_font)
        other_label.grid(row=len(checkbuttons)*2, column=0, padx=10, pady=5, sticky='nw')
        self.other_problems_text = Entry(problems_frame, textvariable=self.other_problems_var, width=30)
        self.other_problems_text.grid(row=len(checkbuttons)*2+1, column=0, padx=10, pady=5, sticky='nsew')

        # Add text widget to Notes Labelframe
        self.notes_text = Text(notes_frame, width=39, height=5)
        self.notes_text.grid(column=1, row=13, padx=10, pady=10, sticky='nsew')

        for i, (frame, text, var) in enumerate(checkbuttons):
            cb = ttk.Checkbutton(frame, text=text, variable=var)
            cb.grid(row=i*2, column=0, padx=10, pady=5, sticky='nsew')

        # Make the Panedwindow visible
        self.ilp_goal_pane_left.grid(row=0, column=0, padx=10, pady=10, sticky='nsew')
        self.ilp_goal_pane_right.grid(row=0, column=1, padx=10, pady=10, sticky='nsew')

        # Configure the grid to expand properly
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)

        # Ensure the PanedWindow and its components expand
        for pane in [self.ilp_goal_pane_left, self.ilp_goal_pane_right]:
            pane.grid_rowconfigure(0, weight=1)
            pane.grid_columnconfigure(0, weight=1)

class StudentDemographicsTab(ttk.Frame):
    def __init__(self, parent, selected_student_id, students_tab, default_font):
        self.default_font = default_font
        self.selected_student_id = selected_student_id
        self.students_tab = students_tab
        super().__init__(parent)

        self.fields = [
            ("Last Name:", StringVar()),
            ("First Name:", StringVar()),
            ("Gender:", StringVar()),
            ("Mobile:", StringVar()),
            ("Landline:", StringVar()),
            ("Email:", StringVar()),
            ("Date of Birth:", StringVar()),
            ("Country:", StringVar()),
            ("Native Language:", StringVar()),
            ("Other Languages:", StringVar()),
            ("Job:", StringVar()),
            ("Home Address:", StringVar()),
            ("NRS Level:", StringVar()),
            ("Orientation Date:", StringVar()),
            ("Student Notes:", StringVar()),
            ("Residency Document:", StringVar()),
            ("LACES ID:", StringVar())
        ]

        self.site_vars = {}   # site_name -> BooleanVar
        self.labels = {}      # label_text -> ttk.Label (left column)
        self.data_labels = {} # label_text -> ttk.Label (read-only view, middle column)
        self.entries = {}     # label_text -> ttk.Entry (edit view, right column)
        self.edit_mode = False
        self.student_id = None

        self.create_widgets()
        self.populate_data(selected_student_id)

    # ---------- DB helpers ----------
    def query_sites(self):
        """Query the database to get all site names and their primary keys."""
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT site_pk, site_name
            FROM Sites
            ORDER BY site_name ASC
        """)
        sites = cursor.fetchall()
        conn.close()
        return {site_pk: site_name for site_pk, site_name in sites}

    # ---------- UI helpers ----------
    def _bind_mousewheel(self, widget, target):
        # Windows/macOS
        def _on_wheel(e):
            target.yview_scroll(int(-1*(e.delta/120)), "units")
        widget.bind("<Enter>", lambda e: widget.bind_all("<MouseWheel>", _on_wheel))
        widget.bind("<Leave>", lambda e: widget.unbind_all("<MouseWheel>"))
        # Linux (Button-4/5)
        def _on_up(_e):   target.yview_scroll(-1, "units")
        def _on_down(_e): target.yview_scroll( 1, "units")
        widget.bind("<Enter>", lambda e: (widget.bind_all("<Button-4>", _on_up),
                                          widget.bind_all("<Button-5>", _on_down)))
        widget.bind("<Leave>", lambda e: (widget.unbind_all("<Button-4>"),
                                          widget.unbind_all("<Button-5>")))

    def create_widgets(self):
        """Creates demographic fields in a scrollable area and places a scrollable site selection frame."""
        # ---- Title ----
        title = ttk.Label(self, text="Demographics", font=self.default_font)
        title.grid(row=0, column=0, columnspan=3, padx=10, pady=(10, 6), sticky="w")

        # ---- Scrollable fields (shows ~10 rows) ----
        ROW_HEIGHT = 34
        VISIBLE_ROWS = 10
        canvas_height = ROW_HEIGHT * VISIBLE_ROWS

        demo_container = ttk.Frame(self)
        demo_container.grid(row=1, column=0, columnspan=3, sticky="nsew", padx=10, pady=(0, 8))

        self.demo_canvas = tk.Canvas(demo_container, highlightthickness=0, height=canvas_height)
        demo_scroll = ttk.Scrollbar(demo_container, orient="vertical", command=self.demo_canvas.yview)
        self.demo_canvas.configure(yscrollcommand=demo_scroll.set)

        self.demo_canvas.grid(row=0, column=0, sticky="nsew")
        demo_scroll.grid(row=0, column=1, sticky="ns")
        demo_container.grid_rowconfigure(0, weight=1)
        demo_container.grid_columnconfigure(0, weight=1)

        # Inner frame that actually holds label, data_label, entry widgets
        self.fields_frame = ttk.Frame(self.demo_canvas)
        self._fields_window_id = self.demo_canvas.create_window((0, 0), window=self.fields_frame, anchor="nw")

        # Build the three-column grid inside the scrollable frame
        for i, (label_text, string_var) in enumerate(self.fields):
            lbl = ttk.Label(self.fields_frame, text=label_text, font=self.default_font)
            dat = ttk.Label(self.fields_frame, textvariable=string_var, font=self.default_font)
            ent = ttk.Entry(self.fields_frame, textvariable=string_var, font=self.default_font)

            lbl.grid(row=i, column=0, padx=(6, 10), pady=4, sticky="w")
            dat.grid(row=i, column=1, padx=(0, 10), pady=4, sticky="w")
            ent.grid(row=i, column=2, padx=(0, 10), pady=4, sticky="ew")
            ent.grid_remove()  # start hidden (read-only mode)

            self.labels[label_text] = lbl
            self.data_labels[label_text] = dat
            self.entries[label_text] = ent

        # columns sizing inside fields_frame
        self.fields_frame.grid_columnconfigure(1, weight=1)  # data label column
        self.fields_frame.grid_columnconfigure(2, weight=1)  # entry column expands

        # Keep scrollregion updated & sync inner width to canvas width
        def _on_configure(_event=None):
            self.demo_canvas.configure(scrollregion=self.demo_canvas.bbox("all"))
            self.demo_canvas.itemconfig(self._fields_window_id, width=self.demo_canvas.winfo_width())

        self.fields_frame.bind("<Configure>", _on_configure)
        self.demo_canvas.bind("<Configure>", _on_configure)

        # Mouse wheel scrolling
        self._bind_mousewheel(self.fields_frame, self.demo_canvas)
        self._bind_mousewheel(self.demo_canvas, self.demo_canvas)

        # ---- Sites (label + scrollable checklist) ----
        site_label = ttk.Label(self, text="Sites:", font=self.default_font)
        site_label.grid(row=2, column=0, padx=10, pady=5, sticky="w")

        self.site_container = ttk.Frame(self)
        self.site_container.grid(row=2, column=1, columnspan=2, padx=10, pady=5, sticky="nsew")

        self.canvas = tk.Canvas(self.site_container, height=200)  # fixed visible height for sites
        self.site_frame = ttk.Frame(self.canvas)
        self.scrollbar = ttk.Scrollbar(self.site_container, orient="vertical", command=self.canvas.yview)

        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        self.canvas.create_window((0, 0), window=self.site_frame, anchor="nw")
        self.site_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))

        # Mouse wheel for sites list
        self._bind_mousewheel(self.site_frame, self.canvas)
        self._bind_mousewheel(self.canvas, self.canvas)

        # ---- Buttons ----
        self.edit_button = tk.Button(self, text="Edit", command=self.enable_editing, font=self.default_font)
        self.edit_button.grid(row=3, column=0, padx=10, pady=10, sticky="w")

        self.save_button = tk.Button(self, text="Save", command=self.save_changes, font=self.default_font)
        self.save_button.grid(row=3, column=1, padx=10, pady=10, sticky="w")
        self.save_button.grid_remove()  # hidden initially

        self.delete_button = tk.Button(self, text="Delete", command=self.delete_student, font=self.default_font)
        self.delete_button.grid(row=3, column=2, padx=10, pady=10, sticky="e")

        # Main grid stretch
        self.grid_rowconfigure(1, weight=1)  # the scrollable demographics area grows
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure(2, weight=1)

    # ---------- Populate / Sites ----------
    def populate_data(self, student_id):
        """Populate data for the selected student and check their associated sites."""
        self.student_id = student_id
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        cursor.execute('''
            SELECT 
                last_name, 
                first_name,
                gender,
                mobile, 
                landline, 
                email, 
                dob, 
                country, 
                first_language, 
                other_languages, 
                job, 
                home_address, 
                nrs_level, 
                orientation_date, 
                student_notes,
                residency_document,
                laces_id
            FROM 
                Student 
            WHERE student_id = ?
        ''', (student_id,))
        student_data = cursor.fetchone()

        cursor.execute('SELECT site_pk FROM StudentSites WHERE student_id = ?', (student_id,))
        student_sites = [row[0] for row in cursor.fetchall()]
        conn.close()

        if student_data:
            for (label_text, string_var), value in zip(self.fields, student_data):
                value = '' if value is None else value
                if label_text in ['Mobile:', 'Landline:']:
                    string_var.set(format_mobile(value))
                elif label_text in ['Date of Birth:', 'Orientation Date:']:
                    string_var.set(format_date(value))
                else:
                    string_var.set(value)

        self.populate_site_checkbuttons(student_sites)

    def populate_site_checkbuttons(self, selected_sites):
        """Dynamically create checkbuttons for available sites within a scrollable frame."""
        sites = self.query_sites()

        # Clear previous checkbuttons
        for w in self.site_frame.winfo_children():
            w.destroy()

        self.site_vars = {}
        for site_pk, site_name in sites.items():
            var = BooleanVar(value=(site_pk in selected_sites))
            cb = tk.Checkbutton(self.site_frame, text=site_name, variable=var, font=self.default_font, anchor="w")
            cb.pack(anchor="w", padx=5, pady=2, fill="x")
            self.site_vars[site_name] = var

    # ---------- Edit / Save / Delete ----------
    def enable_editing(self):
        student_window = self.master.master
        center_window(student_window, width=900, height=1000)

        for label_text in self.labels.keys():
            self.data_labels[label_text].grid_remove()
            self.entries[label_text].grid()

        self.edit_button.config(state='disabled')
        self.save_button.config(state='normal')
        self.edit_mode = True
        self.save_button.grid()

    def save_changes(self):
        if self.edit_mode:
            updated_data = {label: var.get() for label, var in self.fields}
            selected_sites = [site for site, var in self.site_vars.items() if var.get()]
            self.update_database(updated_data, selected_sites)
            self.populate_data(self.selected_student_id)

            for label_text in self.labels.keys():
                self.entries[label_text].grid_remove()
                self.data_labels[label_text].grid()

            self.edit_button.config(state='normal')
            self.save_button.config(state='disabled')
            self.edit_mode = False
        self.save_button.grid_remove()

    def update_database(self, updated_data, selected_sites):
        """Update student data and associated sites in the database."""
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        cursor.execute('''
            UPDATE 
                Student 
            SET 
                last_name = ?, 
                first_name = ?,
                gender = ?,
                mobile = ?, 
                landline = ?, 
                email = ?, 
                dob = ?, 
                country = ?, 
                first_language = ?, 
                other_languages = ?, 
                job = ?, 
                home_address = ?, 
                nrs_level = ?, 
                orientation_date = ?, 
                student_notes = ?,
                residency_document = ?,
                laces_id = ?
            WHERE 
                student_id = ?
        ''', (
            updated_data['Last Name:'].strip(), 
            updated_data['First Name:'].strip(),
            updated_data['Gender:'].strip(),
            format_mobile(updated_data['Mobile:'], out=True), 
            format_mobile(updated_data['Landline:'], out=True), 
            updated_data['Email:'].strip(), 
            format_date(updated_data['Date of Birth:'], out=True), 
            updated_data['Country:'].strip(), 
            updated_data['Native Language:'].strip(), 
            updated_data['Other Languages:'].strip(), 
            updated_data['Job:'].strip(), 
            updated_data['Home Address:'].strip(), 
            updated_data['NRS Level:'].strip(), 
            format_date(updated_data['Orientation Date:'], out=True), 
            updated_data['Student Notes:'].strip(),
            updated_data['Residency Document:'].strip(),
            updated_data['LACES ID:'].strip(),
            self.student_id
        ))

        # Reset site associations
        cursor.execute('DELETE FROM StudentSites WHERE student_id = ?', (self.student_id,))
        for site_name in selected_sites:
            cursor.execute('SELECT site_pk FROM Sites WHERE site_name = ?', (site_name,))
            site_pk = cursor.fetchone()
            if site_pk:
                cursor.execute('INSERT INTO StudentSites (student_id, site_pk) VALUES (?, ?)', (self.student_id, site_pk[0]))

        conn.commit()
        conn.close()

        # Repopulate the main Students tree
        self.students_tab.populate_treeview()

    def delete_student(self):
        confirm_delete = messagebox.askyesno("Delete Student", "Are you sure you want to delete this student? This action cannot be undone.")
        if not confirm_delete:
            return

        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        try:
            cursor.execute('DELETE FROM Student WHERE student_id = ?', (self.student_id,))
            cursor.execute('DELETE FROM Attendance WHERE student_id = ?', (self.student_id,))
            cursor.execute('DELETE FROM ILP WHERE student_id = ?', (self.student_id,))
            cursor.execute('DELETE FROM Meetings WHERE student_id = ?', (self.student_id,))
            cursor.execute('DELETE FROM Testing WHERE student_id = ?', (self.student_id,))
            conn.commit()
            messagebox.showinfo("Success", "Student record has been deleted successfully.")
            self.students_tab.student_window.destroy()
            self.students_tab.populate_treeview()
        except Exception as e:
            conn.rollback()
            messagebox.showerror("Error", f"An error occurred while deleting the student record: {e}")
        finally:
            conn.close()

class StudentAttendanceTab(ttk.Frame):
    def __init__(self, parent, students_tab, selected_student_id, default_font):
        super().__init__(parent)
        self.default_font = default_font
        self.students_tab = students_tab
        self.selected_student_id = selected_student_id
        self.columns = ['Class Date', 'Site', 'Hours', 'Weekly Total', 'Monthly Total']
        self.visible_columns = self.columns[:]
        self.filter_text = StringVar()  
        self.filter_column = StringVar()  
        self._sort_column = None
        self._sort_reverse = False
        self.string_vars = {}
        self.create_widgets()
        self.populate_treeview(selected_student_id)

    def refresh_students_tab(self):
        if self.students_tab:
            self.students_tab.populate_treeview()

    def create_widgets(self):

        # Add Filter button
        self.filter_button = Button(self, text="Filter", command=self.open_filter_window, font=self.default_font)
        self.filter_button.grid(row=0, column=1, padx=5, pady=5, sticky='w')

        # Add Attendance Record Treeview
        self.tree = ttk.Treeview(self, columns=self.columns, show='headings')
        for col in self.columns:
            self.tree.heading(col, text=col, command=lambda c=col: self.sort_by_column(c, False))
            self.tree.column(col, anchor='center')
        self.tree.grid(row=1, column=1, padx=10, pady=10, sticky='nsew')

        # Add vertical scrollbars
        self.vscrollbar = ttk.Scrollbar(self, orient=VERTICAL, command=self.tree.yview) 
        self.vscrollbar.grid(column=2, row=1, sticky=(N,S)) 
        self.tree['yscrollcommand'] = self.vscrollbar.set

        # Create Add Attendance Record button
        self.add_attendance_button = Button(self, text="Add Attendance", command=self.add_attendance, font=self.default_font)
        self.add_attendance_button.grid(row=2, column=1, padx=5, pady=5, sticky='n')

        # Configure grid to expand properly
        self.grid_rowconfigure(1, weight=1)
        self.grid_rowconfigure(2, weight=1)
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)

        # Bind double click to treeview items (to view student's file)
        self.tree.bind("<Double-1>", self.on_treeview_click)

    def calculate_totals(self, attendance_data):
        def format_number(num):
            if num % 1 == 0:
                return int(num)
            else:
                return num
        weekly_totals = []
        monthly_totals = []
        for i, (current_date_str, current_hours) in enumerate(attendance_data):
            current_date = datetime.strptime(current_date_str, '%m-%d-%Y')

            # Calculate the start of the week (Monday) and end of the week (Sunday)
            start_of_week = current_date - timedelta(days=current_date.weekday())
            end_of_week = start_of_week + timedelta(days=6)

            weekly_total = 0
            monthly_total = 0

            for j, (date_str, hours) in enumerate(attendance_data):
                date_obj = datetime.strptime(date_str, '%m-%d-%Y')

                # Calculate weekly total (Monday to Sunday)
                if start_of_week <= date_obj <= end_of_week:
                    weekly_total += hours

                # Calculate monthly total (30 days from the current date)
                if current_date - timedelta(days=30) < date_obj <= current_date:
                    monthly_total += hours
 
            weekly_totals.append(format_number(weekly_total))
            monthly_totals.append(format_number(monthly_total))
        
        return weekly_totals, monthly_totals

    def populate_treeview(self, selected_student_id, filtered_rows=None):
        # Clear existing rows
        for row in self.tree.get_children():
            self.tree.delete(row)

        # Connect to the SQLite database
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        if filtered_rows is None:
            query = '''
                SELECT 
                    a.attendance_pk, 
                    a.class_date, 
                    a.site, 
                    a.class_hours
                FROM 
                    Attendance a
                WHERE 
                    a.student_id = ?
                ORDER BY 
                    a.class_date DESC
            '''
            cursor.execute(query, (selected_student_id,))
            rows = cursor.fetchall()
        else:
            rows = filtered_rows

        processed_rows = [(row[0], format_date(row[1]), row[2], row[3]) for row in rows]

        # Calculate weekly and monthly totals
        attendance_data = [(row[1], row[3]) for row in processed_rows]
        weekly_totals, monthly_totals = self.calculate_totals(attendance_data)

        # Consruct a map of attendance_pk for later use in modifying rows
        self.attendance_pk_map = {}

        # Insert data into the Treeview
        for row, weekly_total, monthly_total in zip(processed_rows, weekly_totals, monthly_totals):
            item_id = self.tree.insert('', 'end', values=(row[1], row[2], row[3], weekly_total, monthly_total))
            self.attendance_pk_map[item_id] = row[0]

        # Close the database connection
        conn.close()

        # Adjust column widths based on content
        for col in self.tree['columns']:
            self.tree.column(col, width=tkFont.Font().measure(col))  # Start with the header width
            for item in self.tree.get_children():
                cell_value = self.tree.item(item, 'values')[self.tree['columns'].index(col)]
                cell_width = tkFont.Font().measure(cell_value) + 100
                if self.tree.column(col, width=None) < cell_width:
                    self.tree.column(col, width=cell_width)

    def sort_by_column(self, col, reverse):
        # Fetch data from Treeview
        data = [(self.tree.set(child, col), child) for child in self.tree.get_children('')]

        # Convert data to appropriate types for sorting
        if col == 'Class Date':
            data = [(datetime.strptime(d, '%m-%d-%Y') if d else datetime.min, child) for d, child in data]
        elif col in ['Hours', 'Weekly Total', 'Monthly Total']:
            data = [(float(d) if d else float('-inf'), child) for d, child in data]
        else:
            data = [(d, child) for d, child in data]

        # Sort data
        data.sort(reverse=reverse)

        # Rearrange items in sorted positions
        for index, (val, child) in enumerate(data):
            self.tree.move(child, '', index)

        # Reverse sort next time
        self._sort_column = col
        self._sort_reverse = not reverse

        # Update heading with new sort direction
        self.tree.heading(col, command=lambda: self.sort_by_column(col, not reverse))

    def apply_filter(self):
        column = self.filter_column.get()
        filter_text = self.filter_text.get()
        filtered_rows = []

        if column and filter_text:
            # Connect to the SQLite database
            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()

            # Fetch attendance data for the selected student
            query = '''
                SELECT a.class_date, a.site, a.class_hours
                FROM Attendance a
                WHERE a.student_id = ?
            '''
            cursor.execute(query, (self.selected_student_id,))
            rows = cursor.fetchall()
            conn.close()

            processed_rows = [(format_date(row[0]), row[1], row[2]) for row in rows]

            # Apply the filter to the rows
            column_index = self.columns.index(column) if column in self.columns else -1
            if column_index == -1:
                for row in processed_rows:
                    if column == 'Weekly Total' or column == 'Monthly Total':
                        attendance_data = [(r[0], r[2]) for r in processed_rows]
                        weekly_totals, monthly_totals = self.calculate_totals(attendance_data)
                        if column == 'Weekly Total':
                            for row, weekly_total in zip(processed_rows, weekly_totals):
                                if filter_text.lower() in str(weekly_total).lower():
                                    filtered_rows.append(row)
                        elif column == 'Monthly Total':
                            for row, monthly_total in zip(processed_rows, monthly_totals):
                                if filter_text.lower() in str(monthly_total).lower():
                                    filtered_rows.append(row)
            else:
                for row in processed_rows:
                    cell_value = str(row[column_index]).lower()
                    if column == 'Hours':
                        try:
                            filter_value = float(filter_text)
                            if float(row[column_index]) == filter_value:
                                filtered_rows.append(row)
                        except ValueError:
                            continue
                    elif filter_text.lower() in cell_value:
                        filtered_rows.append(row)

            # Calculate weekly and monthly totals for filtered rows
            if filtered_rows:
                attendance_data = [(row[0], row[2]) for row in filtered_rows]
                weekly_totals, monthly_totals = self.calculate_totals(attendance_data)

                # Insert data into the Treeview with totals
                self.tree.delete(*self.tree.get_children())
                for row, weekly_total, monthly_total in zip(filtered_rows, weekly_totals, monthly_totals):
                    self.tree.insert('', 'end', values=(row[0], row[1], row[2], weekly_total, monthly_total))
            else:
                self.populate_treeview(self.selected_student_id, None)

    def reset_filters(self):
        self.filter_text.set('')
        self.filter_column.set('')
        self.populate_treeview(self.selected_student_id)

    def open_filter_window(self):
        self.filter_window = Toplevel(self)
        self.filter_window.title("Filter Records")

        Label(self.filter_window, text="Select Column:", font=self.default_font).grid(row=0, column=0, padx=10, pady=5, sticky='w')
        column_menu = AutocompleteCombobox(self.filter_window, textvariable=self.filter_column, font=self.default_font)
        column_menu['values'] = self.columns  # Populate with column names
        column_menu.grid(row=0, column=1, padx=10, pady=5, sticky='w')

        Label(self.filter_window, text="Enter Filter Text:", font=self.default_font).grid(row=1, column=0, padx=10, pady=5, sticky='w')
        filter_entry = Entry(self.filter_window, textvariable=self.filter_text, font=self.default_font)
        filter_entry.grid(row=1, column=1, padx=10, pady=5, sticky='w')

        Button(self.filter_window, text="Apply Filter", command=self.apply_filter, font=self.default_font).grid(row=2, column=0, columnspan=2, pady=5)
        Button(self.filter_window, text="Reset", command=self.reset_filters, font=self.default_font).grid(row=3, column=0, columnspan=2, pady=5)

    def add_attendance(self):
        # Create the Add Attendance window and center it
        self.add_attendance_window = Toplevel(self)
        self.add_attendance_window.withdraw()
        self.add_attendance_window.title("Add Attendance")
        self.add_attendance_window.grid_rowconfigure(0, weight=1)
        self.add_attendance_window.grid_columnconfigure(0, weight=1)
        center_window(self.add_attendance_window, 1000, 800)

        # Define the attendance record label and entry fields
        self.fields = [
            ("Class Date:", "class_date"),
            ("Site:", "site"),
            ("Hours:", "hours")]

        # Create Save button
        self.save_attendance_button = Button(self.add_attendance_window, text="Save", command=self.save_attendance, state='disabled', font=self.default_font)
        self.save_attendance_button.grid(row=len(self.fields)+2, column=0, columnspan=2, pady=10)
        
        # Method to check if all fields are filled
        def check_fields():
            if all(var.get().strip() for var in self.string_vars.values()):
                self.save_attendance_button.config(state='normal')
            else:
                self.save_attendance_button.config(state='disabled')

        # Create labels and entries
        for i, (label_text, var_name) in enumerate(self.fields):
            i+=1
            label = Label(self.add_attendance_window, text=label_text, font=self.default_font)
            label.grid(row=i, column=0, padx=10, pady=5, sticky='n')

            string_var = StringVar()
            self.string_vars[var_name] = string_var  # Store the StringVar object in the dictionary

            entry = Entry(self.add_attendance_window, textvariable=string_var, font=self.default_font)
            entry.grid(row=i, column=1, padx=10, pady=5, sticky='n')

            if label_text == "Class Date:":
                # Create the calendar widget
                large_font = font.Font(family="Helvetica", size=14, weight="bold")
                self.cal = Calendar(self.add_attendance_window, selectmode='day', 
                                    year=datetime.today().year, month=datetime.today().month, 
                                    day=datetime.today().day, font=large_font, 
                                    headersbackground="lightblue",
                                    background="lightgrey", foreground="black",
                                    selectbackground="blue", selectforeground="white")
                self.cal.grid(row=0, column=1, columnspan=2, padx=10, pady=5)
                # Bind the calendar selection event to update the "Class Date:" entry box
                self.cal.bind("<<CalendarSelected>>", lambda event, var=string_var: var.set(self.cal.get_date()))
            elif label_text == "Site:":
                # Create a combobox for the Site field
                site_names = get_site_names()
                combobox = AutocompleteCombobox(self.add_attendance_window, textvariable=string_var, values=site_names, font=self.default_font)
                combobox.grid(row=i, column=1, padx=10, pady=5, sticky='n')
            elif label_text == "Hours:":
                # Create a combobox for the Hours field with a default value
                combobox = AutocompleteCombobox(self.add_attendance_window, textvariable=string_var, values=[.5, 1, 1.5, 2, 2.5], font=self.default_font)
                combobox.set("2.5")  # Set default value
                combobox.grid(row=i, column=1, padx=10, pady=5, sticky='n')
             
            # Add the trace after the widget creation to ensure button initialization is done
            string_var.trace('w', lambda name, index, mode, var=string_var: check_fields())

        # Configure the grid to expand properly
        for i in range(1, len(self.fields)+1):
            self.add_attendance_window.grid_rowconfigure(i, weight=1)
            self.add_attendance_window.grid_columnconfigure(0, weight=1)

        # Configure the grid to expand properly
        for i in range(1, len(self.fields)+2):
            self.add_attendance_window.grid_rowconfigure(i, weight=1)
        self.add_attendance_window.grid_columnconfigure(0, weight=1)
        self.add_attendance_window.grid_columnconfigure(1, weight=1)

        self.add_attendance_window.deiconify()

    def get_string_vars(self):
        return self.string_vars

    def save_attendance(self):
        msg = f"Save attendance changes?"
        if messagebox.askyesno("Add Attendance", msg):
            # Get the values from the entry fields
            class_date = self.string_vars['class_date'].get()
            site = self.string_vars['site'].get()
            hours = self.string_vars['hours'].get()

            # Connect to the SQLite database
            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()
            query = '''
                INSERT INTO Attendance (student_id, class_date, site, class_hours)
                VALUES (?, ?, ?, ?)
            '''
            # Execute the insert query
            cursor.execute(query, (
                self.selected_student_id, 
                format_date(class_date, out=True), 
                site, 
                hours
            ))
            # Commit the changes and close the connection
            conn.commit()
            conn.close()

            # Close the Add Attendance window
            self.add_attendance_window.destroy()

            # Refresh the attendance Treeview
            self.populate_treeview(self.selected_student_id)

            # Refresh the StudentsTab treeview
            self.refresh_students_tab()
    
    def on_treeview_click(self, event):
        region = self.tree.identify('region', event.x, event.y)
        if region != 'heading':
            self.on_double_click(event)

    def on_double_click(self, event):
        curItem = self.tree.focus()
        item_values = self.tree.item(curItem)['values']
        self.selected_attendance_pk = self.attendance_pk_map.get(curItem)  # Retrieve attendance_pk using the item_id
        class_date, site, hours = item_values[0], item_values[1], item_values[2]
        title_string = f"Record: {class_date} @ {site}"

        self.attendance_record_window = Toplevel(self)
        self.attendance_record_window.title(title_string)
        center_window(self.attendance_record_window, width=500, height=300)

        # Define fields for editing
        self.fields = [
            ("Class Date:", StringVar(value=class_date)),
            ("Site:", StringVar(value=site)),
            ("Hours:", StringVar(value=hours))
        ]
        self.labels = {}
        self.data_labels = {}
        self.entries = {}
        self.edit_mode = False

        for i, (label_text, string_var) in enumerate(self.fields):
            label = Label(self.attendance_record_window, text=label_text, font=self.default_font)
            data_label = Label(self.attendance_record_window, textvariable=string_var, font=self.default_font)
            entry = Entry(self.attendance_record_window, textvariable=string_var, font=self.default_font)

            label.grid(row=i, column=0, padx=5, pady=10, sticky='nsew')
            data_label.grid(row=i, column=1, padx=5, pady=10, sticky='nsew')
            entry.grid(row=i, column=2, padx=5, pady=10, sticky='nsew')
            entry.grid_remove()  # Hide the entry widget initially

            self.labels[label_text] = label
            self.data_labels[label_text] = data_label
            self.entries[label_text] = entry

        # Add Edit, Save, and Delete buttons
        self.edit_button = Button(self.attendance_record_window, text="Edit", command=self.enable_editing, font=self.default_font)
        self.edit_button.grid(row=len(self.fields), column=0, padx=10, pady=10, sticky='w')

        self.save_button = Button(self.attendance_record_window, text="Save", command=self.save_changes, font=self.default_font)
        self.save_button.grid(row=len(self.fields), column=1, padx=10, pady=10, sticky='w')
        self.save_button.grid_remove()  # Hide the Save button initially

        self.delete_button = Button(self.attendance_record_window, text="Delete", command=self.delete_attendance, font=self.default_font)
        self.delete_button.grid(row=len(self.fields), column=2, padx=10, pady=10, sticky='e')

        # Configure the grid to expand properly
        for i in range(1, 4):
            self.attendance_record_window.grid_rowconfigure(i, weight=1)
        self.attendance_record_window.grid_columnconfigure(0, weight=1)
        self.attendance_record_window.grid_columnconfigure(1, weight=1)
        self.attendance_record_window.grid_columnconfigure(2, weight=1)

    def enable_editing(self):
        for label_text in self.labels.keys():
            self.data_labels[label_text].grid_remove()
            self.entries[label_text].grid()
        self.edit_button.config(state='disabled')
        self.save_button.config(state='normal')
        self.edit_mode = True
        self.save_button.grid()  # Show the Save button

    def save_changes(self):
        if self.edit_mode:
            updated_data = {label: var.get() for label, var in self.fields}
            self.update_database(self.selected_attendance_pk, updated_data)

            # Refresh the treeview in Attendance tab
            self.populate_treeview(self.selected_student_id)
            
            # Refresh the StudentsTab treeview
            self.refresh_students_tab()

            self.attendance_record_window.destroy()

    def update_database(self, attendance_pk, updated_data):
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        try:
            query = '''
                UPDATE Attendance
                SET class_date = ?, site = ?, class_hours = ?
                WHERE attendance_pk = ?
            '''
            cursor.execute(query, (
                format_date(updated_data['Class Date:'], out=True), 
                updated_data['Site:'], 
                updated_data['Hours:'], 
                attendance_pk
            ))
            conn.commit()
            messagebox.showinfo("Success", "Attendance record has been updated successfully.")
                    
            # Refresh the StudentsTab treeview
            self.refresh_students_tab()
            
            # Refresh the treeview in Attendance tab
            self.populate_treeview(self.selected_student_id)
            
        except Exception as e:
            conn.rollback()
            messagebox.showerror("Error", f"An error occurred while updating the attendance record: {e}")
        finally:
            conn.close()

    def delete_attendance(self):
        confirm_delete = messagebox.askyesno("Delete Attendance", "Are you sure you want to delete this attendance record? This action cannot be undone.")
        if not confirm_delete:
            return
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        try:
            # Delete from Attendance table using attendance_pk
            cursor.execute('DELETE FROM Attendance WHERE attendance_pk = ?', (self.selected_attendance_pk,))
            conn.commit()
            messagebox.showinfo("Success", "Attendance record has been deleted successfully.")
            # Update the treeview in StudentAttendanceTab
            self.populate_treeview(self.selected_student_id)
            # Refresh the StudentsTab treeview
            self.refresh_students_tab()
        except Exception as e:
            conn.rollback()
            messagebox.showerror("Error", f"An error occurred while deleting the attendance record: {e}")
        finally:
            conn.close()
        self.attendance_record_window.destroy()

class StudentILPTab(ttk.Frame):
    def __init__(self, parent, student_id, default_font):
        super().__init__(parent)
        self.default_font = default_font
        self.student_id = student_id
        self.edit_mode = False  # Initialize edit mode as False
        self.other_text = None  # Initialize the other_text widget
        self.create_widgets()
        self.populate_checkboxes()  # Populate the checkboxes based on the database values

    def create_widgets(self):
        self.fields = [
            "long_goal_job", "long_goal_college", "long_goal_citizenship", 
            "short_goal_speaking", "short_goal_listening", "short_goal_reading", 
            "short_goal_writing", "progress_meet", "progress_increase_one_level", 
            "progress_pass_citizenship", "progress_job", "how_attend", "how_homework", 
            "how_practice_at_home", "how_job_specialist", "problems_childcare", 
            "problems_transportation", "problems_time", "problems_difficult", 
            "problems_location", "problems_moving", "problems_other", "time_less_six", 
            "time_six_to_twelve", "time_more_twelve", "modality_visual", "modality_verbal", 
            "modality_written", "modality_auditory", "modality_kinesthetic"
        ]

        self.string_vars = {field: StringVar(value="0") for field in self.fields}  # Default all checkboxes to unchecked
        self.other_problems_var = StringVar()

        # Create Edit and Save buttons
        self.edit_button = Button(self, text="Edit", command=self.enable_editing, font=self.default_font)
        self.edit_button.grid(row=len(self.fields), column=0, padx=10, pady=10, sticky='w')

        self.save_button = Button(self, text="Save", command=self.save_changes, font=self.default_font)
        self.save_button.grid(row=len(self.fields), column=1, padx=10, pady=10, sticky='w')
        self.save_button.grid_remove()  # Hide the Save button initially

        # Create Panedwindows to separate the checkboxes 
        self.ilp_goal_pane_left = ttk.Panedwindow(self, orient=VERTICAL)
        self.ilp_goal_pane_right = ttk.Panedwindow(self, orient=VERTICAL)

        # Create a bold font
        bold_font = tkFont.Font(weight="bold")

        # Create dictionaries for labels, data_labels, and entries
        self.labels = {}
        self.data_labels = {}
        self.entries = {}

        # Define the style for the bold Labelframes
        style = ttk.Style()
        style.configure("Bold.TLabelframe.Label", font=bold_font)

        # Helper function to create a scrollable Labelframe
        def create_scrollable_labelframe(parent, text):
            container = ttk.Labelframe(parent, text=text, style="Bold.TLabelframe", width=100, height=100)

            # Create a Canvas widget inside the Labelframe
            canvas = Canvas(container)
            scrollbar = ttk.Scrollbar(container, orient=VERTICAL, command=canvas.yview)

            # Create a frame inside the Canvas to hold the widgets
            frame = ttk.Frame(canvas)

            # Configure the canvas
            canvas.configure(yscrollcommand=scrollbar.set)

            # Place the canvas and scrollbar in the Labelframe
            canvas.grid(row=0, column=0, sticky="nsew")
            scrollbar.grid(row=0, column=1, sticky="ns")

            # Add the frame to the canvas
            canvas.create_window((0, 0), window=frame, anchor="nw")

            # Make the frame resize with the canvas
            frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

            # Configure grid in Labelframe to make the canvas expandable
            container.grid_rowconfigure(0, weight=1)
            container.grid_columnconfigure(0, weight=1)

            return container, frame

        # Create scrollable Labelframes
        long_term_goals, long_term_goals_frame = create_scrollable_labelframe(self.ilp_goal_pane_left, 'Long-Term Goals')
        short_term_goals, short_term_goals_frame = create_scrollable_labelframe(self.ilp_goal_pane_left, 'Short-Term Goals')
        progress, progress_frame = create_scrollable_labelframe(self.ilp_goal_pane_left, 'Progress')
        how, how_frame = create_scrollable_labelframe(self.ilp_goal_pane_left, "How I'll Succeed")
        problems, problems_frame = create_scrollable_labelframe(self.ilp_goal_pane_right, 'Problems')
        time, time_frame = create_scrollable_labelframe(self.ilp_goal_pane_right, 'Time')
        modality, modality_frame = create_scrollable_labelframe(self.ilp_goal_pane_right, 'Learning Modalities')
        notes, notes_frame = create_scrollable_labelframe(self.ilp_goal_pane_right, 'Notes')

        # Add Labelframes to Panedwindows
        for labelframe in (long_term_goals, short_term_goals, progress, how):
            self.ilp_goal_pane_left.add(labelframe, weight=1)
        for labelframe in (problems, time, modality, notes):
            self.ilp_goal_pane_right.add(labelframe, weight=1)

        # Create Checkbuttons with the given variables and text
        checkbuttons = [
            (long_term_goals_frame, "Get a job or better job", self.string_vars["long_goal_job"]),
            (long_term_goals_frame, "Enter college/training program", self.string_vars["long_goal_college"]),
            (long_term_goals_frame, "Get my USA citizenship", self.string_vars["long_goal_citizenship"]),
            (short_term_goals_frame, "Increase my speaking skills", self.string_vars["short_goal_speaking"]),
            (short_term_goals_frame, "Increase my listening skills", self.string_vars["short_goal_listening"]),
            (short_term_goals_frame, "Increase my reading skills", self.string_vars["short_goal_reading"]),
            (short_term_goals_frame, "Increase my writing skills", self.string_vars["short_goal_writing"]),
            (progress_frame, "Meet with my teacher every 90 days", self.string_vars["progress_meet"]),
            (progress_frame, "Increase my English by 1 level with a progress test", self.string_vars["progress_increase_one_level"]),
            (progress_frame, "Pass my USA citizenship test", self.string_vars["progress_pass_citizenship"]),
            (progress_frame, "Get a job or enter college/training program", self.string_vars["progress_job"]),
            (how_frame, "Attend every class", self.string_vars["how_attend"]),
            (how_frame, "Complete my homework", self.string_vars["how_homework"]),
            (how_frame, "Practice English at home or line for 20-30 minutes 3x/week", self.string_vars["how_practice_at_home"]),
            (how_frame, "Meet with a job/career specialist", self.string_vars["how_job_specialist"]),
            (problems_frame, "Childcare", self.string_vars["problems_childcare"]),
            (problems_frame, "Transportation", self.string_vars["problems_transportation"]),
            (problems_frame, "Not enough time", self.string_vars["problems_time"]),
            (problems_frame, "Too difficult", self.string_vars["problems_difficult"]),
            (problems_frame, "A different location", self.string_vars["problems_location"]),
            (problems_frame, "Moving", self.string_vars["problems_moving"]),
            (time_frame, "Fewer than 6 months", self.string_vars["time_less_six"]),
            (time_frame, "6-12 months", self.string_vars["time_six_to_twelve"]),
            (time_frame, "More than 1 year", self.string_vars["time_more_twelve"]),
            (modality_frame, "Visual", self.string_vars["modality_visual"]),
            (modality_frame, "Verbal", self.string_vars["modality_verbal"]),
            (modality_frame, "Written", self.string_vars["modality_written"]),
            (modality_frame, "Auditory", self.string_vars["modality_auditory"]),
            (modality_frame, "Kinesthetic (Hands-On)", self.string_vars["modality_kinesthetic"])
        ]

        for i, (frame, text, var) in enumerate(checkbuttons):
            cb = ttk.Checkbutton(frame, text=text, variable=var, style="Custom.TCheckbutton")
            cb.grid(row=i*2, column=0, padx=10, pady=5, sticky='nsew')
            cb.state(['disabled'])

            # Store checkbutton in labels dictionary for editing
            self.labels[text] = cb

        # Add text field for "Other" problems
        other_label = Label(problems_frame, text="Other problems:")
        other_label.grid(row=len(checkbuttons)*2, column=0, padx=10, pady=5, sticky='nw')
        self.other_problems_text = Entry(problems_frame, textvariable=self.other_problems_var, width=40, state="disabled")
        self.other_problems_text.grid(row=len(checkbuttons)*2+1, column=0, padx=10, pady=5, sticky='nsew')
        self.other_problems_text.configure(state='readonly')

        # Add text widget to Notes Labelframe
        self.notes_text = Text(notes_frame, width=30, height=5)
        self.notes_text.grid(column=1, row=13, padx=10, pady=10, sticky='nsew')
        self.notes_text.configure(state='disabled')

        # Make the Panedwindow visible
        self.ilp_goal_pane_left.grid(row=0, column=0, padx=10, pady=10, sticky='nsew')
        self.ilp_goal_pane_right.grid(row=0, column=1, padx=10, pady=10, sticky='nsew')

        # Configure the grid to expand properly
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)

        # Ensure the PanedWindow and its components expand
        for pane in [self.ilp_goal_pane_left, self.ilp_goal_pane_right]:
            pane.grid_rowconfigure(0, weight=1)
            pane.grid_columnconfigure(0, weight=1)


    def populate_checkboxes(self):
        # Connect to the SQLite database
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        query = '''
            SELECT 
                goal_job, 
                goal_college, 
                goal_citizenship, 
                goal_speaking, 
                goal_listening, 
                goal_reading, 
                goal_writing,
                progress_meet_teacher, 
                progress_increase_one_level, 
                progress_citizenship, 
                progress_get_job,
                how_attend, 
                how_homework, 
                how_practice_home, 
                how_meet_job_specialist,
                problem_childcare, 
                problem_transport, 
                problem_time, 
                problem_difficult, 
                problem_location, 
                problem_moving, 
                problem_other,
                time_less_six, 
                time_six_to_twelve, 
                time_more_twelve,
                modality_visual, 
                modality_verbal, 
                modality_written, 
                modality_auditory, 
                modality_kinesthetic, 
                notes
            FROM 
                ILP 
            WHERE 
                student_id = ?
        '''
        # Retrieve the ILP data for the selected student
        cursor.execute(query, (self.student_id,))
        
        ilp_data = cursor.fetchone()
        conn.close()

        if ilp_data:
            # Set the checkboxes based on the data
            for i, field in enumerate(self.fields):
                self.string_vars[field].set(str(ilp_data[i]))  # Update each checkbox state

            # Set the "Other" problems text field
            self.other_problems_var.set(ilp_data[21])  # Assuming the "problem_other" column is at index 21

            # Set notes text if applicable
            notes_index = len(self.fields)  # Assuming notes is the last field
            self.notes_text.insert(END, ilp_data[notes_index])

    def toggle_text_widget(self, var, frame, row):
        if var.get() == "1":
            if self.other_text is None:  # Create the text widget only if it doesn't exist
                self.other_text = Text(frame, width=30, height=5)
                self.other_text.grid(row=row+1, column=0, padx=10, pady=5, sticky='nsew')
                frame.grid_rowconfigure(row+1, weight=1)  # Make the new row expandable
        else:
            if self.other_text is not None:
                self.other_text.grid_forget()
                self.other_text = None
                frame.grid_rowconfigure(row+1, weight=0)  # Reset the row weight when the widget is removed

    def enable_editing(self):
        self.edit_mode = True
        self.edit_button.config(state='disabled')
        self.save_button.config(state='normal')
        self.save_button.grid()  # Show the Save button

        # Enable editing for checkboxes
        for checkbox in self.labels.values():
            checkbox.state(['!disabled'])  # Enable the checkboxes

        # Enable editing for text widgets
        self.other_problems_text.configure(state='normal')
        self.notes_text.configure(state='normal')

    def save_changes(self):
        # Check if the user wants to save changes
        if messagebox.askyesno("Save Changes", "Do you want to save the changes?"):
            # Gather the current checkbox selections
            updated_data = {field: var.get() for field, var in self.string_vars.items()}
            updated_data['other_problems'] = self.other_problems_var.get()  # Include other problems text

            # Update the database with the new ILP data
            self.update_database(updated_data)

            # Disable editing after saving
            self.edit_mode = False
            for checkbox in self.labels.values():
                checkbox.state(['disabled'])  # Disable the checkboxes
            self.save_button.grid_remove()  # Hide the save button

            # Disable text widgets after saving
            self.other_problems_text.configure(state='readonly')
            self.notes_text.configure(state='disabled')

            self.edit_button.config(state='normal')  # Re-enable the edit button


    def update_database(self, updated_data):
        # Connect to the SQLite database
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        try:
            query = '''
                UPDATE ILP SET 
                    goal_job = ?,
                    goal_college = ?,
                    goal_citizenship = ?,
                    goal_speaking = ?,
                    goal_listening = ?,
                    goal_reading = ?,
                    goal_writing = ?,
                    progress_meet_teacher = ?,
                    progress_increase_one_level = ?,
                    progress_citizenship = ?,
                    progress_get_job = ?,
                    how_attend = ?,
                    how_homework = ?,
                    how_practice_home = ?,
                    how_meet_job_specialist = ?,
                    problem_childcare = ?,
                    problem_transport = ?,
                    problem_time = ?,
                    problem_difficult = ?,
                    problem_location = ?,
                    problem_moving = ?,
                    problem_other = ?,
                    time_less_six = ?,
                    time_six_to_twelve = ?,
                    time_more_twelve = ?,
                    modality_visual = ?,
                    modality_verbal = ?,
                    modality_written = ?,
                    modality_auditory = ?,
                    modality_kinesthetic = ?,
                    notes = ?
                WHERE 
                    student_id = ?
            '''
            # Prepare the update query for the ILP table
            cursor.execute(query, (
                updated_data['long_goal_job'],
                updated_data['long_goal_college'],
                updated_data['long_goal_citizenship'],
                updated_data['short_goal_speaking'],
                updated_data['short_goal_listening'],
                updated_data['short_goal_reading'],
                updated_data['short_goal_writing'],
                updated_data['progress_meet'],
                updated_data['progress_increase_one_level'],
                updated_data['progress_pass_citizenship'],
                updated_data['progress_job'],
                updated_data['how_attend'],
                updated_data['how_homework'],
                updated_data['how_practice_at_home'],
                updated_data['how_job_specialist'],
                updated_data['problems_childcare'],
                updated_data['problems_transportation'],
                updated_data['problems_time'],
                updated_data['problems_difficult'],
                updated_data['problems_location'],
                updated_data['problems_moving'],
                updated_data['other_problems'],
                updated_data['time_less_six'],
                updated_data['time_six_to_twelve'],
                updated_data['time_more_twelve'],
                updated_data['modality_visual'],
                updated_data['modality_verbal'],
                updated_data['modality_written'],
                updated_data['modality_auditory'],
                updated_data['modality_kinesthetic'],
                self.notes_text.get("1.0", END).strip(),
                self.student_id
            ))

            conn.commit()
            messagebox.showinfo("Success", "ILP record has been updated successfully.")

        except Exception as e:
            conn.rollback()
            messagebox.showerror("Error", f"An error occurred while updating the ILP record: {e}")

        finally:
            conn.close()
            
class StudentMeetingTab(ttk.Frame):
    def __init__(self, parent, students_tab, selected_student_id, default_font):
        super().__init__(parent)
        self.default_font = default_font
        self.students_tab = students_tab
        self.selected_student_id = selected_student_id
        self.columns = ['Meeting Date', 'Site', 'Notes']
        self.visible_columns = self.columns[:]
        self.filter_text = StringVar()  
        self.filter_column = StringVar()  
        self._sort_column = None
        self._sort_reverse = False
        self.string_vars = {}
        self.create_widgets()
        self.populate_treeview(selected_student_id)

    def refresh_students_tab(self):
        if self.students_tab:
            self.students_tab.populate_treeview()

    def create_widgets(self):
        # Add Filter button
        self.filter_button = Button(self, text="Filter", command=self.open_filter_window, font=self.default_font)
        self.filter_button.grid(row=0, column=0, padx=5, pady=5, sticky='w')

        # Add Meeting Record Treeview
        self.tree = ttk.Treeview(self, columns=self.columns, show='headings')
        for col in self.columns:
            self.tree.heading(col, text=col, command=lambda c=col: self.sort_by_column(c, False))
            # Adjust the column width for "Site" and others
            if col == "Site":
                self.tree.column(col, width=20, anchor='center')  # Set a fixed width for "Site"
            else:
                self.tree.column(col, anchor='center', stretch=True)  # Allow other columns to stretch
                
        self.tree.grid(row=1, column=0, columnspan=3, padx=10, pady=10, sticky='nsew')

        # Add vertical scrollbar
        self.vscrollbar = ttk.Scrollbar(self, orient=VERTICAL, command=self.tree.yview)
        self.vscrollbar.grid(column=3, row=1, sticky='ns')  # Note: Adjusted grid placement
        self.tree['yscrollcommand'] = self.vscrollbar.set

        # Create Add Meeting Record button
        self.add_meeting_button = Button(self, text="Add Meeting", command=self.add_meeting, font=self.default_font)
        self.add_meeting_button.grid(row=2, column=0, padx=5, pady=5, sticky='nw')

        # Configure grid to expand properly
        self.grid_rowconfigure(1, weight=1)
        self.grid_rowconfigure(2, weight=0)
        self.grid_columnconfigure(0, weight=1)  # Allow the first column to expand
        self.grid_columnconfigure(1, weight=0)  # Weight for the scrollbar column


        # Bind double click to treeview items (to view meeting details)
        self.tree.bind("<Double-1>", self.on_treeview_click)

    def populate_treeview(self, selected_student_id, filtered_rows=None):
        # Clear existing rows
        for row in self.tree.get_children():
            self.tree.delete(row)

        # Connect to the SQLite database
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        if filtered_rows is None:
            query = '''
                SELECT m.meeting_pk, m.meeting_date, m.meeting_site, m.meeting_notes
                FROM Meetings m
                WHERE m.student_id = ?
                ORDER BY m.meeting_date DESC
            '''
            cursor.execute(query, (selected_student_id,))
            rows = cursor.fetchall()
        else:
            rows = filtered_rows

        processed_rows = [(row[0], format_date(row[1]), row[2], row[3]) for row in rows]

        # Create pk map for editing row
        self.meeting_pk_map = {}

        # Insert data into the Treeview
        for row in processed_rows:
            item_id = self.tree.insert('', 'end', values=(row[1], row[2], row[3].replace("\n", " ")))
            self.meeting_pk_map[item_id] = row[0]

        # Close the database connection
        conn.close()

        # Adjust column widths based on content
        for col in self.tree['columns']:
            self.tree.column(col, width=tkFont.Font().measure(col))  # Start with the header width
            for item in self.tree.get_children():
                cell_value = self.tree.item(item, 'values')[self.tree['columns'].index(col)]
                cell_width = tkFont.Font().measure(cell_value) + 100
                if self.tree.column(col, width=None) < cell_width:
                    self.tree.column(col, width=cell_width)

    def sort_by_column(self, col, reverse):
        # Fetch data from Treeview
        data = [(self.tree.set(child, col), child) for child in self.tree.get_children('')]

        # Convert data to appropriate types for sorting
        if col == 'Meeting Date':
            data = [(datetime.strptime(d, '%m-%d-%Y') if d else datetime.min, child) for d, child in data]
        else:
            data = [(d, child) for d, child in data]

        # Sort data
        data.sort(reverse=reverse)

        # Rearrange items in sorted positions
        for index, (val, child) in enumerate(data):
            self.tree.move(child, '', index)

        # Reverse sort next time
        self._sort_column = col
        self._sort_reverse = not reverse

        # Update heading with new sort direction
        self.tree.heading(col, command=lambda: self.sort_by_column(col, not reverse))

    def apply_filter(self):
        column = self.filter_column.get()
        filter_text = self.filter_text.get()
        filtered_rows = []

        if column and filter_text:
            # Connect to the SQLite database
            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()

            # Fetch meeting data for the selected student
            query = '''
                SELECT m.meeting_date, m.meeting_site, m.meeting_notes
                FROM Meetings m
                WHERE m.student_id = ?
            '''
            cursor.execute(query, (self.selected_student_id,))
            rows = cursor.fetchall()
            conn.close()

            processed_rows = [(format_date(row[0]), row[1], row[2]) for row in rows]

            # Apply the filter to the rows
            column_index = self.columns.index(column) if column in self.columns else -1
            for row in processed_rows:
                cell_value = str(row[column_index]).lower()
                if filter_text.lower() in cell_value:
                    filtered_rows.append(row)

            # Insert filtered data into the Treeview
            self.tree.delete(*self.tree.get_children())
            for row in filtered_rows:
                self.tree.insert('', 'end', values=row)

    def reset_filters(self):
        self.filter_text.set('')
        self.filter_column.set('')
        self.populate_treeview(self.selected_student_id)

    def open_filter_window(self):
        self.filter_window = Toplevel(self)
        self.filter_window.title("Filter Records")

        Label(self.filter_window, text="Select Column:", font=self.default_font).grid(row=0, column=0, padx=10, pady=5, sticky='w')
        column_menu = AutocompleteCombobox(self.filter_window, textvariable=self.filter_column, font=self.default_font)
        column_menu['values'] = self.columns  # Populate with column names
        column_menu.grid(row=0, column=1, padx=10, pady=5, sticky='w')

        Label(self.filter_window, text="Enter Filter Text:", font=self.default_font).grid(row=1, column=0, padx=10, pady=5, sticky='w')
        filter_entry = Entry(self.filter_window, textvariable=self.filter_text, font=self.default_font)
        filter_entry.grid(row=1, column=1, padx=10, pady=5, sticky='w')

        Button(self.filter_window, text="Apply Filter", font=self.default_font, command=self.apply_filter).grid(row=2, column=0, columnspan=2, pady=5)
        Button(self.filter_window, text="Reset", font=self.default_font, command=self.reset_filters).grid(row=3, column=0, columnspan=2, pady=5)

    def add_meeting(self):
        # Create the Add Meeting window and center it
        self.add_meeting_window = Toplevel(self.master)
        self.add_meeting_window.withdraw()  # Hide the window initially
        self.add_meeting_window.title("Add Meeting")
        self.add_meeting_window.grid_rowconfigure(0, weight=1)
        self.add_meeting_window.grid_columnconfigure(0, weight=1)
        center_window(self.add_meeting_window, 675, 750)

        # Define the meeting record label and entry fields
        self.fields = [
            ("Meeting Date:", "meeting_date"),
            ("Site:", "site"),
            ("Notes:", "notes")
        ]

        # Create Save button
        self.save_meeting_button = Button(self.add_meeting_window, font=self.default_font, text="Save", command=self.save_meeting, state='disabled')
        self.save_meeting_button.grid(row=len(self.fields)+2, column=0, columnspan=2, pady=10)

        # Method to check if all fields are filled
        def check_fields():
            notes_text = self.notes_text.get("1.0", "end-1c").strip()
            # Ensure all fields except notes are filled, and notes_text is not empty
            if all(var.get().strip() for key, var in self.string_vars.items() if key != 'notes') and notes_text:
                self.save_meeting_button.config(state='normal')
            else:
                self.save_meeting_button.config(state='disabled')

        # Create labels and entries
        for i, (label_text, var_name) in enumerate(self.fields):
            i += 1
            label = Label(self.add_meeting_window, text=label_text, font=self.default_font)
            label.grid(row=i, column=0, padx=10, pady=5, sticky='n')

            if label_text == "Notes:":
                # Use Text widget for multiline input
                self.notes_text = Text(self.add_meeting_window, wrap=WORD, height=10, width=40)
                self.notes_text.grid(row=i, column=1, padx=10, pady=5, sticky='nsew')
                self.notes_text.bind('<KeyRelease>', lambda event: check_fields())  # Check fields on input
            else:
                string_var = StringVar()
                self.string_vars[var_name] = string_var  # Store the StringVar object in the dictionary
                entry = Entry(self.add_meeting_window, textvariable=string_var, font=self.default_font)
                entry.grid(row=i, column=1, padx=10, pady=5, sticky='n')
                if label_text == "Meeting Date:":
                    # Create the calendar widget
                    large_font = font.Font(family="Helvetica", size=14, weight="bold")
                    self.cal = Calendar(self.add_meeting_window, selectmode='day',
                                        year=datetime.today().year, month=datetime.today().month,
                                        day=datetime.today().day, font=large_font,
                                        headersbackground="lightblue",
                                        background="lightgrey", foreground="black",
                                        selectbackground="blue", selectforeground="white")
                    self.cal.grid(row=0, column=1, columnspan=2, padx=10, pady=5)
                    # Bind the calendar selection event to update the "Meeting Date:" entry box
                    self.cal.bind("<<CalendarSelected>>", lambda event, var=string_var: var.set(self.cal.get_date()))
                elif label_text == "Site:":
                    # Create a combobox for the Site field
                    site_names = get_site_names()
                    combobox = AutocompleteCombobox(self.add_meeting_window, textvariable=string_var, values=site_names, font=self.default_font)
                    combobox.grid(row=i, column=1, padx=10, pady=5, sticky='n')

                # Add the trace after the widget creation to ensure button initialization is done
                string_var.trace('w', lambda name, index, mode, var=string_var: check_fields())

        # Configure the grid to expand properly
        for i in range(1, len(self.fields)+1):
            self.add_meeting_window.grid_rowconfigure(i, weight=1)
            self.add_meeting_window.grid_columnconfigure(0, weight=1)

        # Configure the grid to expand properly
        for i in range(1, len(self.fields)+2):
            self.add_meeting_window.grid_rowconfigure(i, weight=1)
        self.add_meeting_window.grid_columnconfigure(0, weight=1)
        self.add_meeting_window.grid_columnconfigure(1, weight=1)

        # Now that the window is fully set up, show it
        self.add_meeting_window.deiconify()

    def update_notes_text(self, string_var):
        content = self.notes_text.get("1.0", "end-1c")  # Get text content
        string_var.set(content.strip())  # Update StringVar with stripped content

    def save_meeting(self):
        msg = "Save meeting record?"
        if messagebox.askyesno("Add Meeting", msg):
            # Get the values from the entry fields
            meeting_date = self.string_vars['meeting_date'].get()
            site = self.string_vars['site'].get()
            notes = self.notes_text.get("1.0", "end-1c").strip()  # Get text from Text widget directly

            # Connect to the SQLite database
            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()

            # Execute the insert query
            query = '''
                INSERT INTO Meetings (student_id, meeting_date, meeting_site, meeting_notes)
                VALUES (?, ?, ?, ?)
            '''
            cursor.execute(query, (
                self.selected_student_id, 
                format_date(meeting_date, out=True), 
                site, 
                notes
            ))

            # Commit the changes and close the connection
            conn.commit()
            conn.close()

            # Close the Add Meeting window
            self.add_meeting_window.destroy()

            # Refresh the meeting Treeview
            self.populate_treeview(self.selected_student_id)

            # Refresh the StudentsTab treeview
            self.refresh_students_tab()
    
    def on_treeview_click(self, event):
        region = self.tree.identify('region', event.x, event.y)
        if region != 'heading':
            self.on_double_click(event)

    def on_double_click(self, event):
        curItem = self.tree.focus()
        item_values = self.tree.item(curItem)['values']
        self.selected_meeting_pk = self.meeting_pk_map.get(curItem)  # Retrieve attendance_pk using the item_id
        meeting_date = item_values[0]
        site = item_values[1]
        # Fetch the notes directly from the database using the meeting_date and site as keys
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        # Fetch the meeting record to get the full notes text
        query = '''
            SELECT meeting_notes
            FROM Meetings
            WHERE meeting_pk = ? 
        '''
        cursor.execute(query, (self.selected_meeting_pk,))
        notes_row = cursor.fetchone()
        conn.close()

        # Check if the notes were retrieved successfully
        notes = notes_row[0] if notes_row else ""

        title_string = f"Record: {meeting_date} @ {site}"

        self.meeting_record_window = Toplevel(self)
        self.meeting_record_window.withdraw()  # Hide the window initially
        self.meeting_record_window.title(title_string)
        center_window(self.meeting_record_window, width=800, height=500)

        # Define fields for editing
        self.fields = [
            ("Meeting Date:", StringVar(value=meeting_date)),
            ("Site:", StringVar(value=site)),
        ]
        self.labels = {}
        self.data_labels = {}
        self.entries = {}
        self.edit_mode = False

        for i, (label_text, string_var) in enumerate(self.fields):
            label = Label(self.meeting_record_window, text=label_text, font=self.default_font)
            data_label = Label(self.meeting_record_window, textvariable=string_var, font=self.default_font)
            entry = Entry(self.meeting_record_window, textvariable=string_var, font=self.default_font)

            label.grid(row=i, column=0, padx=5, pady=10, sticky='nsew')
            data_label.grid(row=i, column=1, padx=5, pady=10, sticky='nsew')
            entry.grid(row=i, column=2, padx=5, pady=10, sticky='nsew')
            entry.grid_remove()  # Hide the entry widget initially

            self.labels[label_text] = label
            self.data_labels[label_text] = data_label
            self.entries[label_text] = entry

        # Add a Text widget for notes with line wrap
        notes_label = Label(self.meeting_record_window, text="Notes:", font=self.default_font)
        notes_label.grid(row=len(self.fields), column=0, padx=5, pady=10, sticky='nsew')

        self.notes_text = Text(self.meeting_record_window, wrap=WORD, height=10, width=50)
        self.notes_text.grid(row=len(self.fields), column=1, columnspan=2, padx=5, pady=10, sticky='nsew')

        # Insert existing notes with line breaks
        self.notes_text.insert(END, notes)

        # Add Edit, Save, and Delete buttons
        self.edit_button = Button(self.meeting_record_window, text="Edit", command=self.enable_editing, font=self.default_font)
        self.edit_button.grid(row=len(self.fields) + 1, column=0, padx=10, pady=10, sticky='w')

        self.save_button = Button(self.meeting_record_window, text="Save", command=self.save_changes, font=self.default_font)
        self.save_button.grid(row=len(self.fields) + 1, column=1, padx=10, pady=10, sticky='w')
        self.save_button.grid_remove()  # Hide the Save button initially

        self.delete_button = Button(self.meeting_record_window, text="Delete", command=self.delete_meeting(self.selected_meeting_pk), font=self.default_font)
        self.delete_button.grid(row=len(self.fields) + 1, column=2, padx=10, pady=10, sticky='e')

        # Configure the grid to expand properly
        for i in range(len(self.fields) + 1):
            self.meeting_record_window.grid_rowconfigure(i, weight=1)
        self.meeting_record_window.grid_columnconfigure(0, weight=1)
        self.meeting_record_window.grid_columnconfigure(1, weight=1)
        self.meeting_record_window.grid_columnconfigure(2, weight=1)

        # Now that the window is fully set up, show it
        self.meeting_record_window.deiconify()

    def enable_editing(self):
        for label_text in self.labels.keys():
            self.data_labels[label_text].grid_remove()
            self.entries[label_text].grid()
        
        self.notes_text.config(state='normal')  # Enable editing for the notes field
        self.edit_button.config(state='disabled')
        self.save_button.config(state='normal')
        self.edit_mode = True
        self.save_button.grid()  # Show the Save button

    def save_changes(self):
        if self.edit_mode:
            old_data = {label: self.tree.item(self.tree.focus())['values'][i] for i, (label, _) in enumerate(self.fields)}
            old_data['Notes:'] = self.notes_text.get('1.0', END).strip()  # Capture notes separately
            updated_data = {label: var.get() for label, var in self.fields}
            updated_data['Notes:'] = self.notes_text.get("1.0", "end-1c")  # Get updated notes
            self.update_database(self.selected_meeting_pk, updated_data)

            # Refresh the treeview in Meeting tab
            self.populate_treeview(self.selected_student_id)

            # Refresh the StudentsTab treeview
            self.refresh_students_tab()

            self.meeting_record_window.destroy()

    def update_database(self, meeting_pk, updated_data):
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        try:
            query = '''
                UPDATE 
                    Meetings
                SET 
                    meeting_date = ?, 
                    meeting_site = ?, 
                    meeting_notes = ?
                WHERE 
                    meeting_pk = ? 
            '''
            cursor.execute(query, (
                format_date(updated_data['Meeting Date:'], out=True), 
                updated_data['Site:'], 
                updated_data['Notes:'], 
                meeting_pk
            ))
            conn.commit()
            
            messagebox.showinfo("Success", "Meeting record has been updated successfully.")
                    
            # Refresh the StudentsTab treeview
            self.refresh_students_tab()
            
            # Refresh the treeview in Meeting tab
            self.populate_treeview(self.selected_student_id)
            
        except Exception as e:
            conn.rollback()
            messagebox.showerror("Error", f"An error occurred while updating the meeting record: {e}")
        finally:
            conn.close()

    def delete_meeting(self, meeting_pk):
        # Confirm deletion
        confirm_delete = messagebox.askyesno("Delete Meeting", "Are you sure you want to delete this meeting record? This action cannot be undone.")
        if not confirm_delete:
            return
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        try:
            # Delete from Meetings table
            cursor.execute('DELETE FROM Meetings WHERE meeting_pk = ?', (meeting_pk,))
            conn.commit()
            messagebox.showinfo("Success", "Meeting record has been deleted successfully.")
            # Update the treeview in StudentMeetingTab
            self.populate_treeview(self.selected_student_id)
            # Refresh the StudentsTab treeview
            self.refresh_students_tab()
        except Exception as e:
            conn.rollback()
            messagebox.showerror("Error", f"An error occurred while deleting the meeting record: {e}")
        finally:
            conn.close()
        self.meeting_record_window.destroy()

class StudentTestingTab(ttk.Frame):
    def __init__(self, parent, default_font):
        super().__init__(parent)
        self.default_font = default_font
        self.create_widgets()

    def create_widgets(self):
        pass

class AttendanceTab(ttk.Frame):
    def __init__(self, parent, default_font):
        super().__init__(parent)
        self.default_font = default_font
        self.create_widgets()

    def create_widgets(self):  
        # Create a Notebook
        self.notebook = ttk.Notebook(self)

        # Create and add tabs
        self.all_tab = AttendanceAllTab(self.notebook, self.default_font)
        # self.site_tab = AttendanceSiteTab(self.notebook)
        self.add_tab = AttendanceAddTab(self.notebook, self.default_font, self.all_tab)

        self.notebook.add(self.all_tab, text='All')
        # self.notebook.add(self.site_tab, text='Site')
        self.notebook.add(self.add_tab, text='Add')

        # Pack the Notebook widget to make it visible
        self.notebook.grid(row=0, column=0, padx=10, pady=10, sticky='nsew')

        # Bind the tab change event
        self.notebook.bind("<<NotebookTabChanged>>", self.on_tab_changed)

        # Configure the master window to expand properly
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

    def on_tab_changed(self, event):
        # Check if the AttendanceAll tab is selected
        selected_tab = self.notebook.tab(self.notebook.select(), "text")
        if selected_tab == 'All':
            self.all_tab.populate_treeview()

class AttendanceAllTab(ttk.Frame):
    def __init__(self, parent, default_font):
        super().__init__(parent)
        self.parent = parent
        self.default_font = default_font
        self.columns = ["Student ID", "Last Name", "First Name", "Site", "Class Date", "Hours"]
        self.sort_directions = {col: False for col in self.columns}
        self.filter_text = StringVar()
        self.filter_column = StringVar()
        self.site_filter = StringVar()
        self.date_filter = StringVar()
        self.total_attendance_records_counter = IntVar()
        self.create_widgets()
        self.populate_treeview()

    def on_treeview_click(self, event):
        region = self.tree.identify('region', event.x, event.y)
        if region != 'heading':
            self.on_double_click(event)

    def create_widgets(self):
        # Labels and dropdowns for filters
        filter_frame = ttk.Frame(self)
        filter_frame.grid(row=0, column=0, columnspan=4, padx=5, pady=(0, 2), sticky='e')

        # Site Filter
        site_label = Label(filter_frame, text="Filter by Site:", font=self.default_font)
        site_label.grid(row=0, column=0, padx=2, pady=1, sticky='e')

        self.site_dropdown = AutocompleteCombobox(filter_frame, textvariable=self.site_filter, width=15, font=self.default_font)
        self.site_dropdown['values'] = self.fetch_unique_values('site')
        self.site_dropdown.bind("<<ComboboxSelected>>", lambda _: self.apply_filters())
        self.site_dropdown.grid(row=0, column=1, padx=2, pady=1, sticky='w')

        # Date Filter
        date_label = Label(filter_frame, text="Filter by Class Date:", font=self.default_font)
        date_label.grid(row=1, column=0, padx=2, pady=1, sticky='e')

        self.date_dropdown = AutocompleteCombobox(filter_frame, textvariable=self.date_filter, width=15, font=self.default_font)
        self.date_dropdown['values'] = self.fetch_unique_values('class_date')
        self.date_dropdown.bind("<<ComboboxSelected>>", lambda _: self.apply_filters())
        self.date_dropdown.grid(row=1, column=1, padx=2, pady=1, sticky='w')

        # Treeview for displaying attendance records
        self.tree = ttk.Treeview(self, columns=self.columns, show='headings')
        for col in self.columns:
            self.tree.heading(col, text=col, command=lambda c=col: self.sort_records(c))
            self.tree.column(col, anchor='center')

        self.tree.grid(row=1, column=0, columnspan=4, padx=5, pady=(1, 2), sticky='nsew')

        # Scrollbars
        vsb = ttk.Scrollbar(self, orient="vertical", command=self.tree.yview)
        vsb.grid(row=1, column=4, sticky='ns')
        self.tree.configure(yscrollcommand=vsb.set)

        # Button frame
        button_frame = ttk.Frame(self)
        button_frame.grid(row=2, column=0, columnspan=4, padx=5, pady=(2, 0), sticky='e')

        # Filter button
        self.filter_button = Button(button_frame, text="Filter", command=self.open_filter_window, font=self.default_font)
        self.filter_button.grid(row=0, column=0, padx=2, pady=1, sticky='e')

        # Clear filter button
        self.clear_filter_button = Button(button_frame, text="Clear Filter", command=self.reset_filters, font=self.default_font)
        self.clear_filter_button.grid(row=0, column=1, padx=2, pady=1, sticky='e')

        # Total Attendance Records label frame
        total_record_frame = ttk.Frame(self)
        total_record_frame.grid(row=2, column=0, columnspan=4, padx=5, pady=(2, 0), sticky='w')

        # Total Attendance Records label
        self.total_attendance_records_prefix = Label(total_record_frame, text="Total Attendance Records:", font=self.default_font)
        self.total_attendance_records_prefix.grid(row=0, column=0, padx=2, pady=1, sticky='w')

        # Total Attendance Records label
        self.total_attendance_records = Label(total_record_frame, textvariable=self.total_attendance_records_counter, font=self.default_font)
        self.total_attendance_records.grid(row=0, column=1, padx=2, pady=1, sticky='w')

        # Configure grid with equal weight for all rows and columns
        self.configure_grid()

        self.tree.bind("<Double-1>", self.on_treeview_click)

    def configure_grid(self):
        total_rows = 3  # Number of rows used in the main grid (including filter_frame, tree, and button_frame)
        total_columns = 5  # Number of columns used in the main grid

        # Set weight for each row
        self.grid_rowconfigure(0, weight=0)  # Filter row
        self.grid_rowconfigure(1, weight=1)  # Treeview row
        self.grid_rowconfigure(2, weight=0)  # Button row

        # Set weight for each column
        for col in range(total_columns):
            self.grid_columnconfigure(col, weight=1)

    def fetch_unique_values(self, column_name):
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        if column_name == 'class_date':
            query = f"SELECT DISTINCT {column_name} FROM Attendance ORDER BY class_date DESC"
        else:
            query = f"SELECT DISTINCT {column_name} FROM Attendance ORDER BY site"
        cursor.execute(query)
        values = [row[0] for row in cursor.fetchall()]
        conn.close()
        return values

    def populate_treeview(self, filtered_rows=None):
        # Clear existing rows
        for row in self.tree.get_children():
            self.tree.delete(row)
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        if filtered_rows is None:
            query = '''
                SELECT
                    a.attendance_pk,
                    s.student_id,
                    s.last_name,
                    s.first_name,
                    a.site,
                    a.class_date,
                    a.class_hours
                FROM
                    Attendance a
                JOIN 
                    Student s 
                ON 
                    a.student_id = s.student_id
                WHERE 
                    strftime('%m', a.class_date) = strftime('%m', CURRENT_DATE)
                    AND strftime('%Y', a.class_date) = strftime('%Y', CURRENT_DATE)
                ORDER BY
                    a.class_date DESC,
                    a.site ASC,
                    s.last_name ASC
            '''
            cursor.execute(query)
            rows = cursor.fetchall()
        else:
            rows = filtered_rows

        # Create attendance_pk map for later editing
        self.attendance_pk_map = {}

        # Insert data into the Treeview
        for row in rows:
            if any(element is None for element in row):
                continue  # Skip this row if any element is None            
            # Reformat date: YYYY-MM-DD -> MM-DD-YYYY (no leading zeros)
            item_id = self.tree.insert('', 'end', values=(row[1], row[2], row[3], row[4], format_date(row[5]), row[6]))
            self.attendance_pk_map[item_id] = row[0]

        # Update record counter for displaying total number of records
        self.total_attendance_records_counter.set(len(rows))

        # Adjust column widths based on content
        for col in self.tree['columns']:
            self.tree.column(col, width=tkFont.Font().measure(col))  # Start with the header width
            for item in self.tree.get_children():
                cell_value = self.tree.item(item, 'values')[self.tree['columns'].index(col)]
                cell_width = tkFont.Font().measure(cell_value) + 100
                if self.tree.column(col, width=None) < cell_width:
                    self.tree.column(col, width=cell_width)        
        
        # Close the database connection
        conn.close()

    def on_tab_changed(self, event):
        if self.parent.select() == self:
            self.populate_treeview()

    def sort_records(self, col):
        # Toggle the sorting direction
        self.sort_directions[col] = not self.sort_directions[col]
        data = [(self.tree.set(child, col), child) for child in self.tree.get_children('')]

        if col == 'Class Date':
            # Handle sorting for date column
            data.sort(key=lambda x: datetime.strptime(x[0], '%m-%d-%Y'), reverse=self.sort_directions[col])
        elif col == 'Hours':
            # Handle sorting for numeric column
            data.sort(key=lambda x: float(x[0]), reverse=self.sort_directions[col])
        elif col == 'Student ID':
            # Handle sorting for numeric column
            data.sort(key=lambda x: int(x[0]), reverse=self.sort_directions[col])
        else:
            # Handle sorting for text columns
            data.sort(key=lambda x: str(x[0]), reverse=self.sort_directions[col])

        for index, (val, child) in enumerate(data):
            self.tree.move(child, '', index)

        # Update heading with new sort direction
        self.tree.heading(col, command=lambda: self.sort_records(col))

    def apply_filter(self):
        column = self.filter_column.get()
        filter_text = self.filter_text.get().lower().strip()
        filtered_rows = []
        
        if column and filter_text:
            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()

            query = '''
                SELECT 
                    a.attendance_pk,
                    s.student_id,  
                    s.last_name,
                    s.first_name, 
                    a.site, 
                    a.class_date, 
                    a.class_hours
                FROM 
                    Attendance a
                JOIN 
                    Student s 
                ON 
                    a.student_id = s.student_id
                ORDER BY
                    a.class_date DESC,
                    a.site ASC,
                    s.last_name ASC
            '''
            cursor.execute(query)
            rows = cursor.fetchall()
            conn.close()

            # Remove leading zeros from the filter text
            filter_text = re.sub(r'\b0+(\d)', r'\1', filter_text)

            # Apply the filter to the rows
            for row in rows:
                # formatted_row = list(row)
                formatted_row = tuple(row)

                # Format dates for filtering
                class_date = formatted_row[5]

                # Convert date to different searchable formats
                def date_matches(date_str, search_text):
                    try:
                        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                        formatted_date = [
                            date_obj.strftime('%Y-%m-%d'),  # Full date
                            date_obj.strftime('%m-%d-%Y'),  # Alternative full date format
                            date_obj.strftime('%y-%m-%d'),  # Short year with full date
                            date_obj.strftime('%m-%d'),     # Month-day only
                            date_obj.strftime('%Y-%m'),     # Month
                            date_obj.strftime('%m'),        # Year and month
                            date_obj.strftime('%Y'),        # Year only
                            date_obj.strftime('%y'),        # Short year only
                            date_obj.strftime('%B'),        # full month name
                            date_obj.strftime('%b'),        # short month name
                            date_obj.strftime('%B %Y'),     # full month name with year
                            date_obj.strftime('%b %Y')      # short month name with yar                        
                        ]

                        # Normalize month and day by removing leading zeros
                        normalized_date = [
                            re.sub(r'\b0+(\d)', r'\1', d) for d in formatted_date
                        ]
                        # Check if any normalized version matches the search text
                        return any(search_text in nd.lower() for nd in normalized_date)
                    except ValueError:
                        return False

                # Check if the filter text matches in any format
                if column == 'Class Date':
                    if date_matches(class_date, filter_text):
                        filtered_rows.append(formatted_row)
                else:
                    col_index = self.columns.index(column) + 1
                    if filter_text in str(formatted_row[col_index]).lower():
                        filtered_rows.append(formatted_row)
        self.filter_window.destroy()
        self.populate_treeview(filtered_rows if filtered_rows else None)

    def apply_filters(self):
        site = self.site_filter.get()
        class_date = self.date_filter.get()
        filtered_rows = []

        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        query = '''
            SELECT
                a.attendance_pk,
                s.student_id,
                s.last_name, 
                s.first_name, 
                a.site, 
                a.class_date, 
                a.class_hours
            FROM 
                Attendance a
            JOIN 
                Student s 
            ON 
                a.student_id = s.student_id
            WHERE 
                (? = '' OR a.site = ?) AND (? = '' OR a.class_date = ?)
            ORDER BY
                a.class_date DESC,
                a.site ASC,
                s.last_name ASC
        '''
        cursor.execute(query, (site, site, class_date, class_date))
        filtered_rows = cursor.fetchall()
        conn.close()

        self.populate_treeview(filtered_rows)

    def reset_filters(self):
        self.filter_text.set('')
        self.filter_column.set('')
        self.populate_treeview()

    def open_filter_window(self):
        self.filter_window = Toplevel(self)
        self.filter_window.title("Filter Records")

        Label(self.filter_window, text="Select Column:", font=self.default_font).grid(row=0, column=0, padx=10, pady=5, sticky='w')
        column_menu = AutocompleteCombobox(self.filter_window, textvariable=self.filter_column, font=self.default_font)
        column_menu['values'] = self.columns  # Populate with column names
        column_menu.grid(row=0, column=1, padx=10, pady=5, sticky='w')

        Label(self.filter_window, text="Enter Filter Text:", font=self.default_font).grid(row=1, column=0, padx=10, pady=5, sticky='w')
        filter_entry = Entry(self.filter_window, textvariable=self.filter_text, font=self.default_font)
        filter_entry.grid(row=1, column=1, padx=10, pady=5, sticky='w')

        Button(self.filter_window, text="Apply Filter", command=self.apply_filter, font=self.default_font).grid(row=2, column=0, columnspan=2, pady=5)
        Button(self.filter_window, text="Reset", command=self.reset_filters, font=self.default_font).grid(row=3, column=0, columnspan=2, pady=5)

    def on_double_click(self, event):
        curItem = self.tree.focus()
        if not curItem:
            return
        item_values = self.tree.item(curItem)['values']
        self.selected_attendance_pk = self.attendance_pk_map.get(curItem)  # Retrieve attendance_pk using the item_id
        student_id, first_name, last_name, site, class_date, class_hours = item_values
        title_string = f"Record: {first_name} {last_name} @ {site} on {class_date}"

        self.attendance_record_window = Toplevel(self)
        self.attendance_record_window.title(title_string)
        center_window(self.attendance_record_window, width=450, height=250)

        # Define fields for editing
        self.fields = [
            ("Class Date:", StringVar(value=class_date)),
            ("Site:", StringVar(value=site)),
            ("Hours:", StringVar(value=class_hours))
        ]
        self.labels = {}
        self.data_labels = {}
        self.entries = {}
        self.edit_mode = False

        for i, (label_text, string_var) in enumerate(self.fields):
            label = Label(self.attendance_record_window, text=label_text, font=self.default_font)
            data_label = Label(self.attendance_record_window, textvariable=string_var, font=self.default_font)
            entry = Entry(self.attendance_record_window, textvariable=string_var, font=self.default_font)

            label.grid(row=i, column=0, padx=5, pady=10, sticky='nsew')
            data_label.grid(row=i, column=1, padx=5, pady=10, sticky='nsew')
            entry.grid(row=i, column=2, padx=5, pady=10, sticky='nsew')
            entry.grid_remove()  # Hide the entry widget initially

            self.labels[label_text] = label
            self.data_labels[label_text] = data_label
            self.entries[label_text] = entry

        # Add Edit, Save, and Delete buttons
        self.edit_button = Button(self.attendance_record_window, text="Edit", command=self.enable_editing, font=self.default_font)
        self.edit_button.grid(row=len(self.fields), column=0, padx=10, pady=10, sticky='w')

        self.save_button = Button(self.attendance_record_window, text="Save", command=self.save_changes, font=self.default_font)
        self.save_button.grid(row=len(self.fields), column=1, padx=10, pady=10, sticky='w')
        self.save_button.grid_remove()  # Hide the Save button initially

        self.delete_button = Button(self.attendance_record_window, text="Delete", command=lambda: self.delete_attendance(self.selected_attendance_pk), font=self.default_font)
        self.delete_button.grid(row=len(self.fields), column=2, padx=10, pady=10, sticky='e')

        # Configure the grid to expand properly
        for i in range(1, 4):
            self.attendance_record_window.grid_rowconfigure(i, weight=1)
        self.attendance_record_window.grid_columnconfigure(0, weight=1)
        self.attendance_record_window.grid_columnconfigure(1, weight=1)
        self.attendance_record_window.grid_columnconfigure(2, weight=1)

    def enable_editing(self):
        for label_text in self.labels.keys():
            self.data_labels[label_text].grid_remove()
            self.entries[label_text].grid()
        self.edit_button.config(state='disabled')
        self.save_button.config(state='normal')
        self.edit_mode = True
        self.save_button.grid()  # Show the Save button

    def save_changes(self):
        if self.edit_mode:
            curItem = self.tree.focus()
            item_values = self.tree.item(curItem)['values']
            
            old_data = {
                'Class Date:': item_values[4],
                'Site:': item_values[3],
                'Hours:': item_values[5]
            }

            selected_student_id = item_values[0]
            
            updated_data = {label: var.get() for label, var in self.fields}
            self.update_database(self.selected_attendance_pk, updated_data, selected_student_id)

            # Refresh the treeview
            self.populate_treeview()
            
            self.attendance_record_window.destroy()

    def update_database(self, attendance_pk, updated_data, selected_student_id):
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        try:
            query = '''
                UPDATE 
                    Attendance
                SET 
                    class_date = ?, 
                    site = ?, 
                    class_hours = ?
                WHERE 
                    attendance_pk = ?
            '''
            cursor.execute(query, (
                format_date(updated_data['Class Date:'], out=True), 
                updated_data['Site:'], 
                updated_data['Hours:'], 
                attendance_pk
            ))
            conn.commit()
            messagebox.showinfo("Success", "Attendance record has been updated successfully.")
                    
            # Refresh the treeview
            self.populate_treeview()
            
        except Exception as e:
            conn.rollback()
            messagebox.showerror("Error", f"An error occurred while updating the attendance record: {e}")
        finally:
            conn.close()

    def delete_attendance(self, attendance_pk):
        # Confirm deletion
        confirm_delete = messagebox.askyesno("Delete Attendance", "Are you sure you want to delete this attendance record? This action cannot be undone.")
        if not confirm_delete:
            return
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        try:
            # Delete from Attendance table
            cursor.execute('DELETE FROM Attendance WHERE attendance_pk = ?', (attendance_pk,))
            conn.commit()
            messagebox.showinfo("Success", "Attendance record has been deleted successfully.")
            # Update the treeview
            self.populate_treeview()
        except Exception as e:
            conn.rollback()
            messagebox.showerror("Error", f"An error occurred while deleting the attendance record: {e}")
        finally:
            conn.close()
        self.attendance_record_window.destroy()

class AttendanceSiteTab(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.create_widgets()

    def create_widgets(self):
        pass

class AttendanceAddTab(ttk.Frame):
    def __init__(self, parent, default_font, attendance_all_tab):
        super().__init__(parent)
        self.parent = parent
        self.default_font = default_font
        self.attendance_all_tab = attendance_all_tab
        self.selected_site = StringVar()
        self.selected_date = StringVar(value=datetime.today().strftime('%m-%d-%Y'))
        self.total_attendance_records_counter = IntVar(value=0)
        self.checked_items = set()  # Track checked items
        self.sort_directions = {"Select": False, "Last Name": True, "First Name": True, "Last Class": False}
        self.create_widgets()
        self.populate_sites()

    def create_widgets(self):
        Label(self, text="Select Site:", font=self.default_font).grid(row=0, column=0, padx=10, pady=5, sticky='w')
        self.site_combobox = AutocompleteCombobox(self, textvariable=self.selected_site, state='readonly', font=self.default_font)
        self.site_combobox.grid(row=0, column=1, padx=10, pady=5, sticky='w')
        self.site_combobox.bind("<<ComboboxSelected>>", self.load_students)

        # Date selection
        Label(self, text="Select Date:", font=self.default_font).grid(row=1, column=0, padx=10, pady=5, sticky='w')
        large_font = font.Font(family="Helvetica", size=14, weight="bold")
        self.cal = Calendar(self, selectmode='day', 
                            year=datetime.today().year, month=datetime.today().month, 
                            day=datetime.today().day, font=large_font, 
                            headersbackground="lightblue",
                            background="lightgrey", foreground="black",
                            selectbackground="blue", selectforeground="white")
        self.cal.grid(row=2, column=0, columnspan=2, padx=10, pady=5)
        self.cal.bind("<<CalendarSelected>>", self.update_date)

        # Load checkbox images
        self.unchecked_img = ImageTk.PhotoImage(Image.new('RGBA', (15, 15), (255, 255, 255, 0)))
        self.checked_img = ImageTk.PhotoImage(Image.new('RGBA', (15, 15), (0, 0, 255, 255)))

        # Student table with checkboxes
        self.tree = ttk.Treeview(self, columns=("Select", "Last Name", "First Name", "Last Class"), show='headings')
        self.tree.heading("Select", text="Select", command=lambda: self.sort_students("Select"))
        self.tree.heading("Last Name", text="Last Name", command=lambda: self.sort_students("Last Name"))
        self.tree.heading("First Name", text="First Name", command=lambda: self.sort_students("First Name"))
        self.tree.heading("Last Class", text="Last Class", command=lambda: self.sort_students("Last Class"))
        self.tree.grid(row=3, column=0, columnspan=4, padx=10, pady=5, sticky='nsew')

        # Scrollbars
        vsb = ttk.Scrollbar(self, orient="vertical", command=self.tree.yview) 
        vsb.grid(row=3, column=4, sticky='ns')
        self.tree.configure(yscrollcommand=vsb.set)

        # Bind click event to toggle checkboxes
        self.tree.bind('<Button-1>', self.enable_toggle_checkbox)

        # Total records label
        Label(self, text="Total Students:", font=self.default_font).grid(row=4, column=0, padx=10, pady=5, sticky='w')
        Label(self, textvariable=self.total_attendance_records_counter, font=self.default_font).grid(row=4, column=1, padx=10, pady=5, sticky='w')

        # Save button
        self.save_button = Button(self, text="Save Attendance", command=self.save_attendance, font=self.default_font)
        self.save_button.grid(row=4, column=2, padx=10, pady=10, sticky='e')

        # Clear Selection button
        self.clear_button = Button(self, text="Clear Selection", command=self.clear_selection, font=self.default_font)
        self.clear_button.grid(row=4, column=3, padx=10, pady=10, sticky='e')

        # Configure grid weights to handle resizing
        self.grid_rowconfigure(3, weight=1)
        self.grid_columnconfigure(1, weight=1)

    def enable_toggle_checkbox(self, event):
        region = self.tree.identify('region', event.x, event.y)
        if region != 'heading':
            self.toggle_checkbox(event)

    def populate_sites(self):
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        cursor.execute('SELECT DISTINCT site_name FROM Sites ORDER BY site_name ASC')
        sites = [row[0] for row in cursor.fetchall()]
        conn.close()
        self.site_combobox['values'] = sites

    def load_students(self, event=None):
        site = self.selected_site.get()
        if not site:
            return
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        # query = '''
        #     SELECT
        #         s.last_name,
        #         s.first_name,
        #         MAX(a.class_date) AS most_recent_class_date
        #     FROM
        #         StudentSites ss
        #     LEFT JOIN
        #         Sites st ON ss.site_pk = st.site_pk
        #     LEFT JOIN
        #         Student s ON ss.student_id = s.student_id
        #     LEFT JOIN
        #         Attendance a ON ss.student_id = a.student_id
        #     WHERE
        #         st.site_name = ? AND s.last_name IS NOT NULL
        #     GROUP BY
        #         s.student_id
        #     ORDER BY
        #         most_recent_class_date DESC,
        #         s.last_name ASC;
        # '''
        query = '''
            WITH Attended AS (
                SELECT 
                    a.student_id, 
                    MAX(a.class_date) AS last_class
                FROM Attendance a
                WHERE a.site = ?
                GROUP BY a.student_id
            )
            SELECT DISTINCT
                s.last_name, 
                s.first_name, 
                COALESCE(a.last_class, '') AS last_class
            FROM StudentSites ss
            LEFT JOIN Sites st ON st.site_pk = ss.site_pk
            LEFT JOIN Student s ON s.student_id = ss.student_id
            LEFT JOIN Attended a ON s.student_id = a.student_id
            WHERE st.site_name = ? AND s.last_name IS NOT NULL
            ORDER BY last_class DESC, s.first_name ASC;

        '''

        cursor.execute(query, (site, site,))
        students = cursor.fetchall()
        conn.close()

        # Clear existing rows
        for row in self.tree.get_children():
            self.tree.delete(row)

        # Populate students and remove leading zeros
        for student in students:
            first_name, last_name, last_class = student
            if last_class:
                last_class = format_date(last_class)
            item = self.tree.insert('', 'end', values=("", first_name, last_name, last_class if last_class else ""), image=self.unchecked_img)
            self.checked_items.discard(item)

    def update_date(self, event=None):
        # Convert the date string to a datetime object using flexible parsing
        try:
            date_obj = datetime.strptime(self.cal.get_date(), '%m/%d/%y')
            formatted_date = date_obj.strftime('%m-%d-%Y')
            self.selected_date.set(formatted_date)
        except ValueError as e:
            messagebox.showerror("Error", f"Invalid date format: {e}")

    def sort_students(self, col):
        # Toggle the sorting direction
        self.sort_directions[col] = not self.sort_directions[col]
        data = [(self.tree.set(child, col), child) for child in self.tree.get_children('')]
        if col == 'Last Class':
            # Handle sorting for date column
            data.sort(key=lambda x: datetime.strptime(x[0], '%m-%d-%Y') if x[0] else datetime.min, reverse=self.sort_directions[col])
        else:
            # Handle sorting for text columns
            data.sort(reverse=self.sort_directions[col])

        for index, (val, child) in enumerate(data):
            self.tree.move(child, '', index)

        # Update heading with new sort direction
        self.tree.heading(col, command=lambda: self.sort_students(col))

    def toggle_checkbox(self, event):
        item = self.tree.identify_row(event.y)
        if item:
            if item in self.checked_items:
                self.tree.item(item, image=self.unchecked_img, values=("",) + self.tree.item(item, 'values')[1:])
                self.checked_items.discard(item)
                c = self.total_attendance_records_counter.get()
                if c != 0:
                    c -= 1
                    self.total_attendance_records_counter.set(c)
            else:
                self.tree.item(item, image=self.checked_img, values=("✓",) + self.tree.item(item, 'values')[1:])
                self.checked_items.add(item)
                c = self.total_attendance_records_counter.get()
                c += 1
                self.total_attendance_records_counter.set(c)

    def save_attendance(self):
        if not self.checked_items:
            messagebox.showwarning("Warning", "No students selected for attendance.")
            return

        site = self.selected_site.get()
        date = self.selected_date.get()

        # Display a confirmation message box before saving
        confirm = messagebox.askokcancel("Confirm Attendance",
                                        f"Confirm save attendance for: \n\n\tSite: {site} \n\tDate: {format_date(date)}")
        
        if not confirm:
            return  # Exit if the user cancels the action

        # Ensure the date is already formatted as YYYY-MM-DD
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        try:
            for item in self.checked_items:
                last_name, first_name = self.tree.item(item, 'values')[1:3]
                # Assuming student_id is needed to insert into Attendance
                query = 'SELECT student_id FROM Student WHERE last_name = ? AND first_name = ?'
                cursor.execute(query, (last_name, first_name,))
                student_id = cursor.fetchone()[0]
                cursor.execute('''
                    INSERT INTO Attendance (student_id, class_date, site, class_hours)
                    VALUES (?, ?, ?, ?)
                ''', (student_id, format_date(date, out=True), site, 2.5))

            conn.commit()

            # Reload class_date and site filter dropdowns in AttendanceAll
            AttendanceAllTab.create_widgets(self.attendance_all_tab)

            messagebox.showinfo("Success", "Attendance records have been saved successfully.")

        except Exception as e:
            conn.rollback()
            messagebox.showerror("Error", f"An error occurred: {e}")
            
        finally:
            # Uncheck students whose attendance has been saved and set attendance counter to 0
            self.total_attendance_records_counter.set(0)
            self.clear_selection()
            conn.close()

    def clear_selection(self):
        for item in self.checked_items:
            self.tree.item(item, image=self.unchecked_img, values=("",) + self.tree.item(item, 'values')[1:])
        self.checked_items.clear()
        self.total_attendance_records_counter.set(0)

class ClassTab(ttk.Frame):
    def __init__(self, parent, default_font):
        super().__init__(parent)
        self.default_font = default_font
        self.parent = parent
        self.string_vars = {}
        self.fields = [
            ("Class Date:", "class_date"),
            ("Site:", "class_site"),
            ("Class Notes:", "class_notes"),
            ("To Do:", "to_do")
        ]
        self.create_widgets()
        self.populate_treeview()
        self.doc_dir = LP_DIR_PATH
        self.sort_orders = {}  # Initialize sort orders for each column

    def reset_fields(self):
        """Reset the self.fields variable to ensure consistency."""
        self.fields = [
            ("Class Date:", "class_date"),
            ("Site:", "class_site"),
            ("Class Notes:", "class_notes"),
            ("To Do:", "to_do")
        ]
    
    def create_widgets(self):
        # Create Treeview
        self.columns = ["Class Date", "Site", "Class Notes", "To Do"]
        self.tree = ttk.Treeview(self, columns=self.columns, show='headings', height=1)
        
        # Measure text width for headers
        for col in self.columns:
            self.tree.heading(col, text=col, command=lambda c=col: self.sort_by_column(c))
            # Only adjust the first two columns
            if col in self.columns[:2]:
                header_font = Font()
                max_header_width = header_font.measure(col)
                # Calculate the maximum width needed for the header text
                self.tree.column(col, anchor='center', width=max_header_width+80, stretch=NO)
            else:
                self.tree.column(col, anchor='center')

        self.tree.grid(row=1, column=0, columnspan=4, padx=10, pady=10, sticky='nsew')

        # Scrollbars
        vsb = ttk.Scrollbar(self, orient="vertical", command=self.tree.yview)
        vsb.grid(row=1, column=4, sticky='ns')
        self.tree.configure(yscrollcommand=vsb.set)

        # Buttons
        self.add_button = Button(self, text="Add Class", command=self.add_class, font=self.default_font)
        self.add_button.grid(row=0, column=0, padx=10, pady=5, sticky='w')

        self.filter_button = Button(self, text="Filter", command=self.open_filter_window, font=self.default_font)
        self.filter_button.grid(row=0, column=1, padx=10, pady=5, sticky='w')

        self.import_button = Button(self, text="Open Lesson Plan", command=self.import_lesson_plan, font=self.default_font)
        self.import_button.grid(row=0, column=2, padx=10, pady=5, sticky='w')

        self.create_lesson_button = Button(self, text="Create Lesson", command=self.create_lesson, font=self.default_font)
        self.create_lesson_button.grid(row=0, column=3, padx=10, pady=5, sticky='w')

        # Configure grid to expand properly
        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(0, weight=1)

        # Bind double click to treeview items
        self.tree.bind("<Double-1>", self.on_treeview_click)

    def populate_treeview(self, filtered_rows=None):
        # Clear existing rows
        for row in self.tree.get_children():
            self.tree.delete(row)

        if filtered_rows is None:
            # Connect to the SQLite database
            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()
            query = '''
                SELECT
                    class_pk,
                    class_date,
                    class_site,
                    class_notes,
                    to_do
                FROM 
                    Class
                ORDER BY 
                    class_date DESC
            '''
            cursor.execute(query)
            rows = cursor.fetchall()
            conn.close()
        else:
            rows = filtered_rows

        # Define pk map for later editing
        self.class_pk_map = {}

        # Format date: YYYY-MM-DD -> MM-DD-YYYY (no leading zeros)
        for row in rows:
            try:
                # Convert date format
                item_id = self.tree.insert('', 'end', values=(format_date(row[1]), row[2], row[3].replace('\n',' '), row[4].replace('\n',' ')))
                self.class_pk_map[item_id] = row[0]

            except ValueError as e:
                print(f"Error processing row {row}: {e}")
                
    def sort_by_column(self, col):
        # Toggle sort order for the column
        if col not in self.sort_orders:
            self.sort_orders[col] = False  # Default to ascending order
        self.sort_orders[col] = not self.sort_orders[col]  # Reverse sort order

        # Sort logic for treeview columns
        data = [(self.tree.set(child, col), child) for child in self.tree.get_children('')]

        if col == 'Class Date':
            # Convert date strings to datetime objects for sorting
            data.sort(key=lambda t: datetime.strptime(t[0], '%m-%d-%Y'), reverse=self.sort_orders[col])
        else:
            # Sort other columns as strings
            data.sort(key=lambda t: t[0].lower(), reverse=self.sort_orders[col])

        for index, (val, child) in enumerate(data):
            self.tree.move(child, '', index)

    def open_filter_window(self):
        self.filter_window = Toplevel(self)
        self.filter_window.title("Filter Records")

        Label(self.filter_window, text="Select Column:", font=self.default_font).grid(row=0, column=0, padx=10, pady=5, sticky='w')
        self.filter_column = StringVar()
        column_menu = AutocompleteCombobox(self.filter_window, textvariable=self.filter_column, font=self.default_font)
        column_menu['values'] = self.columns
        column_menu.grid(row=0, column=1, padx=10, pady=5, sticky='w')

        Label(self.filter_window, text="Enter Filter Text:", font=self.default_font).grid(row=1, column=0, padx=10, pady=5, sticky='w')
        self.filter_text = StringVar()
        filter_entry = Entry(self.filter_window, textvariable=self.filter_text, font=self.default_font)
        filter_entry.grid(row=1, column=1, padx=10, pady=5, sticky='w')

        Button(self.filter_window, text="Apply Filter", command=self.apply_filter, font=self.default_font).grid(row=2, column=0, columnspan=2, pady=5)
        Button(self.filter_window, text="Reset", command=self.reset_filters, font=self.default_font).grid(row=3, column=0, columnspan=2, pady=5)

    def apply_filter(self):
        column = self.filter_column.get()
        filter_text = self.filter_text.get().lower().strip()
        filtered_rows = []

        if column and filter_text:
            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()

            query = 'SELECT class_date, class_site, class_notes, to_do FROM Class'
            cursor.execute(query)
            rows = cursor.fetchall()
            conn.close()

            # Remove leading zeros from the filter text
            filter_text = re.sub(r'\b0+(\d)', r'\1', filter_text)

            # Apply the filter to the rows
            for row in rows:
                formatted_row = list(row)

                # Format dates for filtering
                class_date = formatted_row[0]

                # Convert date to different searchable formats
                def date_matches(date_str, search_text):
                    try:
                        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                        formatted_date = [
                            date_obj.strftime('%Y-%m-%d'),  # Full date
                            date_obj.strftime('%m-%d-%Y'),  # Alternative full date format
                            date_obj.strftime('%y-%m-%d'),  # Short year with full date
                            date_obj.strftime('%m-%d'),     # Month-day only
                            date_obj.strftime('%Y-%m'),     # Month
                            date_obj.strftime('%m'),        # Year and month
                            date_obj.strftime('%Y'),        # Year only
                            date_obj.strftime('%y'),        # Short year only
                            date_obj.strftime('%B'),        # full month name
                            date_obj.strftime('%b'),        # short month name
                            date_obj.strftime('%B %Y'),     # full month name with year
                            date_obj.strftime('%b %Y')      # short month name with yar                        
                        ]

                        # Normalize month and day by removing leading zeros
                        normalized_date = [
                            re.sub(r'\b0+(\d)', r'\1', d) for d in formatted_date
                        ]

                        # Check if any normalized version matches the search text
                        return any(search_text in nd.lower() for nd in normalized_date)
                    except ValueError:
                        return False

                # Check if the filter text matches in any format
                if column == 'Class Date':
                    if date_matches(class_date, filter_text):
                        filtered_rows.append(formatted_row)
                else:
                    col_index = self.columns.index(column)
                    if filter_text in str(formatted_row[col_index]).lower():
                        filtered_rows.append(formatted_row)

        self.filter_window.destroy()
        self.populate_treeview(filtered_rows if filtered_rows else None)

    def reset_filters(self):
        self.filter_text.set('')
        self.filter_column.set('')
        self.populate_treeview()

    def add_class(self):
        # Reset fields before opening the add class window
        self.reset_fields()

        # Create the Add Class window
        add_class_window = Toplevel(self)
        add_class_window.withdraw()
        add_class_window.title("Add Class")
        center_window(add_class_window, 700, 700)  # Ensure this function is defined

        # Define the class record label and entry fields
        string_vars = {}  # Use local dictionary for this window's variables

        # Create Save button
        save_class_button = Button(
            add_class_window,
            text="Save",
            command=lambda: self.save_class(string_vars, class_notes_text, to_do_text_widget),
            state='disabled', 
            font=self.default_font
        )
        save_class_button.grid(row=len(self.fields) + 2, column=0, columnspan=2, pady=10)

        def check_fields():
            notes_text = class_notes_text.get("1.0", "end-1c").strip()
            to_do_text = to_do_text_widget.get("1.0", "end-1c").strip()
            if all(var.get().strip() for key, var in string_vars.items() if key not in ('class_notes', 'class_site', 'to_do')) and notes_text and to_do_text:
                save_class_button.config(state='normal')
            else:
                save_class_button.config(state='disabled')

        # Create labels and entries
        for i, (label_text, var_name) in enumerate(self.fields):
            label = Label(add_class_window, text=label_text, font=self.default_font)
            label.grid(row=i, column=0, padx=10, pady=5, sticky='n')

            if label_text == "Class Notes:":
                class_notes_text = Text(add_class_window, wrap=WORD, height=5, width=40)
                class_notes_text.grid(row=i, column=1, padx=10, pady=5, sticky='nsew')
                class_notes_text.bind('<KeyRelease>', lambda event: check_fields())
            elif label_text == "To Do:":
                to_do_text_widget = Text(add_class_window, wrap=WORD, height=5, width=40)
                to_do_text_widget.grid(row=i, column=1, padx=10, pady=5, sticky='nsew')
                to_do_text_widget.bind('<KeyRelease>', lambda event: check_fields())
            elif label_text == "Site:":
                string_var = StringVar()
                string_vars[var_name] = string_var
                site_names = get_site_names()
                combobox = AutocompleteCombobox(add_class_window, textvariable=string_var, values=site_names, font=self.default_font)
                combobox.grid(row=i, column=1, padx=10, pady=5, sticky='n')
            else:
                string_var = StringVar()
                string_vars[var_name] = string_var  # Store the StringVar object in the dictionary
                entry = Entry(add_class_window, textvariable=string_var, font=self.default_font)
                entry.grid(row=i, column=1, padx=10, pady=5, sticky='n')

                if label_text == "Class Date:":
                    large_font = font.Font(family="Helvetica", size=14, weight="bold")
                    cal = Calendar(add_class_window, selectmode='day',
                                   year=datetime.today().year, month=datetime.today().month,
                                   day=datetime.today().day, font=large_font,
                                   headersbackground="lightblue",
                                   background="lightgrey", foreground="black",
                                   selectbackground="blue", selectforeground="white")
                    cal.grid(row=0, column=1, columnspan=2, padx=10, pady=5)
                    cal.bind("<<CalendarSelected>>", lambda event, var=string_var: var.set(cal.get_date()))

                string_var.trace('w', lambda name, index, mode, var=string_var: check_fields())

        # Configure the grid to expand properly
        for i in range(1, len(self.fields) + 2):
            add_class_window.grid_rowconfigure(i, weight=1)
        add_class_window.grid_columnconfigure(0, weight=1)
        add_class_window.grid_columnconfigure(1, weight=1)
        add_class_window.deiconify()

    def save_class(self, string_vars, class_notes_text, to_do_text_widget):
        msg = "Save class record?"
        if messagebox.askyesno("Add Class", msg):
            class_date = string_vars['class_date'].get()
            class_site = string_vars['class_site'].get()
            class_notes = class_notes_text.get("1.0", "end-1c").strip()
            to_do = to_do_text_widget.get("1.0", "end-1c").strip()

            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()

            cursor.execute('''
                INSERT INTO Class (class_date, class_site, class_notes, to_do)
                VALUES (?, ?, ?, ?)
            ''', (format_date(class_date, out=True), class_site, class_notes, to_do))

            conn.commit()
            conn.close()

            class_notes_text.master.destroy()  # Destroy the window associated with the widget
            self.populate_treeview()

    def on_treeview_click(self, event):
        region = self.tree.identify('region', event.x, event.y)
        if region != 'heading':
            self.on_double_click(event)

    def on_double_click(self, event):
        self.reset_fields()
        curItem = self.tree.focus()
        item_values = self.tree.item(curItem)['values']
        self.selected_class_pk = self.class_pk_map[curItem]
        class_date = item_values[0]
        class_site = item_values[1]
        # Fetch the notes directly from the database using the class_date as a key
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
    
        # Fetch the class record to get the full notes text
        query = '''
            SELECT class_notes, to_do
            FROM Class
            WHERE class_pk = ?
        '''
        cursor.execute(query, (self.selected_class_pk,))
        class_data = cursor.fetchone()
        conn.close()

        class_notes = class_data[0] if class_data else ""
        to_do = class_data[1] if class_data else ""

        title_string = f"Class on: {class_date}"

        self.class_record_window = Toplevel(self)
        self.class_record_window.title(title_string)
        center_window(self.class_record_window, 800, 500)

        # Define fields for editing
        self.fields = [
            ("Class Date:", StringVar(value=class_date)),
            ("Site:", StringVar(value=class_site))
        ]
        self.labels = {}
        self.data_labels = {}
        self.entries = {}
        self.edit_mode = False

        for i, (label_text, string_var) in enumerate(self.fields):
            label = Label(self.class_record_window, text=label_text, font=self.default_font)
            data_label = Label(self.class_record_window, textvariable=string_var, font=self.default_font)
            entry = Entry(self.class_record_window, textvariable=string_var, font=self.default_font)

            label.grid(row=i, column=0, padx=5, pady=10, sticky='nsew')
            data_label.grid(row=i, column=1, padx=5, pady=10, sticky='nsew')
            entry.grid(row=i, column=2, padx=5, pady=10, sticky='nsew')
            entry.grid_remove()  # Hide the entry widget initially

            self.labels[label_text] = label
            self.data_labels[label_text] = data_label
            self.entries[label_text] = entry

        # Add a Text widget for notes with line wrap
        class_notes_label = Label(self.class_record_window, text="Class Notes:", font=self.default_font)
        class_notes_label.grid(row=len(self.fields), column=0, padx=5, pady=10, sticky='nsew')

        self.class_notes_text = Text(self.class_record_window, wrap=WORD, height=5, width=50)
        self.class_notes_text.grid(row=len(self.fields), column=1, columnspan=2, padx=5, pady=10, sticky='nsew')

        # Insert existing class notes with line breaks
        self.class_notes_text.insert(END, class_notes)

        # Add a Text widget for to_do with line wrap
        to_do_label = Label(self.class_record_window, text="To Do:", font=self.default_font)
        to_do_label.grid(row=len(self.fields)+1, column=0, padx=5, pady=10, sticky='nsew')

        self.to_do_text = Text(self.class_record_window, wrap=WORD, height=5, width=50)
        self.to_do_text.grid(row=len(self.fields)+1, column=1, columnspan=2, padx=5, pady=10, sticky='nsew')

        # Insert existing to_do with line breaks
        self.to_do_text.insert(END, to_do)

        # Store the initial notes and to-do content
        self.initial_class_notes = class_notes
        self.initial_to_do = to_do

        # Add Edit, Save, and Delete buttons
        self.edit_button = Button(self.class_record_window, text="Edit", command=self.enable_editing, font=self.default_font)
        self.edit_button.grid(row=len(self.fields) + 2, column=0, padx=10, pady=10, sticky='w')

        self.save_button = Button(self.class_record_window, text="Save", command=self.save_changes, font=self.default_font)
        self.save_button.grid(row=len(self.fields) + 2, column=1, padx=10, pady=10, sticky='w')
        self.save_button.grid_remove()  # Hide the Save button initially

        self.delete_button = Button(self.class_record_window, text="Delete", command=lambda: self.delete_class(self.selected_class_pk), font=self.default_font)
        self.delete_button.grid(row=len(self.fields) + 2, column=2, padx=10, pady=10, sticky='e')

        # Configure the grid to expand properly
        for i in range(len(self.fields) + 2):
            self.class_record_window.grid_rowconfigure(i, weight=1)
        self.class_record_window.grid_columnconfigure(0, weight=1)
        self.class_record_window.grid_columnconfigure(1, weight=1)
        self.class_record_window.grid_columnconfigure(2, weight=1)

    def enable_editing(self):
        for label_text in self.labels.keys():
            self.data_labels[label_text].grid_remove()
            self.entries[label_text].grid()

        self.class_notes_text.config(state='normal')  # Enable editing for the notes field
        self.to_do_text.config(state='normal')  # Enable editing for the to-do field
        self.edit_button.config(state='disabled')
        self.save_button.config(state='normal')
        self.edit_mode = True
        self.save_button.grid()  # Show the Save button

    def save_changes(self):
        if self.edit_mode:

            updated_data = {
                'Class Date:': self.entries["Class Date:"].get(),
                'Site:': self.entries["Site:"].get(),
                'Class Notes:': self.class_notes_text.get("1.0", "end-1c").strip(),
                'To Do:': self.to_do_text.get("1.0", "end-1c").strip()
            }

            # Update the database
            self.update_database(self.selected_class_pk, updated_data)

            # Refresh the treeview in Class tab
            self.populate_treeview()

            # Close the class record window
            self.class_record_window.destroy()

    def update_database(self, class_pk, updated_data):
        # Connect to the SQLite database
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        try:
            query = '''
                UPDATE 
                    Class
                SET 
                    class_date = ?, 
                    class_site = ?, 
                    class_notes = ?, 
                    to_do = ?
                WHERE 
                    class_pk = ?
            '''
            cursor.execute(query, (
                format_date(updated_data['Class Date:'], out=True),
                updated_data['Site:'],
                updated_data['Class Notes:'],
                updated_data['To Do:'],
                class_pk,
            ))
            conn.commit()
            messagebox.showinfo("Success", "Class record has been updated successfully.")
        except Exception as e:
            conn.rollback()
            messagebox.showerror("Error", f"An error occurred while updating the class record: {e}")
        finally:
            conn.close()

    def delete_class(self, class_pk):
        # Confirm deletion
        confirm_delete = messagebox.askyesno("Delete Class", "Are you sure you want to delete this class record? This action cannot be undone.")
        if not confirm_delete:
            return
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        try:
            # Delete from Class table
            query = '''
                DELETE FROM 
                    Class 
                WHERE 
                    class_pk = ? 
            '''
            cursor.execute(query, (class_pk,))
            conn.commit()
            messagebox.showinfo("Success", "Class record has been deleted successfully.")
            # Update the treeview in ClassTab
            self.populate_treeview()
        except Exception as e:
            conn.rollback()
            messagebox.showerror("Error", f"An error occurred while deleting the class record: {e}")
        finally:
            conn.close()
        self.class_record_window.destroy()

    def import_lesson_plan(self):
        # Create a new Tkinter window
        # window = Tk()
        window = Toplevel(self)
        window.title("Open Lesson Plan")
        center_window(window, 1500, 500)

        # Create a Treeview widget
        tree = ttk.Treeview(window, columns=('Lesson Title', 'Start Date', 'End Date', 'Creation Date'), show='headings')
        tree.heading('Lesson Title', text='Lesson Title', command=lambda: self.sort_column(tree, 'Lesson Title', False))
        tree.heading('Start Date', text='Start Date', command=lambda: self.sort_column(tree, 'Start Date', False))
        tree.heading('End Date', text='End Date', command=lambda: self.sort_column(tree, 'End Date', False))
        tree.heading('Creation Date', text='Creation Date', command=lambda: self.sort_column(tree, 'Creation Date', False))
                
        # Center-align the columns
        for col in tree['columns']:
            tree.column(col, anchor='center')  # Center-align text
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(window, orient="vertical", command=tree.yview)
        tree.configure(yscroll=scrollbar.set)
        scrollbar.pack(side='right', fill='y')

        # Adjust column widths based on content
        for col in tree['columns']:
            tree.column(col, width=tkFont.Font().measure(col))  # Start with the header width
            for item in tree.get_children():
                cell_value = tree.item(item, 'values')[self.tree['columns'].index(col)]
                cell_width = tkFont.Font().measure(cell_value)
                if tree.column(col, width=None) < cell_width:
                    tree.column(col, width=cell_width)

        # Pack the Treeview
        tree.pack(expand=True, fill='both')

        def on_treeview_click(event):
            region = self.tree.identify('region', event.x, event.y)
            if region != 'heading':
                self.open_on_double_click(event)

        # Bind double-click event
        tree.bind('<Double-1>', on_treeview_click)

        # Regex pattern for filtering files
        pattern = r"Hoffert_Philip_Aspire_LP_(.*?)_(\d{1,2}-\d{1,2}-\d{2})_(\d{1,2}-\d{1,2}-\d{2})_(\d{1,2}-\d{1,2}-\d{2})\.docx"

        files = []

        # List all files in the directory
        for file_name in os.listdir(self.doc_dir):
            match = re.match(pattern, file_name)
            if match:
                lesson_title = match.group(1).replace('_', ' ')  # Replace underscores with spaces
                start_date = match.group(2)
                end_date = match.group(3)
                creation_date = match.group(4)
                
                # Convert start_date to a datetime object for accurate sorting
                start_date_obj = datetime.strptime(start_date, '%m-%d-%y')  

                files.append((file_name, lesson_title, start_date_obj, end_date, creation_date))

                # tree.insert('', 'end', iid=file_name, values=(lesson_title, start_date, end_date, creation_date))

        files.sort(key=lambda x: x[2], reverse=True)

        for file_name, lesson_title, start_date_obj, end_date, creation_date in files:
            start_date_str = start_date_obj.strftime('%m-%d-%y')  # Convert back to string if needed
            if os.name == 'nt':
                start_date_str = format_date(start_date_str, display_format="%#m-%#d-%y")
            else:
                start_date_str = format_date(start_date_str, display_format="%-m-%-d-%y")
            tree.insert('', 'end', iid=file_name, values=(lesson_title, start_date_str, end_date, creation_date))


        # Run the Tkinter main loop
        window.mainloop()

    def sort_column(self, tree, col, reverse):
        # Define the columns that contain dates
        date_columns = {"Start Date", "End Date", "Creation Date"}

        # Get all the values in the column
        data = [(tree.set(child, col), child) for child in tree.get_children('')]

        # Determine if the column is a date column
        if col in date_columns:
            # Convert the date strings to datetime objects for proper sorting
            data = [
                (datetime.strptime(val, '%m-%d-%y'), child) if val else (datetime.min, child)
                for val, child in data
            ]
        
        # Sort the data
        data.sort(reverse=reverse)

        # Rearrange the items in sorted positions
        for index, (_, item) in enumerate(data):
            tree.move(item, '', index)

        # Reverse sort order for the next click
        tree.heading(col, command=lambda: self.sort_column(tree, col, not reverse))

    def open_on_double_click(self, event):
        # Get the item that was double-clicked
        tree = event.widget
        item = tree.selection()[0]  # Get selected item
        # file_name = tree.item(item, 'tags')[0]  # Retrieve the file name from the tags
        file_name = self.doc_dir + '/' + item

        if item:
            file_path = os.path.join(self.doc_dir, file_name)

            # Open the file using the default application for .docx files
            if os.name == 'posix':  # macOS or Linux
                subprocess.call(('open', file_path))
            elif os.name == 'nt':  # Windows
                os.startfile(file_path)

    def create_lesson(self):
        self.create_lesson_window = Toplevel(self)
        center_window(self.create_lesson_window, width=screen_width, height=screen_height)
        self.create_lesson_window.title("Create Lesson Plan")
        lesson_planner = LessonPlanner(self.create_lesson_window, self.default_font)
        lesson_planner.create_lesson()

class LessonPlanner:
    def __init__(self, root, default_font):
        self.root = root
        self.default_font = default_font
        self.headers = {
            "[Lesson Title]": "",
            "TEACHER NAME": "",
            "PROGRAM NAME": "",
            "[Unit Title]": "",
            "NRS EFL(s)": "",
            "TIME FRAME": "",
        }

        self.standards = {
            "[Receptive 1]": "",
            "[Receptive 6]": "",
            "[Receptive 7]": "",
            "[Receptive 8]": "",
            "[Productive 3]": "",
            "[Productive 4]": "",
            "[Productive 7]": "",
            "[Productive 9]": "",
            "[Productive 10]": "",
            "[Interactive 2]": "",
            "[Interactive 5]": ""
        }

        self.footers = {
            'LEARNER OUTCOME(S)': "",
            'ASSESSMENT TOOLS/METHODS': "",
            'LEARNER PRIOR KNOWLEDGE': "",
            'INSTRUCTIONAL ACTIVITIES': "",
            'RESOURCES': "",
            "DIFFERENTIATION": "",
            "TEACHER REFLECTION/LESSON EVALUATION": "",
            "ADDITIONAL INFORMATION": ""
        }
        self.c = 0

        # Define your fields
        self.header_fields = [key for key in self.headers.keys() if key not in ('TEACHER NAME', 'PROGRAM NAME')]
        self.standards_fields = [key for key in self.standards.keys()]
        self.footers_fields = [key for key in self.footers.keys()]

        # StringVars for Tkinter
        self.header_string_vars = {field: StringVar() for field in self.header_fields}
        self.standards_string_vars = {field: StringVar() for field in self.standards_fields}
        self.footers_string_vars = {field: StringVar() for field in self.footers_fields}

        # Load the document
        self.doc_path = LP_DIR_PATH + '/Lesson Plan Template for Ohio ESOL Standards rvsd.docx'
        self.doc = Document(self.doc_path)

    def titlecase(self, some_str):
        clean_str = ""
        words = [w.lower().replace('[', '').replace(']', '') for w in some_str.split()]
        for w in words:
            clean_str += w[0].upper() + w[1:] + " "
        clean_str = clean_str.strip()
        clean_str = clean_str.replace('Nrs', 'NRS').replace('Efl(s)', 'EFL(s)')
        clean_str = clean_str.replace('Teacher Reflection/lesson Evaluation', 'Teacher Reflection/Lesson Evaluation')
        clean_str = clean_str.replace('/methods', '/Methods')
        clean_str += ":"
        return clean_str.strip()

    def create_lesson(self):
        # Configure the master window to expand properly
        self.root.grid_rowconfigure(0, weight=1)
        self.root.grid_columnconfigure(0, weight=1)

        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        self.root.state("zoomed")

        # Create a notebook widget
        notebook = ttk.Notebook(self.root)

        # Create frames for each section
        program_info_frame = ttk.Frame(notebook)
        receptive_frame = ttk.Frame(notebook)
        productive_frame = ttk.Frame(notebook)
        interactive_frame = ttk.Frame(notebook)
        reflection_frame = ttk.Frame(notebook)

        # Add frames as tabs in the notebook
        notebook.add(program_info_frame, text='Program Information')
        notebook.add(receptive_frame, text='Receptive')
        notebook.add(productive_frame, text='Productive')
        notebook.add(interactive_frame, text='Interactive')
        notebook.add(reflection_frame, text='Reflection')

        # Pack the Notebook widget to make it visible
        notebook.grid(row=0, column=0, padx=10, pady=10, sticky='nsew')

        # Define fields
        nrs_fields = '1 2 3 4 5 6 6+'.split()

        receptive_fields = ["1.1.1.", "2.1.1.", "3.1.1.", "4.1.1.", "5.1.1.", "6.1.1.", "3.1.2.", "4.1.2.", "5.1.2.", "6.1.2.", "4.1.3.", "5.1.3.", "6.1.3.", "4.1.4.", "5.1.4.", "6.1.4.", "4.1.5.", "2.6.1.", "3.6.1.", "4.6.1.", "5.6.1.", "6.6.1.", "3.6.2.", "4.6.2.", "5.6.2.", "6.6.2.", "5.6.3.", "6.6.3.", "1.7.2.", "2.7.2.", "1.8.1.", "2.8.1.", "3.8.1.", "4.8.1.", "5.8.1.", "6.8.1."]
        productive_fields = ["1.3.1.", "2.3.1.", "3.3.1.", "4.3.1.", "5.3.1.", "6.3.1.", "3.3.2.", "4.3.2.", "5.3.2.", "6.3.2.", "4.3.3.", "5.3.3.", "6.3.3.", "5.3.4.", "6.3.4.", "1.4.1.", "2.4.1.", "3.4.1.", "4.4.1.", "5.4.1.", "6.4.1.", "2.4.2.", "3.4.2.", "4.4.2.", "5.4.2.", "6.4.2.", "3.4.3.", "4.4.3.", "5.4.3.", "6.4.3.", "3.4.4.", "4.4.4.", "5.4.4.", "6.4.4.", "6.4.5.", "1.7.1.", "2.7.1.", "3.7.1.", "4.7.1.", "5.7.1.", "6.7.1.", "3.7.2.", "4.7.2.", "5.7.2.", "6.7.2.", "3.7.3.", "4.7.3.", "5.7.3.", "6.7.3.", "1.9.1.", "2.9.1.", "3.9.1.", "4.9.1.", "5.9.1.", "6.9.1.", "2.9.2.", "3.9.2.", "4.9.2.", "5.9.2.", "6.9.2.", "3.9.3.", "4.9.3.", "5.9.3.", "6.9.3.", "3.9.4.", "4.9.4.", "5.9.4.", "6.9.4.", "1.10.1.", "2.10.1.", "3.10.1.", "4.10.1.", "5.10.1.", "6.10.1.", "1.10.2.", "2.10.2.", "3.10.2.", "4.10.2.", "5.10.2.", "6.10.2.", "4.10.3.", "5.10.3."]
        interactive_fields = ["1.2.1.", "2.2.1.", "3.2.1.", "4.2.1.", "5.2.1.", "6.2.1.", "1.2.2.", "2.2.2.", "3.2.2.", "4.2.2.", "5.2.2.", "6.2.2.", "2.2.3.", "3.2.3.", "4.2.3.", "5.2.3.", "6.2.3.", "2.2.4.", "3.2.4.", "4.2.4.", "5.2.4.", "6.2.4.", "4.2.5.", "5.2.5.", "6.2.5.", "4.2.6.", "5.2.6.", "6.2.6.", "4.2.7.", "4.2.8.", "1.5.1.", "2.5.1.", "3.5.1.", "4.5.1.", "5.5.1.", "6.5.1.", "1.5.2.", "2.5.2.", "3.5.2.", "4.5.2.", "5.5.2.", "6.5.2.", "1.5.3.", "2.5.3.", "3.5.3.", "4.5.3.", "5.5.3.", "6.5.3.", "2.5.4.", "3.5.4.", "4.5.4.", "5.5.4.", "6.5.4.", "4.5.5.", "5.5.5.", "6.5.5.", "5.5.6.", "6.5.6.", "5.5.7.", "6.5.7.", "5.5.8.", "6.5.8."]
        combined_standards_field = receptive_fields + productive_fields + interactive_fields

        nrs_string_vars = {field: StringVar(value=0) for field in nrs_fields}
        self.receptive_string_vars = {field: StringVar(value=0) for field in receptive_fields}
        self.productive_string_vars = {field: StringVar(value=0) for field in productive_fields}
        self.interactive_string_vars = {field: StringVar(value=0) for field in interactive_fields}

        # Define checkbuttons
        self.nrs_checkbuttons = [
            ("1", nrs_string_vars["1"]),
            ("2", nrs_string_vars["2"]),
            ("3", nrs_string_vars["3"]),
            ("4", nrs_string_vars["4"]),
            ("5", nrs_string_vars["5"]),
            ("6", nrs_string_vars["6"]),
            ("6+", nrs_string_vars["6+"])
        ]

        receptive_checkbuttons = [
            ("Identify a few key words and phrases from read alouds, visual images, and oral presentations using a very limited set of strategies, with prompting and support.", self.receptive_string_vars["1.1.1."]),
            ("Identify a few keywords and phrases in oral communications and simple spoken and written texts using a very limited set of strategies.", self.receptive_string_vars["2.1.1."]),
            ("Identify the main topic in oral presentations and simple spoken and written texts using an emerging set of strategies.", self.receptive_string_vars["3.1.1."]),
            ("Determine a central idea or theme in oral presentations and spoken and written texts using a developing set of strategies.", self.receptive_string_vars["4.1.1."]),
            ("Determine a central idea or theme in oral presentations and spoken and written texts using an increasing range of strategies.", self.receptive_string_vars["5.1.1."]),
            ("Determine central ideas or themes in oral presentations and spoken and written texts using a wide range of strategies.", self.receptive_string_vars["6.1.1."]),
            ("Retell a few key details using an emerging set of strategies.", self.receptive_string_vars["3.1.2."]),
            ("Retell key details using a developing set of strategies.", self.receptive_string_vars["4.1.2."]),
            ("Analyze the development of the themes/ideas using an increasing range of strategies.", self.receptive_string_vars["5.1.2."]),
            ("Analyze the development of the themes/ideas using a wide range of strategies.", self.receptive_string_vars["6.1.2."]),
            ("Answer questions about key details using a developing set of strategies.", self.receptive_string_vars["4.1.3."]),
            ("Cite specific details and evidence from texts to support the analysis using an increasing range of strategies.", self.receptive_string_vars["5.1.3."]),
            ("Cite specific details and evidence from texts to support the analysis using a wide range of strategies.", self.receptive_string_vars["6.1.3."]),
            ("Explain how the theme is developed by specific details in texts using a developing set of strategies.", self.receptive_string_vars["4.1.4."]),
            ("Summarize a text using an increasing range of strategies.", self.receptive_string_vars["5.1.4."]),
            ("Summarize a text using a wide range of strategies.", self.receptive_string_vars["6.1.4."]),
            ("Summarize part of a text using a developing set of strategies.", self.receptive_string_vars["4.1.5."]),
            ("Identify a point an author or a speaker makes, with support.", self.receptive_string_vars["2.6.1."]),
            ("Identify the main argument an author or speaker makes, with support.", self.receptive_string_vars["3.6.1."]),
            ("Explain the reasons an author or a speaker gives to support a claim, with support.", self.receptive_string_vars["4.6.1."]),
            ("Analyze the reasoning in persuasive spoken and written texts.", self.receptive_string_vars["5.6.1."]),
            ("Analyze and evaluate the reasoning in persuasive spoken and written texts.", self.receptive_string_vars["6.6.1."]),
            ("Identify one reason an author or a speaker gives to support the argument, with support.", self.receptive_string_vars["3.6.2."]),
            ("Identify one or two reasons an author or a speaker gives to support the main point, with support.", self.receptive_string_vars["4.6.2."]),
            ("Determine whether the evidence is sufficient to support the claim.", self.receptive_string_vars["5.6.2."]),
            ("Determine whether the evidence is sufficient to support the claim.", self.receptive_string_vars["6.6.2."]),
            ("Cite textual evidence to support the analysis.", self.receptive_string_vars["5.6.3."]),
            ("Cite specific textual evidence to thoroughly support the analysis.", self.receptive_string_vars["6.6.3."]),
            ("Recognize the meaning of some words learned through conversations, reading, and being read to.", self.receptive_string_vars["1.7.2."]),
            ("Recognize the meaning of some words learned through conversations, reading, and being read to.", self.receptive_string_vars["2.7.2."]),
            ("Recognize the meaning of a few frequently occurring words and phrases in simple oral presentations and read alouds about familiar topics, experiences, or events, with prompting and support.", self.receptive_string_vars["1.8.1."]),
            ("Recognize the meaning of a few frequently occurring words, simple phrases, and formulaic expressions in spoken and written texts about familiar topics, experiences, or events, relying heavily on context, questioning, and knowledge of morphology in their native language(s).", self.receptive_string_vars["2.8.1."]),
            ("Determine the meaning of frequently occurring words, phrases, and expressions in spoken and written texts about familiar topics, experiences, or events, using context, questioning, and knowledge of morphology in their native language(s).", self.receptive_string_vars["3.8.1."]),
            ("Determine the meaning of general academic and content-specific words and phrases and frequently occurring expressions in spoken and written texts about familiar topics, experiences, or events using context, questioning, and a developing knowledge of English and their native language(s)' morphology.", self.receptive_string_vars["4.8.1."]),
            ("Determine the meaning of general academic and content-specific words and phrases, figurative and connotative language, and a growing number of idiomatic expressions in spoken and written texts about a variety of topics, experiences, or events using context, questioning, and an increasing knowledge of English morphology.", self.receptive_string_vars["5.8.1."]),
            ("Determine the meaning of general academic and content-specific words and phrases, figurative and connotative language, and idiomatic expressions in spoken and written texts about a variety of topics, experiences, or events using context, questioning, and consistent knowledge of English morphology.", self.receptive_string_vars["6.8.1."])
        ]

        productive_checkbuttons = [
            ("Communicate information and feelings about familiar texts, topics, and experiences, with prompting and support.", self.productive_string_vars["1.3.1."]),
            ("Communicate information and feelings about familiar texts, topics, and experiences, with support.", self.productive_string_vars["2.3.1."]),
            ("Deliver short oral presentations about familiar texts, topics, experiences, or events, with support.", self.productive_string_vars["3.3.1."]),
            ("Deliver short oral presentations about familiar texts, topics, or events, with support.", self.productive_string_vars["4.3.1."]),
            ("Deliver oral presentations about a variety of texts, topics, or events.", self.productive_string_vars["5.3.1."]),
            ("Deliver oral presentations about a variety of texts, topics, or events.", self.productive_string_vars["6.3.1."]),
            ("Compose simple written narratives or informational texts about familiar texts, topics, experiences, or events, with support.", self.productive_string_vars["3.3.2."]),
            ("Compose written informational texts about familiar texts, topics, or events, with support.", self.productive_string_vars["4.3.2."]),
            ("Compose written informational texts about a variety of texts, topics, or events.", self.productive_string_vars["5.3.2."]),
            ("Compose written informational texts about a variety of texts, topics, or events.", self.productive_string_vars["6.3.2."]),
            ("Develop the topic with a few details about familiar texts, topics, or events, with support.", self.productive_string_vars["4.3.3."]),
            ("Develop the topic with some relevant details, concepts, examples, and information about a variety of texts, topics, or events.", self.productive_string_vars["5.3.3."]),
            ("Develop the topic fully with relevant details, concepts, examples, and information about a variety of texts, topics, or events.", self.productive_string_vars["6.3.3."]),
            ("Integrate graphics or multimedia when useful about a variety of texts, topics, or events.", self.productive_string_vars["5.3.4."]),
            ("Integrate graphics or multimedia when useful about a variety of texts, topics, or events.", self.productive_string_vars["6.3.4."]),
            ("Express a preference or opinion about a familiar topic.", self.productive_string_vars["1.4.1."]),
            ("Express an opinion about a familiar topic, experience or event.", self.productive_string_vars["2.4.1."]),
            ("Construct a claim about familiar topics, experiences, or events.", self.productive_string_vars["3.4.1."]),
            ("Construct a claim about familiar topics.", self.productive_string_vars["4.4.1."]),
            ("Construct a claim about a variety of topics.", self.productive_string_vars["5.4.1."]),
            ("Construct a substantive claim about a variety of topics.", self.productive_string_vars["6.4.1."]),
            ("Give a reason for the opinion.", self.productive_string_vars["2.4.2."]),
            ("Introduce the topic, experience, or event.", self.productive_string_vars["3.4.2."]),
            ("Introduce the topic.", self.productive_string_vars["4.4.2."]),
            ("Introduce the topic.", self.productive_string_vars["5.4.2."]),
            ("Introduce the claim.", self.productive_string_vars["6.4.2."]),
            ("Give a reason to support the claim.", self.productive_string_vars["3.4.3."]),
            ("Provide sufficient reasons or facts to support the claim.", self.productive_string_vars["4.4.3."]),
            ("Provide logically ordered reasons or facts that effectively support the claim.", self.productive_string_vars["5.4.3."]),
            ("Distinguish the claim from a counter-claim.", self.productive_string_vars["6.4.3."]),
            ("Provide a concluding statement.", self.productive_string_vars["3.4.4."]),
            ("Provide a concluding statement.", self.productive_string_vars["4.4.4."]),
            ("Provide a concluding statement.", self.productive_string_vars["5.4.4."]),
            ("Provide logically ordered and relevant reasons and evidence to support the claim and to refute the counter-claim.", self.productive_string_vars["6.4.4."]),
            ("Provide a conclusion that summarizes the argument presented.", self.productive_string_vars["6.4.5."]),
            ("Show limited awareness of differences between informal and formal language use.", self.productive_string_vars["1.7.1."]),
            ("Show emerging awareness of differences between informal and formal language use.", self.productive_string_vars["2.7.1."]),
            ("Show increasing awareness of differences between informal and formal language use.", self.productive_string_vars["3.7.1."]),
            ("Adapt language choices and style according to purpose, task, and audience with developing ease in various social and academic contexts.", self.productive_string_vars["4.7.1."]),
            ("Adapt language choices and style according to purpose, task, and audience in various social and academic contexts.", self.productive_string_vars["5.7.1."]),
            ("Adapt language choices and style according to purpose, task, and audience with ease in various social and academic contexts.", self.productive_string_vars["6.7.1."]),
            ("Adapt language choices to task and audience with emerging control in various social and academic contexts.", self.productive_string_vars["3.7.2."]),
            ("Use an increasing number of general academic and content-specific words and expressions in spoken and written texts.", self.productive_string_vars["4.7.2."]),
            ("Use a wider range of complex general academic and content-specific words and phrases.", self.productive_string_vars["5.7.2."]),
            ("Use a wide variety of complex general academic and content-specific words and phrases.", self.productive_string_vars["6.7.2."]),
            ("Begin to use some frequently occurring general academic and content-specific words.", self.productive_string_vars["3.7.3."]),
            ("Show developing control of style and tone in spoken and written texts.", self.productive_string_vars["4.7.3."]),
            ("Adopt and maintain a formal and informal style and tone in spoken and written texts, as appropriate.", self.productive_string_vars["5.7.3."]),
            ("Employ both formal and more informal styles and tones effectively in spoken and written texts, as appropriate.", self.productive_string_vars["6.7.3."]),
            ("Use a narrow range of vocabulary and syntactically simple sentences, with support.", self.productive_string_vars["1.9.1."]),
            ("Communicate basic information about an event or topic, with support.", self.productive_string_vars["2.9.1."]),
            ("Recount a short sequence of events in order, with support.", self.productive_string_vars["3.9.1."]),
            ("Recount a sequence of events, with a beginning, middle, and end, with support.", self.productive_string_vars["4.9.1."]),
            ("Recount a longer, more detailed sequence of events or steps in a process, with a clear sequential or chronological structure.", self.productive_string_vars["5.9.1."]),
            ("Recount a complex and detailed sequence of events or steps in a process, with an effective sequential or chronological order.", self.productive_string_vars["6.9.1."]),
            ("Use a narrow range of vocabulary and syntactically simple sentences, with support.", self.productive_string_vars["2.9.2."]),
            ("Introduce an informational topic, with support.", self.productive_string_vars["3.9.2."]),
            ("Introduce and develop an informational topic with facts and details, with support.", self.productive_string_vars["4.9.2."]),
            ("Introduce and develop an informational topic with facts, details, and evidence.", self.productive_string_vars["5.9.2."]),
            ("Introduce and effectively develop an informational topic with facts, details, and evidence.", self.productive_string_vars["6.9.2."]),
            ("Provide one or two facts about the topic, with support.", self.productive_string_vars["3.9.3."]),
            ("Use common transitional words and phrases to connect events, ideas, and opinions, with support.", self.productive_string_vars["4.9.3."]),
            ("Use a variety of more complex transitions to link the major sections of speech and text and to clarify relationships among events and ideas.", self.productive_string_vars["5.9.3."]),
            ("Use complex and varied transitions to link the major sections of speech and text and to clarify relationships among events and ideas.", self.productive_string_vars["6.9.3."]),
            ("Use common linking words to connect events and ideas, with support.", self.productive_string_vars["3.9.4."]),
            ("Provide a conclusion, with support.", self.productive_string_vars["4.9.4."]),
            ("Provide a concluding section or statement.", self.productive_string_vars["5.9.4."]),
            ("Provide a concluding section or statement.", self.productive_string_vars["6.9.4."]),
            ("Recognize and use a small number of frequently occurring nouns and verbs, with support.", self.productive_string_vars["1.10.1."]),
            ("Recognize and use a small number of frequently occurring nouns, noun phrases, verbs, conjunctions, and prepositions, with support.", self.productive_string_vars["2.10.1."]),
            ("Use frequently occurring verbs, nouns, adjectives, adverbs, prepositions, and conjunctions, with support.", self.productive_string_vars["3.10.1."]),
            ("Use simple phrases, with support.", self.productive_string_vars["4.10.1."]),
            ("Use increasingly complex phrases.", self.productive_string_vars["5.10.1."]),
            ("Use complex phrases and clauses.", self.productive_string_vars["6.10.1."]),
            ("Understand and respond to simple questions, with support.", self.productive_string_vars["1.10.2."]),
            ("Understand and respond to simple questions, with support.", self.productive_string_vars["2.10.2."]),
            ("Produce simple and compound sentences, with support.", self.productive_string_vars["3.10.2."]),
            ("Use simple clauses, with support.", self.productive_string_vars["4.10.2."]),
            ("Use increasingly complex clauses.", self.productive_string_vars["5.10.2."]),
            ("Produce and expand simple, compound, and complex sentences.", self.productive_string_vars["6.10.2."]),
            ("Produce and expand simple, compound, and a few complex sentences, with support.", self.productive_string_vars["4.10.3."]),
            ("Produce and expand simple, compound, and complex sentences.", self.productive_string_vars["5.10.3."])
        ]

        interactive_checkbuttons = [
            ("Participate in short conversations and written exchanges about familiar topics and in familiar contexts, with limited involvement.", self.interactive_string_vars["1.2.1."]),
            ("Listen actively to others.", self.interactive_string_vars["2.2.1."]),
            ("Participate in short conversations and written exchanges about familiar topics and texts.", self.interactive_string_vars["3.2.1."]),
            ("Participate in conversations, discussions, and written exchanges about familiar topics, texts, and issues.", self.interactive_string_vars["4.2.1."]),
            ("Participate in conversations, discussions, and written exchanges about a range of topics, texts, and issues.", self.interactive_string_vars["5.2.1."]),
            ("Participate in conversations, extended discussions, and written exchanges about a range of substantive topics, texts, and issues.", self.interactive_string_vars["6.2.1."]),
            ("Respond to simple yes/no questions and some wh- questions, with limited involvement.", self.interactive_string_vars["1.5.1."]),
            ("Participate in short conversations and written exchanges about familiar topics and in familiar contexts.", self.interactive_string_vars["2.5.1."]),
            ("Present information and ideas.", self.interactive_string_vars["3.5.1."]),
            ("Build on the ideas of others.", self.interactive_string_vars["4.5.1."]),
            ("Build on the ideas of others.", self.interactive_string_vars["5.5.1."]),
            ("Build on the ideas of others.", self.interactive_string_vars["6.5.1."]),
            ("Present simple information.", self.interactive_string_vars["1.2.2."]),
            ("Take turns appropriately in interactions with others.", self.interactive_string_vars["2.2.2."]),
            ("Express his or her own ideas.", self.interactive_string_vars["3.2.2."]),
            ("Express his or her own ideas.", self.interactive_string_vars["4.2.2."]),
            ("Express his or her own ideas clearly and persuasively.", self.interactive_string_vars["5.2.2."]),
            ("Respond to simple yes/no questions and some wh- questions.", self.interactive_string_vars["6.2.2."]),
            ("Respond to simple questions and wh- questions.", self.interactive_string_vars["1.5.2."]),
            ("Ask and answer relevant questions.", self.interactive_string_vars["2.5.2."]),
            ("Support points clearly with specific and relevant evidence.", self.interactive_string_vars["3.5.2."]),
            ("Refer to specific and relevant evidence from texts or research to support his or her ideas.", self.interactive_string_vars["4.5.2."]),
            ("Add relevant information and evidence.", self.interactive_string_vars["5.5.2."]),
            ("Ask and answer questions to clarify ideas and conclusions.", self.interactive_string_vars["6.5.2."]),
            ("Ask and answer questions that probe reasoning and claims.", self.interactive_string_vars["2.2.3."]),
            ("Restate some of the key ideas expressed.", self.interactive_string_vars["3.2.3."]),
            ("Summarize the key points expressed.", self.interactive_string_vars["4.2.3."]),
            ("Summarize the key points and evidence discussed.", self.interactive_string_vars["5.2.3."]),
            ("Follow rules for discussion.", self.interactive_string_vars["6.2.3."]),
            ("Ask questions to gain information or clarify understanding.", self.interactive_string_vars["1.5.3."]),
            ("Participate in short, shared research projects, with prompting and support.", self.interactive_string_vars["2.5.3."]),
            ("Carry out short, shared research projects, with support.", self.interactive_string_vars["3.5.3."]),
            ("Carry out short individual or shared research projects, with support.", self.interactive_string_vars["4.5.3."]),
            ("Carry out short research projects to answer a question, with support.", self.interactive_string_vars["5.5.3."]),
            ("Carry out both short and more sustained research projects to answer a question.", self.interactive_string_vars["6.5.3."]),
            ("Carry out both short and more sustained research projects to answer a question or solve a problem.", self.interactive_string_vars["2.2.4."]),
            ("Gather information from a few provided sources, with prompting and support.", self.interactive_string_vars["3.2.4."]),
            ("Gather information from a few provided print and digital sources, with support.", self.interactive_string_vars["4.2.4."]),
            ("Gather information from provided print and digital sources, with support.", self.interactive_string_vars["5.2.4."]),
            ("Gather information from multiple provided print and digital sources, with support.", self.interactive_string_vars["6.2.4."]),
            ("Gather information from multiple print and digital sources.", self.interactive_string_vars["2.5.4."]),
            ("Gather information from multiple print and digital sources.", self.interactive_string_vars["3.5.4."]),
            ("Label some key information, with prompting and support.", self.interactive_string_vars["4.5.4."]),
            ("Label collected information, experiences, or events, with support.", self.interactive_string_vars["5.5.4."]),
            ("Record information in simple notes, with support.", self.interactive_string_vars["6.5.4."]),
            ("Paraphrase key information in a short written or oral report, with support.", self.interactive_string_vars["4.2.5."]),
            ("Evaluate the reliability of each source.", self.interactive_string_vars["5.2.5."]),
            ("Evaluate the reliability of each source.", self.interactive_string_vars["6.2.5."]),
            ("Recall information from experience or from a provided source, with support.", self.interactive_string_vars["4.5.5."]),
            ("Summarize data and information, with support.", self.interactive_string_vars["5.5.5."]),
            ("Include illustrations, diagrams, or other graphics as appropriate, with support.", self.interactive_string_vars["6.5.5."]),
            ("Use search terms effectively.", self.interactive_string_vars["4.2.6."]),
            ("Use advanced search terms effectively.", self.interactive_string_vars["5.2.6."]),
            ("Provide a list of sources, with support.", self.interactive_string_vars["6.2.6."]),
            ("Synthesize information from multiple print and digital sources.", self.interactive_string_vars["5.5.6."]),
            ("Synthesize information from multiple print and digital sources.", self.interactive_string_vars["6.5.6."]),
            ("Integrate information into an organized oral or written report.", self.interactive_string_vars["4.2.7."]),
            ("Analyze and integrate information into clearly organized spoken and written texts.", self.interactive_string_vars["5.5.7."]),
            ("Include illustrations, diagrams, or other graphics as appropriate.", self.interactive_string_vars["6.5.7."]),
            ("Include illustrations, diagrams, or other graphics as appropriate.", self.interactive_string_vars["4.2.8."]),
            ("Cite sources appropriately.", self.interactive_string_vars["5.5.8."]),
            ("Cite sources appropriately.", self.interactive_string_vars["6.5.8."]),
        ]

        # Assign variables to root for easier access later
        for field in nrs_string_vars:
            setattr(self.root, field, nrs_string_vars[field])
        for field in self.receptive_string_vars:
            setattr(self.root, field, self.receptive_string_vars[field])
        for field in self.productive_string_vars:
            setattr(self.root, field, self.productive_string_vars[field])
        for field in self.interactive_string_vars:
            setattr(self.root, field, self.interactive_string_vars[field])

        # Create scrollable canvases for each frame
        receptive_canvas, receptive_scrollable_frame = self.create_scrollable_canvas(receptive_frame, width=screen_width - 50, height=400)
        productive_canvas, productive_scrollable_frame = self.create_scrollable_canvas(productive_frame, width=screen_width - 50, height=400)
        interactive_canvas, interactive_scrollable_frame = self.create_scrollable_canvas(interactive_frame, width=screen_width - 50, height=400)

        # For program_info_frame and reflection_frame, directly use the frames without scrollbars
        program_info_scrollable_frame = program_info_frame
        reflection_scrollable_frame = reflection_frame

        # Track the current tab and update bindings
        def update_scroll_binding(event):
            current_tab = notebook.index(notebook.select())
            if current_tab == 1:
                self.root.bind_all("<MouseWheel>", lambda e: self.on_mouse_wheel(e, receptive_canvas))
            elif current_tab == 2:
                self.root.bind_all("<MouseWheel>", lambda e: self.on_mouse_wheel(e, productive_canvas))
            elif current_tab == 3:
                self.root.bind_all("<MouseWheel>", lambda e: self.on_mouse_wheel(e, interactive_canvas))

        # Bind the tab change event to update the scroll binding
        notebook.bind("<<NotebookTabChanged>>", update_scroll_binding)

        # Initially bind to the first tab
        update_scroll_binding(None)

        def highlight_range(start, end):
            """Highlight the range of dates between start and end."""
            clear_highlight()  # Clear any previous highlights
            current_date = start
            while current_date <= end:
                self.cal.calevent_create(current_date, '', 'highlight')
                current_date += timedelta(days=1)

            # Configure highlight appearance
            self.cal.tag_config('highlight', background='blue', foreground='white')

        def clear_highlight():
            """Clear highlighted dates."""
            for event in self.cal.get_calevents(tag='highlight'):
                self.cal.calevent_remove(event)

        def remove_leading_zeros(date_str):
            """Remove leading zeros from a date string in the format mm/dd/yyyy."""
            # Split the date string and convert each part to an integer to remove leading zeros
            month, day, year = date_str.split('/')
            return f"{int(month)}/{int(day)}/{year}"

        # Create widgets for program information
        for i, (label_text, string_var) in enumerate(self.header_string_vars.items()):
            if label_text == "NRS EFL(s)":
                label = ttk.Label(program_info_scrollable_frame, text=self.titlecase(label_text), style="Bold.TLabel", font=self.default_font)
                label.grid(row=i, column=1, padx=10, pady=5, sticky='e')

                for j, (text, var) in enumerate(self.nrs_checkbuttons):
                    cb = Checkbutton(program_info_scrollable_frame, text=text, variable=var, font=self.default_font)
                    cb.grid(row=i, column=j + 2, padx=5, pady=5, sticky='w')
            elif label_text == "TIME FRAME":
                # Create the calendar widget
                label = ttk.Label(program_info_scrollable_frame, text="Lesson Date Range:", style="Bold.TLabel", font=self.default_font)
                label.grid(row=i, column=0, padx=10, pady=5, sticky='e')
                large_font = font.Font(family="Helvetica", size=14, weight="bold")
                self.cal = Calendar(program_info_scrollable_frame, selectmode='day',
                                    year=datetime.today().year, month=datetime.today().month,
                                    day=datetime.today().day, font=large_font,
                                    headersbackground="lightblue",
                                    background="lightgrey", foreground="black",
                                    selectbackground="blue", selectforeground="white")
                self.cal.grid(row=i+1, column=0, columnspan=2, padx=10, pady=5)

                # Variables to store start and end dates
                self.start_date = None
                self.end_date = None

                # StringVar to hold the range string
                self.range_var = StringVar()
                self.range_var.set("Select two dates to form a range.")
                # self.range_var.set()

                def on_date_click(event):
                    global start_date, end_date
                    selected_date = self.cal.get_date()
                    date_obj = datetime.strptime(selected_date, "%m/%d/%y")

                    if not self.start_date:
                        # Select the start date
                        self.start_date = date_obj
                        self.range_var.set(f"Start Date Selected: {remove_leading_zeros(selected_date)}")
                    elif not self.end_date:
                        # Select the end date
                        self.end_date = date_obj
                        if self.end_date < self.start_date:
                            self.start_date, self.end_date = self.end_date, self.start_date
                        highlight_range(self.start_date, self.end_date)
                        start_date_str = self.start_date.strftime("%m/%d/%Y")
                        end_date_str = self.end_date.strftime("%m/%d/%Y")
                        self.range_var.set(f"Selected Date Range: {remove_leading_zeros(start_date_str)} - {remove_leading_zeros(end_date_str)}")
                    else:
                        # Reset selection if both dates are already selected
                        self.start_date = date_obj
                        self.end_date = None
                        clear_highlight()
                        self.range_var.set(f"Start Date Selected: {remove_leading_zeros(selected_date)}")
                    check_fields()

                # Add Label with larger font using StringVar
                # date_label = Label(program_info_scrollable_frame, textvariable=self.range_var, font=large_font, width=35)
                date_label = ttk.Label(program_info_scrollable_frame, textvariable=self.range_var, font=self.default_font, width=40)
                date_label.grid(row=i + 2, column=0, columnspan=2, padx=20, pady=10, sticky='w')

                # Bind click event to calendar widget
                self.cal.bind("<<CalendarSelected>>", on_date_click)

            else:
                label = ttk.Label(program_info_scrollable_frame, text=self.titlecase(label_text), style="Bold.TLabel", font=self.default_font)
                entry = Entry(program_info_scrollable_frame, textvariable=string_var, font=self.default_font)

                label.grid(row=i, column=0, padx=10, pady=5, sticky='nsew')
                entry.grid(row=i, column=1, padx=10, pady=5, sticky='nswe')

        # Method to check if all required fields are filled
        def check_fields():
            # Check if all header fields are filled
            headers_filled = all(var.get().strip() for key, var in self.header_string_vars.items() if key != 'NRS EFL(s)' and key != 'TIME FRAME')
            # Check if a date range is selected
            date_range_selected = self.start_date is not None and self.end_date is not None
            # Ensure all necessary conditions are met
            if headers_filled and date_range_selected:
                save_button.config(state='normal')
            else:
                save_button.config(state='disabled')

        # Bind the trace to all header fields to check conditions
        for var in self.header_string_vars.values():
            var.trace('w', lambda name, index, mode, var=var: check_fields())

        # Create the save button
        save_button = Button(self.root, text="Save", command=self.save_document, state='disabled', font=self.default_font)
        save_button.grid(row=1, column=0, pady=10)
        
        # Create a bold font
        bold_font = tkFont.Font(weight="bold")

        # Define the style for the bold labels
        style = ttk.Style()
        style.configure("Bold.TLabel", font=bold_font)

        # Function to create sections with checkbuttons for standards
        def create_standards_section(scrollable_frame, label_text, row_start, checkbutton_indices, checkbuttons, max_width=screen_width - 150, section_name=None):
            # Create the section label
            section_label = ttk.Label(scrollable_frame, text=label_text, style="Bold.TLabel", wraplength=max_width, font=self.default_font)
            if section_name == "interactive":
                section_label.grid(row=row_start, column=1, columnspan=2, padx=10, pady=5, sticky='w')
            else:
                section_label.grid(row=row_start, column=0, columnspan=2, padx=10, pady=5, sticky='w')

            for i, idx in enumerate(checkbutton_indices, start=1):
                # Get the checkbutton text and variable
                text, var = checkbuttons[idx]
                text = " ".join((combined_standards_field[self.c], text))

                # Create the checkbutton without text
                cb = Checkbutton(scrollable_frame, variable=var, font=self.default_font)
                cb.grid(row=row_start + i, column=0, padx=5, pady=5, sticky='w')

                # Create a label with wrapped text next to the checkbutton
                label = ttk.Label(scrollable_frame, text=text, wraplength=max_width, font=self.default_font)
                label.grid(row=row_start + i, column=1, padx=10, pady=5, sticky='w')
                self.c+=1

        create_standards_section(
            receptive_scrollable_frame,
            "1. Construct meaning from oral presentations and literary and informational text through level-appropriate listening, reading, and viewing.",
            0,
            range(0, 17),
            receptive_checkbuttons
        )

        create_standards_section(
            receptive_scrollable_frame,
            "6. Analyze and critique the arguments of others orally and in writing.",
            20,
            range(17, 28),
            receptive_checkbuttons
        )

        create_standards_section(
            receptive_scrollable_frame,
            "7. Adapt language choices to purpose, task, and audience when speaking and writing.",
            40,
            range(28, 30),
            receptive_checkbuttons
        )

        create_standards_section(
            receptive_scrollable_frame,
            "8. Determine the meaning of words and phrases in oral presentations and literary and informational text.",
            50,
            range(30, len(receptive_checkbuttons)),
            receptive_checkbuttons
        )

        create_standards_section(
            productive_scrollable_frame,
            "3. Speak and write about level-appropriate complex literary and informational texts and topics.",
            0,
            range(0, 15),
            productive_checkbuttons
        )

        create_standards_section(
            productive_scrollable_frame,
            "4. Construct level-appropriate oral and written claims and support them with reasoning and evidence.",
            20,
            range(15, 35),
            productive_checkbuttons
        )

        create_standards_section(
            productive_scrollable_frame,
            "7. Adapt language choices to purpose, task, and audience when speaking and writing.",
            41,
            range(35, 49),
            productive_checkbuttons
        )

        create_standards_section(
            productive_scrollable_frame,
            "9. Create clear and coherent level-appropriate speech and text.",
            60,
            range(49, 68),
            productive_checkbuttons
        )

        create_standards_section(
            productive_scrollable_frame,
            "10. Demonstrate command of the conventions of standard English to communicate in level-appropriate speech and writing.",
            81,
            range(68, len(productive_checkbuttons)),
            productive_checkbuttons
        )

        create_standards_section(
            interactive_scrollable_frame,
            "2. Participate in level-appropriate oral and written exchanges of information, ideas, and analyses, in various social and academic contexts, responding to peer, audience, or reader comments and questions.",
            0,
            range(0, 30),
            interactive_checkbuttons,
            section_name="interactive"
        )

        create_standards_section(
            interactive_scrollable_frame,
            "5. Conduct research and evaluate and communicate findings to answer questions or solve problems.",
            40,
            range(30, len(interactive_checkbuttons)),
            interactive_checkbuttons,
            section_name="interactive"
        )

        # Create text entries for the footers fields in the reflection tab
        def create_footer_entries(reflection_frame, footers):
            text_widgets = {}
            for i, (label_text, _) in enumerate(footers.items()):
                # Create a label using titlecase
                label = ttk.Label(reflection_frame, text=self.titlecase(label_text), style="Bold.TLabel", font=self.default_font)
                label.grid(row=i, column=0, padx=10, pady=5, sticky='nw')

                # Create a text widget for each footer field
                text_entry = Text(reflection_frame, wrap=WORD, height=5, width=50)
                text_entry.grid(row=i, column=1, padx=10, pady=5, sticky='nsew')

                # Configure grid weights for resizing
                reflection_frame.grid_rowconfigure(i, weight=1)
                reflection_frame.grid_columnconfigure(1, weight=1)

                text_widgets[label_text] = text_entry

            return text_widgets

        # Call the function to create footer entries in the reflection tab
        self.footer_text_widgets = create_footer_entries(reflection_frame, self.footers)
        # # Method to check if all required fields are filled
        # def check_fields():
        #     # Check if all header fields are filled
        #     headers_filled = all(var.get().strip() for key, var in self.header_string_vars.items() if key != 'NRS EFL(s)' and key != 'TIME FRAME')
        #     # Check if a date range is selected
        #     date_range_selected = self.start_date is not None and self.end_date is not None
        #     # Ensure all necessary conditions are met
        #     if headers_filled and date_range_selected:
        #         save_button.config(state='normal')
        #     else:
        #         save_button.config(state='disabled')

        # # Bind the trace to all header fields to check conditions
        # for var in self.header_string_vars.values():
        #     var.trace('w', lambda name, index, mode, var=var: check_fields())

        # # Create the save button
        # save_button = ttk.Button(self.root, text="Save", command=self.save_document, state='disabled')
        # save_button.grid(row=1, column=0, pady=10)
        
    # Function to create a canvas with vertical scrollbars for a given frame
    def create_scrollable_canvas(self, frame, width, height, frame_name=None):
        # Check if the frame is the reflection_frame
        if frame_name in ['reflection_frame', 'program_info']:
            # Directly return the frame without scrollbars
            return frame

        # Create a canvas and vertical scrollbar
        canvas = Canvas(frame, width=width, height=height)
        vscrollbar = ttk.Scrollbar(frame, orient=VERTICAL, command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        # Bind configuration changes to update the canvas scroll region
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all"),
                width=width - 20  # Ensure it extends to the scrollbar
            )
        )

        # Add scrolling functionality
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw", width=width - 20)
        canvas.configure(yscrollcommand=vscrollbar.set)

        # Grid placement for canvas and scrollbar
        canvas.grid(row=0, column=0, sticky="nsew")
        vscrollbar.grid(row=0, column=1, sticky="ns")

        # Configure grid weights for proper resizing
        frame.grid_rowconfigure(0, weight=1)
        frame.grid_columnconfigure(0, weight=1)

        return canvas, scrollable_frame
    
    # Function to handle mouse wheel events
    def on_mouse_wheel(self, event, canvas):
        # Determine platform-specific scroll speed
        if event.delta:
            # Windows
            scroll_speed = int(event.delta / 120)
        else:
            # Mac, using event.num
            scroll_speed = event.num if event.num != 0 else 0

        # Reverse scroll direction to match natural scrolling behavior
        scroll_speed = -scroll_speed

        if scroll_speed == 0:
            scroll_speed = -1 if event.delta < 0 else 1

        # Apply scroll to canvas
        canvas.yview_scroll(scroll_speed, "units")

        # Debug print statements (optional)
        # print(f"Mouse Wheel Event: delta={event.delta}, num={event.num}")
        # print(f"Scrolled by: {scroll_speed} units")


    # Function to collect NRS values from checkboxes
    def collect_nrs_values(self):
        nrs_selected_values = [text for text, var in self.nrs_checkbuttons if var.get()!="0"]
        self.headers["NRS EFL(s)"] = ", ".join(nrs_selected_values)

    def parse_time_frame(self):
        try:
            # Strip the time_frame string to remove leading/trailing spaces
            time_frame = self.range_var.get()[self.range_var.get().index('Selected Date Range: ')+len('Selected Date Range: '):].strip()
            # Split on dash, allowing for spaces, and strip each part to remove spaces
            start_date, end_date = map(str.strip, re.split(r'\s*-\s*', time_frame))
            start_date = datetime.strptime(start_date, "%m/%d/%Y")
            end_date = datetime.strptime(end_date, "%m/%d/%Y")
            # Format dates with hyphens between month, day, and year, underscores between different dates
            formatted_start = f"{start_date.month}-{start_date.day}-{start_date.year % 100}"
            formatted_end = f"{end_date.month}-{end_date.day}-{end_date.year % 100}"
            return f"{formatted_start}_{formatted_end}"
        except ValueError:
            # Return an empty string if parsing fails
            return ""

    # Function to save the document
    def save_document(self):
        self.collect_nrs_values()

        # Get the values from header entries
        for key, var in self.header_string_vars.items():
            if key not in ["NRS EFL(s)", "TIME FRAME"]:
                if var.get():
                    self.headers[key] = var.get()

        # Get the values from standards checkbuttons
        for key, var in self.receptive_string_vars.items():
            level = int(key[2])
            if level == 1:
                if var.get() and var.get()!="0":
                    self.standards["[Receptive 1]"] += key + "\n"
            elif level == 6:
                if var.get() and var.get()!="0":
                    self.standards["[Receptive 6]"] += key + "\n"
            elif level == 7:
                if var.get() and var.get()!="0":
                    self.standards["[Receptive 7]"] += key + "\n"
            elif level == 8:
                if var.get() and var.get()!="0":
                    self.standards["[Receptive 8]"] += key + "\n"

        # Get the values from standards checkbuttons
        for key, var in self.productive_string_vars.items():
            level = int(key[2])
            if level == 3:
                if var.get() and var.get()!="0":
                    self.standards["[Productive 3]"] += key + "\n"
            elif level == 4:
                if var.get() and var.get()!="0":
                    self.standards["[Productive 4]"] += key + "\n"
            elif level == 7:
                if var.get() and var.get()!="0":
                    self.standards["[Productive 7]"] += key + "\n"
            elif level == 9:
                if var.get() and var.get()!="0":
                    self.standards["[Productive 9]"] += key + "\n"
            elif level == 1:
                if var.get() and var.get()!="0":
                    self.standards["[Productive 10]"] += key + "\n"

        for key, var in self.interactive_string_vars.items():
            level = int(key[2])
            if level == 2:
                if var.get() and var.get()!="0":
                    self.standards["[Interactive 2]"] += key + "\n"
            if level == 5:
                if var.get() and var.get()!="0":
                    self.standards["[Interactive 5]"] += key + "\n"

        # Get the values from text widgets in the reflection frame
        for key, text_widget in self.footer_text_widgets.items():
            self.footers[key] = text_widget.get("1.0", END).strip()

        # Parse and format the time frame, or omit if parsing fails
        formatted_time_frame = self.parse_time_frame()

        used_keys = []
        # Iterate over all tables and cells to replace placeholders with user input
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in self.headers.items():
                        if key in cell.text and key not in used_keys:
                            # Clear cell text
                            cell.text = ''
                            if key == "TIME FRAME":
                                # Add formatted key and value
                                run = cell.paragraphs[0].add_run(key)
                                run.bold = True
                                cell.paragraphs[0].add_run(f"\n\n{formatted_time_frame.replace('-','/').replace('_', ' - ')}")
                                used_keys.append(key)
                            else:
                                # Add formatted key and value
                                run = cell.paragraphs[0].add_run(key)
                                run.bold = True
                                cell.paragraphs[0].add_run(f"\n\n{value}")
                                used_keys.append(key)

                    for key, value in self.footers.items():
                        if key in cell.text and key not in used_keys:
                            # Clear cell text
                            cell.text = ''
                            # Add formatted key and value
                            run = cell.paragraphs[0].add_run(key)
                            run.bold = True
                            cell.paragraphs[0].add_run(f"\n\n{value}")
                            used_keys.append(key)

                    for key, value in self.standards.items():
                        if key in cell.text and key not in used_keys:
                            cell.text = cell.text.replace(key, f"{value}")
                            used_keys.append(key)

        lesson_title = self.headers.get("[Lesson Title]", "No_Lesson_Title").replace(' ', '_')
        unit_title = self.headers.get("[Unit Title]", "No_Unit_Title").replace(' ', '_')
        time_frame = self.headers.get("TIME FRAME", "")

        # Date for the filename
        today_date = datetime.now().strftime("%m-%d-%y")
        today_date = '-'.join(str(int(part)) for part in today_date.split('-'))

        # Construct filename, only adding underscores if necessary components exist
        filename_components = ["lesson_plan", lesson_title, formatted_time_frame, today_date]
        # Remove empty strings and join with underscore
        filename = "_".join(filter(None, filename_components)) + ".docx"

        # Let user select the directory for saving the document
        # output_path = join(self.file_stem, "student_reporting_system/lesson_plans", filename)

        output_path = LP_DIR_PATH + "/" + filename

        # Save the document
        self.doc.save(output_path)
        messagebox.showinfo("Success", f"Document saved as {output_path}")


class RegistrationFormWindow(ttk.Frame):
    def __init__(self, parent, default_font):
        super().__init__(parent)
        self.default_font = default_font
        self.grid(padx=20, pady=20, sticky="nsew")  # Ensure the frame appears
        self.checked_items = set() 
        self.sort_directions_left = {
            "Select": False, "Last": True, "First": True
        }
        self.sort_directions_right = {
            "Last": True, "First": True, "Date": False, "Site": False, "Type": False,
            "Level": False, "Reading": False,
            "Writing": False, "Combined": False, "LACES": False, "StudentID": False, "DOB": False,
            "Gender": False, "Teacher": False, "Residency": False
        }
        self.create_widgets()

    def create_widgets(self):

        def update_test_dates(*args):
            """ Update test dates based on selected site. """
            site = self.registration_site.get()
            self.test_date_combobox['values'] = get_test_dates_by_site(site)
            self.test_date.set("")  # Clear previous selection

        # Variables
        self.registration_site = StringVar()
        self.registration_site.trace_add("write", update_test_dates)  # Trigger update on site change
        self.test_date = StringVar()
        # self.manual_test = IntVar()
        self.online_test = IntVar(value=1)

        # Configure grid layout to manage column proportions
        self.grid_columnconfigure(0, weight=1)  # Left frame (limited width)
        self.grid_columnconfigure(1, weight=1)  # Middle frame
        self.grid_columnconfigure(2, weight=1)  # Right frame   
        
        self.left_frame = ttk.Frame(self)
        self.left_frame.grid(row=0, column=0, padx=10, pady=10, sticky='w')

        self.middle_frame = ttk.Frame(self)
        self.middle_frame.grid(row=0, column=1, padx=10, pady=10, sticky='w')

        self.right_frame = ttk.Frame(self)
        self.right_frame.grid(row=0, column=2, padx=10, pady=10, sticky='w')

        # Bind resizing event to limit left frame width
        self.bind("<Configure>", self.limit_left_frame_size)

        self.left_frame_row_counter = 0
        self.left_frame_column_counter = 0

        self.site_label = ttk.Label(self.left_frame, text="Test Site: ", font=self.default_font)
        self.site_label.grid(row=self.left_frame_row_counter, column=0, padx=20, pady=10, sticky='w')

        self.site_combobox = AutocompleteCombobox(self.left_frame, textvariable=self.registration_site, values=get_site_names(), font=self.default_font)
        self.site_combobox.grid(row=self.left_frame_row_counter, column=1, padx=20, pady=10, sticky='w')
        self.left_frame_row_counter += 1

        self.test_date_label = ttk.Label(self.left_frame, text="Test Date: ", font=self.default_font)
        self.test_date_label.grid(row=self.left_frame_row_counter, column=0, padx=20, pady=10, sticky='w')

        self.test_date_combobox = AutocompleteCombobox(self.left_frame, textvariable=self.test_date, values=get_test_dates_by_site(self.registration_site.get()), font=self.default_font)
        self.test_date_combobox.grid(row=self.left_frame_row_counter, column=1, padx=20, pady=10, sticky='w')
        self.test_date_combobox.bind("<<ComboboxSelected>>", self.load_students_left)
        self.left_frame_row_counter += 1

        # Load checkbox images
        self.unchecked_img = ImageTk.PhotoImage(Image.new('RGBA', (15, 15), (255, 255, 255, 0)))
        self.checked_img = ImageTk.PhotoImage(Image.new('RGBA', (15, 15), (0, 0, 255, 255)))

        # Student table with checkboxes
        self.left_tree_columns = ("Select", "Last", "First", "Test ID")
        self.left_tree = ttk.Treeview(self.left_frame, columns=self.left_tree_columns, show='headings')
        for col in self.left_tree_columns:
            if col == 'Test ID':
                self.left_tree.column(col, width=0, stretch=False)
            elif col == 'Select':
                self.left_tree.column(col, width=75)
                self.left_tree.heading(col, text=col, command=lambda: self.sort_students_left(col))
            else:
                self.left_tree.heading(col, text=col, command=lambda: self.sort_students_left(col))
        self.left_tree.grid(row=self.left_frame_row_counter, column=0, columnspan=len(self.left_tree_columns), padx=10, pady=5, sticky='nsew')

        # Scrollbars
        vsb_left = ttk.Scrollbar(self.left_frame, orient="vertical", command=self.left_tree.yview) 
        vsb_left.grid(row=self.left_frame_row_counter, column=len(self.left_tree_columns)+1, sticky='ns')
        self.left_tree.configure(yscrollcommand=vsb_left.set)

        self.left_frame_row_counter += 1

        hsb_left = ttk.Scrollbar(self.left_frame, orient="horizontal", command=self.left_tree.xview) 
        hsb_left.grid(row=self.left_frame_row_counter, column=0, columnspan=len(self.left_tree_columns)+1, sticky='ew')
        self.left_tree.configure(xscrollcommand=hsb_left.set)

        self.test_modality_manual = Radiobutton(self.left_frame, text="Manual", variable=self.online_test, value=0)
        self.test_modality_manual.grid(row=self.left_frame_row_counter-0, column=len(self.left_tree_columns), padx=10, pady=5, sticky='nsew')

        self.test_modality_online = Radiobutton(self.left_frame, text="Online", variable=self.online_test, value=1)
        self.test_modality_online.grid(row=self.left_frame_row_counter-1, column=len(self.left_tree_columns), padx=10, pady=5, sticky='nsew')

        self.check_all_button = Button(self.left_frame, text="Check All", command=self.check_all_left)
        self.check_all_button.grid(row=self.left_frame_row_counter-2, column=len(self.left_tree_columns)-1, padx=10, pady=5, sticky='nsew')

        # Bind click event to toggle checkboxes
        self.left_tree.bind('<Button-1>', self.enable_toggle_checkbox)

        # Add students from left tree to right tree

        self.add_students_button = Button(self.middle_frame, text="Add Students =>", command=self.load_students_right)
        self.add_students_button.grid(row=0, column=0, padx=10, pady=5, sticky='nsew')

        # Right tree

        self.right_tree_columns = ("Last", "First", "Date", "Site", "Type", "Level", "Reading", "Writing", "Combined", "LACES", "StudentID", "DOB", "Gender", "Teacher", "Residency")

        self.staff_entry_var = tk.StringVar()

        self.right_frame_row_counter = 0
        self.right_frame_column_counter = 0

        self.empty_label = ttk.Label(self.right_frame, text="", font=self.default_font)
        self.empty_label.grid(row=self.right_frame_row_counter, column=0, padx=20, pady=10, sticky='w')

        self.right_frame_row_counter += 1

        self.staff_label = ttk.Label(self.right_frame, text="Staff: ", font=self.default_font)
        self.staff_label.grid(row=self.right_frame_row_counter, column=0, padx=20, pady=10, sticky='w')

        self.staff_entry = ttk.Entry(self.right_frame, textvariable=self.staff_entry_var, font=self.default_font)
        self.staff_entry.grid(row=self.right_frame_row_counter, column=1, padx=20, pady=10, sticky='w')

        self.clear_button = Button(self.right_frame, text='Clear', command=self.clear_right_tree, font=self.default_font)
        self.clear_button.grid(row=self.right_frame_row_counter, column=len(self.right_tree_columns)-2, padx=20, pady=10, sticky='w')

        self.form_button = Button(self.right_frame, text='Form', command=self.create_form, font=self.default_font)
        self.form_button.grid(row=self.right_frame_row_counter, column=len(self.right_tree_columns)-1, padx=20, pady=10, sticky='w')

        self.right_frame_row_counter += 1

        self.right_tree = ttk.Treeview(self.right_frame, columns=self.right_tree_columns, show='headings')
        for col in self.right_tree_columns:
            if col in ('LACES', 'StudentID', 'DOB', "Gender", "Teacher", "Residency"):
                self.right_tree.column(col, width=0, stretch=False)
            else:
                self.right_tree.heading(col, text=col, command=lambda: self.sort_students_right(col))
                self.right_tree.column(col, width=100, anchor='center')
        self.right_tree.grid(row=self.right_frame_row_counter, column=0, columnspan=len(self.right_tree_columns), padx=10, pady=5, sticky='nsew')
        # self.right_tree.configure(xscrollcommand=self.hsb_right.set)

        # Scrollbars
        self.vsb_right = ttk.Scrollbar(self.right_frame, orient="vertical", command=self.right_tree.yview) 
        self.vsb_right.grid(row=self.right_frame_row_counter, column=len(self.right_tree_columns)+1, sticky='ns')
        self.right_tree.configure(yscrollcommand=self.vsb_right.set)

        self.right_frame_row_counter += 1

        self.hsb_right = ttk.Scrollbar(self.right_frame, orient="horizontal", command=self.right_tree.xview) 
        self.hsb_right.grid(row=self.right_frame_row_counter, column=0, columnspan=len(self.right_tree_columns)+1, sticky='ew')
        self.right_tree.configure(xscrollcommand=self.hsb_right.set)

        self.master.state('zoomed')

    def check_all_left(self):
        for row in self.left_tree.get_children():
            # values = self.left_tree.item(row, "values")

            if row in self.checked_items:
                self.left_tree.item(row, image=self.unchecked_img, values=("",) + self.left_tree.item(row, 'values')[1:])
                self.checked_items.discard(row)
                # c = self.total_attendance_records_counter.get()
                # if c != 0:
                #     c -= 1
                #     self.total_attendance_records_counter.set(c)
            else:
                self.left_tree.item(row, image=self.checked_img, values=("✓",) + self.left_tree.item(row, 'values')[1:])
                self.checked_items.add(row)
                # c = self.total_attendance_records_counter.get()
                # c += 1
                # self.total_attendance_records_counter.set(c)

    def create_form(self):
        staff = self.staff_entry_var.get()
        site = self.registration_site.get()
        student_data = []

        # Get column headings from Treeview
        raw_column_names = self.right_tree["columns"]
        column_names = [re.sub(r'\W|^(?=\d)', '_', col) for col in raw_column_names]

        # Define a NamedTuple class dynamically
        Student = namedtuple("Student", column_names)

        for row in self.right_tree.get_children():
            values = self.right_tree.item(row, "values")
            # Convert the tuple into a NamedTuple instance
            student_data.append(Student(*values))

        student_data = sorted(student_data, key=lambda x: x.Last)

        days = sorted(list({s.Date for s in student_data}))
        test_types = set([s.Type for s in student_data])

        for s in student_data:
            if s.Date not in days:
                days.append(s.Date)

        if len(days) == 1:
            day_str = datetime.strptime(days[0], "%m/%d/%y").strftime("%#m-%#d-%y")

        else:
            formatted_days = [datetime.strptime(d, "%m/%d/%y").strftime("%#m-%#d-%y") for d in days]
            day_str = ' - '.join(formatted_days)

        pages_needed = math.ceil(len(self.right_tree.get_children()) / 10)

        if self.online_test.get():
        
            pdf_template_path = PROJECT_ROOT / "forms" / "ESOL Orientation Sign Up Sheet.pdf"
            
            doc = fitz.open(pdf_template_path)
            
            # If more pages are needed, insert from a fresh copy of the template
            while len(doc) < pages_needed:
                template = fitz.open(pdf_template_path)  # Open a new instance of the template
                doc.insert_pdf(template, from_page=0, to_page=0)  # Append a fresh page
                template.close()  # Close the extra instance
                
            for student_idx, student in enumerate(student_data):
                page_num = student_idx // 10  # Determine which page this student belongs to
                page = doc[page_num]

                row_position = student_idx % 10  # Determine which row on the page
                text_height = 12
                font = fitz.Font('helv') 

                site_heading_x1 = 64
                site_heading_x2 = 169

                instructor_heading_x1 = 345
                instructor_heading_x2 = 450

                date_heading_x1 = 495
                date_heading_x2 = 735
                
                heading_y1 =  107
                heading_y2 = 115

                y1 = 194 + (row_position * 37)
                y2 = 231 + (row_position * 37)

                site_column_x1 = 36
                site_column_x2 = 116

                name_column_x1 = 116
                name_column_x2 = 338

                day1_column_x1 = 338
                day1_column_x2 = 374

                day2_column_x1 = 374
                day2_column_x2 = 410

                reading_form_level_column_x1 = 410
                reading_form_level_column_x2 = 474

                reading_scaled_column_x1 = 474
                reading_scaled_column_x2 = 539

                writing_form_level_column_x1 = 539
                writing_form_level_column_x2 = 604

                writing_scaled_column_x1 = 604
                writing_scaled_column_x2 = 669

                combined_column_x1 = 669
                combined_column_x2 = 739

                residency_column_x1 = 742
                residency_column_x2 = 789

                if student_idx % 10 == 0:
                    
                    if len(test_types) > 1:
                        type_str = 'Pre- and Post-Tests'
                    else:
                        type_str = student.Type + "-Tests"

                    type_str = f'({type_str})'

                    ## Site Heading ##

                    site_heading_font_size = 16

                    while ((site_heading_text_length := font.text_length(site, fontsize=site_heading_font_size)) + site_heading_x1) > site_heading_x2 - 5:
                        site_heading_font_size -= 1

                    site_heading_x_offset = (((site_heading_x2 - site_heading_x1) - site_heading_text_length) / 2) + site_heading_x1
                    site_heading_y_offset = (((heading_y2 - heading_y1) - text_height) / 2) + heading_y1 + text_height

                    page.insert_text((site_heading_x_offset, site_heading_y_offset), site, fontsize=site_heading_font_size)
                    
                    ## Instructor Heading ##

                    instructor_heading_font_size = 16

                    while ((instructor_heading_text_length := font.text_length(staff, fontsize=instructor_heading_font_size)) + instructor_heading_x1) > instructor_heading_x2 - 5:
                        instructor_heading_font_size -= 1

                    instructor_heading_x_offset = (((instructor_heading_x2 - instructor_heading_x1) - instructor_heading_text_length) / 2) + instructor_heading_x1
                    instructor_heading_y_offset = (((heading_y2 - heading_y1) - text_height) / 2) + heading_y1 + text_height

                    page.insert_text((instructor_heading_x_offset, instructor_heading_y_offset), staff, fontsize=instructor_heading_font_size)

                    ## Day Heading ##

                    date_heading_font_size = 16

                    while ((date_heading_text_length := font.text_length(day_str, fontsize=date_heading_font_size)) + date_heading_x1) > date_heading_x2 - 5:
                        date_heading_font_size -= 1

                    date_heading_x_offset = (((date_heading_x2 - date_heading_x1) - date_heading_text_length) / 2) + date_heading_x1
                    date_heading_y_offset = (((heading_y2 - heading_y1) - text_height) / 2) + heading_y1 + text_height
                    day_str = day_str.replace('/','-')
                    if ' - ' in day_str:
                        day_str_split = day_str.split(' - ')
                        day_1 = datetime.strptime(day_str_split[0], "%m-%d-%y").strftime("%#m/%#d/%y")
                        day_2 = datetime.strptime(day_str_split[1], "%m-%d-%y").strftime("%#m/%#d/%y")
                        day_str = f"{day_1} - {day_2}"
                    else:
                        day_str = datetime.strptime(day_str, "%m-%d-%y").strftime("%#m/%#d/%y")
                    page.insert_text((date_heading_x_offset, date_heading_y_offset), day_str, fontsize=date_heading_font_size)

                    ## Type Heading ##

                    type_heading_font_size = 16

                    while ((type_heading_text_length := font.text_length(type_str, fontsize=type_heading_font_size)) + date_heading_x1) > date_heading_x2 - 5:
                        type_heading_text_length -= 1

                    date_heading_x_offset = (((date_heading_x2 - date_heading_x1) - type_heading_text_length) / 2) + date_heading_x1
                    date_heading_y_offset = (((53 - 40) - text_height) / 2) + 40 + text_height

                    page.insert_text((date_heading_x_offset, date_heading_y_offset), type_str, fontsize=type_heading_font_size)

                ## Site ##

                # If newline in student.Site, student has multiple sites; handle them successively
                if '\n' in student.Site:
                    sites_list = student.Site.split('\n')
                    for site_idx, site_name in enumerate(sites_list):
                        site_idx += 1
                        site_font_size = 16

                        while ((site_text_length := font.text_length(site_name, fontsize=site_font_size)) + site_column_x1) > site_column_x2 - 5:
                            site_font_size -= 1

                        site_x_offset = (((site_column_x2 - site_column_x1) - site_text_length) / 2) + site_column_x1
                        total_text_height = len(sites_list) * text_height
                        site_y_offset = (((y2 - y1) - total_text_height) / 2) + y1 + (site_idx*text_height)

                        page.insert_text((site_x_offset, site_y_offset), site_name, fontsize=site_font_size)
                else:
                    site_font_size = 16

                    while ((site_text_length := font.text_length(student.Site, fontsize=site_font_size)) + site_column_x1) > site_column_x2 - 5:
                        site_font_size -= 1

                    site_x_offset = (((site_column_x2 - site_column_x1) - site_text_length) / 2) + site_column_x1
                    site_y_offset = (((y2 - y1) - text_height) / 2) + y1 + text_height

                    page.insert_text((site_x_offset, site_y_offset), student.Site, fontsize=site_font_size)

                ## Name ## 
                name_font_size = 16

                student_name_str = f'{student.Last}, {student.First}'

                # If there are both pre- and post-tests on the form, append each test type to the student's name
                if len(test_types) > 1:
                    student_name_str += f' ({student.Type.lower()})'

                while ((name_text_length := font.text_length(student_name_str, fontsize=name_font_size)) + name_column_x1) > name_column_x2 - 5:
                    name_font_size -= 1

                name_x_offset = (((name_column_x2 - name_column_x1) - name_text_length) / 2) + name_column_x1
                name_y_offset = (((y2 - y1) - text_height) / 2) + y1 + text_height

                page.insert_text((name_x_offset, name_y_offset), student_name_str, fontsize=name_font_size)

                ## Day 1 ##

                # Determine which day the student's test was (day 1 or 2)
                if len(days) == 1:
                    day_1 = 'X'
                    day_2 = ''
                elif len(days) == 2:
                    if student.Date == days[0]:
                        day_1 = 'X'
                        day_2 = ''
                    else:
                        day_1 = ''
                        day_2 = 'X'
                
                day1_font_size = 16

                while ((day1_text_length := font.text_length(day_1, fontsize=day1_font_size)) + day1_column_x1) > day1_column_x2 - 5:
                    day1_font_size -= 1

                day1_x_offset = (((day1_column_x2 - day1_column_x1) - day1_text_length) / 2) + day1_column_x1
                day1_y_offset = (((y2 - y1) - text_height) / 2) + y1 + text_height

                page.insert_text((day1_x_offset, day1_y_offset), day_1, fontsize=day1_font_size)

                ## Day 2 ##

                day2_font_size = 16

                while ((day2_text_length := font.text_length(day_2, fontsize=day2_font_size)) + day2_column_x1) > day2_column_x2 - 5:
                    day2_font_size -= 1

                day2_x_offset = (((day2_column_x2 - day2_column_x1) - day2_text_length) / 2) + day2_column_x1
                day2_y_offset = (((y2 - y1) - text_height) / 2) + y1 + text_height

                page.insert_text((day2_x_offset, day2_y_offset), day_2, fontsize=day2_font_size)

                ## Reading Form Level ##

                reading_form_level_font_size = 16

                while ((reading_form_level_text_length := font.text_length(student.Level, fontsize=reading_form_level_font_size)) + reading_form_level_column_x1) > reading_form_level_column_x2 - 5:
                    reading_form_level_font_size -= 1

                reading_form_level_x_offset = (((reading_form_level_column_x2 - reading_form_level_column_x1) - reading_form_level_text_length) / 2) + reading_form_level_column_x1
                reading_form_level_y_offset = (((y2 - y1) - text_height) / 2) + y1 + text_height

                page.insert_text((reading_form_level_x_offset, reading_form_level_y_offset), student.Level, fontsize=reading_form_level_font_size)

                ## Reading Scaled ##

                reading_scaled_font_size = 16

                while ((reading_scaled_text_length := font.text_length(student.Reading, fontsize=reading_scaled_font_size)) + reading_scaled_column_x1) > reading_scaled_column_x2 - 5:
                    reading_scaled_font_size -= 1

                reading_scaled_x_offset = (((reading_scaled_column_x2 - reading_scaled_column_x1) - reading_scaled_text_length) / 2) + reading_scaled_column_x1
                reading_scaled_y_offset = (((y2 - y1) - text_height) / 2) + y1 + text_height

                if self.online_test.get():
                    page.insert_text((reading_scaled_x_offset, reading_scaled_y_offset), student.Reading, fontsize=reading_scaled_font_size)

                ## Writing Form Level ##

                writing_form_level_font_size = 16

                while ((writing_form_level_text_length := font.text_length(student.Level, fontsize=writing_form_level_font_size)) + writing_form_level_column_x1) > writing_form_level_column_x2 - 5:
                    writing_form_level_font_size -= 1

                writing_form_level_x_offset = (((writing_form_level_column_x2 - writing_form_level_column_x1) - writing_form_level_text_length) / 2) + writing_form_level_column_x1
                writing_form_level_y_offset = (((y2 - y1) - text_height) / 2) + y1 + text_height

                page.insert_text((writing_form_level_x_offset, writing_form_level_y_offset), student.Level, fontsize=writing_form_level_font_size)

                ## Writing Scaled ##

                writing_scaled_font_size = 16

                while ((writing_scaled_text_length := font.text_length(student.Writing, fontsize=writing_scaled_font_size)) + writing_scaled_column_x1) > writing_scaled_column_x2 - 5:
                    writing_scaled_font_size -= 1

                writing_scaled_x_offset = (((writing_scaled_column_x2 - writing_scaled_column_x1) - writing_scaled_text_length) / 2) + writing_scaled_column_x1
                writing_scaled_y_offset = (((y2 - y1) - text_height) / 2) + y1 + text_height
                
                if self.online_test.get():
                    page.insert_text((writing_scaled_x_offset, writing_scaled_y_offset), student.Writing, fontsize=writing_scaled_font_size)

                ## Combined ##

                combined_font_size = 16

                while ((combined_text_length := font.text_length(student.Combined, fontsize=combined_font_size)) + combined_column_x1) > combined_column_x2 - 5:
                    combined_font_size -= 1

                combined_x_offset = (((combined_column_x2 - combined_column_x1) - combined_text_length) / 2) + combined_column_x1
                combined_y_offset = (((y2 - y1) - text_height) / 2) + y1 + text_height
                
                if self.online_test.get():
                    page.insert_text((combined_x_offset, combined_y_offset), student.Combined, fontsize=combined_font_size)
                
                ## Residency ##

                residency_font_size = 16

                while ((residency_text_length := font.text_length(student.Residency, fontsize=residency_font_size)) + residency_column_x1) > residency_column_x2 - 5:
                    residency_font_size -= 1

                residency_x_offset = (((residency_column_x2 - residency_column_x1) - residency_text_length) / 2) + residency_column_x1
                residency_y_offset = (((y2 - y1) - text_height) / 2) + y1 + text_height
                
                if self.online_test.get():
                    # In the future, we'll need to differentiate between new students who need Residency info on form
                    # and existing students who don't need Residency...
                    # for now, just check that they have Residency info, which, at the moment (10-31-25), only new students have
                    if "None" not in student.Residency.strip():
                        page.insert_text((residency_x_offset, residency_y_offset), student.Residency, fontsize=residency_font_size)
        
        else:

            # manually scored test
            pdf_template_path = PROJECT_ROOT / "forms" / "ESOL Orientation Sign-In Sheet_Updated (2025).pdf"
            reader = PdfReader(pdf_template_path)
            writer = PdfWriter()                        

            # clone the whole document into writer
            writer.clone_document_from_reader(reader)

            # grab the first page of the template
            template_page = reader.pages[0]

            # append fresh copies of page 0 until we reach pages_needed
            while len(writer.pages) < pages_needed:
                writer.add_page(template_page)

            for student_idx, student in enumerate(student_data):
                page_num = student_idx // 10  # Determine which page this student belongs to
                page = writer.pages[page_num]

                data_dict = {}

                if student_idx % 10 == 0:
                    
                    if len(test_types) > 1:
                        type_str = 'Pre- and Post-Tests'
                    else:
                        type_str = student.Type + "-Tests"

                    type_str = f'({type_str})'
                    day_str = datetime.strptime(day_str, "%m-%d-%y").strftime("%#m/%#d/%y")

                    data_dict.update({
                        'Site': site,
                        'Orientation Counselor Instructor': staff,
                        'Dates': day_str,
                        'Type': type_str,
                        "Date of Birth": datetime.strptime(format_date(student.DOB,out=True), "%Y-%m-%d").strftime("%#m/%#d/%Y")
                    })

                # if '\n' in student.Site:
                #     sites_list = student.Site.split('\n')
                #     # for site_idx, site_name in enumerate(sites_list):
                #     #     site_idx += 1

                #     #     site_field = 
                site_field = f'ESOL Class Placement SiteRow{student_idx+1}'
                data_dict[site_field] = student.Site

                student_name_str = f'{student.Last}, {student.First}'

                # If there are both pre- and post-tests on the form, append each test type to the student's name
                if len(test_types) > 1:
                    student_name_str += f' ({student.Type.lower()})'

                student_name_field = f'Print Name Last FirstRow{student_idx+1}'
                data_dict[student_name_field] = student_name_str


                # Determine which day the student's test was (day 1 or 2)
                if len(days) == 1:
                    day_1 = 'X'
                    day_2 = ''
                elif len(days) == 2:
                    if student.Date == days[0]:
                        day_1 = 'X'
                        day_2 = ''
                    else:
                        day_1 = ''
                        day_2 = 'X'

                day_1_field = f'Day 1Row{student_idx+1}'
                data_dict[day_1_field] = day_1

                day_2_field = f'Day 2Row{student_idx+1}'
                data_dict[day_2_field] = day_2

                reading_form_field = f'Reading Form LevelRow{student_idx+1}'
                data_dict[reading_form_field] = student.Level

                writing_form_field = f'Writing Form LevelRow{student_idx+1}'
                data_dict[writing_form_field] = student.Level

                if "None" not in student.Residency.strip(): 
                    residency_form_field = f'Residency{student_idx+1}'
                    data_dict[residency_form_field] = student.Residency

                # Tell the PDF viewer to regenerate appearances
                acro_form = writer._root_object["/AcroForm"]
                acro_form.update({
                    NameObject("/NeedAppearances"): BooleanObject(True),
                })

                # For every field, set
                #   /Q = 1   ← center justification
                #   /DA = "/Helv 0 Tf 0 g"  ← Helvetica, size=0 → auto‐scale
                for field_ref in acro_form["/Fields"]:
                    field = field_ref.get_object()
                    # if field['/T'] in ("Site", "Orientation Counselor Instructor", "Dates", "Type"):
                    #     field.update({
                    #         NameObject("/Q"):  NumberObject(0),
                    #         NameObject("/DA"): TextStringObject("/Helv 0 Tf 0 g"),
                    #     })
                    # else:
                    field.update({
                        NameObject("/Q"):  NumberObject(1),
                        NameObject("/DA"): TextStringObject("/Helv 0 Tf 0 g"),
                    })

                writer.update_page_form_field_values(page, data_dict)
        if ' - ' in day_str:
            day_str_split = day_str.split(' - ')
            day_1 = datetime.strptime(day_str_split[0], "%m/%d/%y").strftime("%#m-%#d-%y")
            day_2 = datetime.strptime(day_str_split[1], "%m/%d/%y").strftime("%#m-%#d-%y")
            day_str = f"{day_1}-{day_2}"
        else:
            day_str = datetime.strptime(day_str, "%m/%d/%y").strftime("%#m-%#d-%y")
        # Save PDF
        save_path = filedialog.asksaveasfilename(
            defaultextension=".pdf", 
            filetypes=[("PDF files", "*.pdf")], 
            initialdir=PROJECT_ROOT / "output" / "Orientation Sign Up Sheets",
            initialfile=f'''Orientation Sign-In Sheet {day_str} @ {site} ({datetime.now().strftime("%Y-%m-%d_%H-%M-%S")})'''
        )
        
        if save_path:
            if self.online_test.get():
                doc.save(save_path)
            else:
                writer.write(save_path)
            # messagebox.showinfo("Success", "PDF has been created successfully!")

            # Open the PDF after saving
            if os.name == "nt":  # Windows
                os.startfile(save_path)
            elif os.name == "posix":  # macOS/Linux
                os.system(f"open {save_path}" if sys.platform == "darwin" else f"xdg-open {save_path}")

        # Create pdfs of student progress forms
        pages_needed = len(student_data)

        progress_form_pdf_fn = PROJECT_ROOT / "forms" / "ESOL Student Progress Form (Fillable).pdf"

        gender_field_map = {
            "male": "Male",
            "female": "Female",
            "non-binary": "Non-Binary",
            "nonbinary": "Non-Binary",
            "prefer not to answer": "Prefer Not to Answer",
        }

        tmp_pdf_files = []

        with tempfile.TemporaryDirectory() as tmp_dir:

            # Make one separate temporary PDF per student
            for i, student in enumerate(student_data, start=1):

                reader = PdfReader(progress_form_pdf_fn)
                writer = PdfWriter()

                # clone the whole document into writer
                writer.clone_document_from_reader(reader)

                # Since this temp PDF is for ONE student only,
                # do NOT append extra pages.
                # Just use the first page of this one-student writer.
                page = writer.pages[0]

                data_dict = {
                    'Site': student.Site.replace('\n', '; '),
                    'Aspire staff': student.Teacher,
                    'Last Name': student.Last,
                    'First Name': student.First,
                    "Date of Birth": datetime.strptime(student.DOB, "%Y-%m-%d").strftime("%#m/%#d/%Y"),
                    "LACES ID": student.LACES
                }

                if student.Type.lower() == "pre":
                    data_dict['Date_Initial'] = datetime.strptime(student.Date, "%m/%d/%y").strftime("%#m/%#d/%y")
                    data_dict['LevelFormReading'] = student.Level
                    data_dict['LevelFormWriting'] = student.Level
                    if student.Reading:
                        data_dict['Scale ScoreReading'] = student.Reading
                    if student.Writing:
                        data_dict['Scale ScoreWriting'] = student.Writing
                    if student.Reading and student.Writing:
                        data_dict['Total Reading/Writing'] = student.Combined

                else:
                    data_dict['Date_Post1'] = datetime.strptime(student.Date, "%m/%d/%y").strftime("%#m/%#d/%y")
                    data_dict['LevelFormReading_2'] = student.Level
                    data_dict['LevelFormWriting_2'] = student.Level
                    if student.Reading:
                        data_dict['Scale ScoreReading_2'] = student.Reading
                    if student.Writing:
                        data_dict['Scale ScoreWriting_2'] = student.Writing
                    if student.Reading and student.Writing:
                        data_dict['Total Reading/Writing_2'] = student.Combined

                # Tell the PDF viewer to regenerate appearances
                acro_form = writer._root_object["/AcroForm"]
                acro_form.update({
                    NameObject("/NeedAppearances"): BooleanObject(True),
                })

                # For every field, set
                #   /Q = 1   ← center justification
                #   /DA = "/Helv 0 Tf 0 g"  ← Helvetica, size=0 → auto‐scale
                for field_ref in acro_form["/Fields"]:
                    field = field_ref.get_object()
                    if field['/T'] in ("Last Name", "First Name"):
                        field.update({
                            NameObject("/Q"):  NumberObject(0),
                            NameObject("/DA"): TextStringObject("/Helv 0 Tf 0 g"),
                        })
                    else:
                        field.update({
                            NameObject("/Q"):  NumberObject(1),
                            NameObject("/DA"): TextStringObject("/Helv 0 Tf 0 g"),
                        })

                # Now fill in the values; they’ll pick up DA and Q settings
                writer.update_page_form_field_values(page, data_dict)

                # Write this ONE student's filled PDF to memory
                buf = io.BytesIO()
                writer.write(buf)
                buf.seek(0)

                # Open this ONE student's filled PDF with fitz
                doc_tmp = fitz.open(stream=buf.read(), filetype="pdf")

                # primary class and gender will be rectangle
                for page, student_for_rect in zip(doc_tmp, [student]):
                    padding = 2

                    # Fields to draw red rectangles around on this page
                    wanted_field_names = ["ESOL"]

                    gender_key = str(student_for_rect.Gender).strip().lower()
                    gender_field_name = gender_field_map.get(gender_key)

                    if gender_field_name:
                        wanted_field_names.append(gender_field_name)
                    else:
                        print(f"Unknown gender value: {student_for_rect.Gender!r}")

                    found_fields = set()

                    for widget in page.widgets() or []:
                        if widget.field_name in wanted_field_names:
                            padded_rect = widget.rect + (-padding, -padding, padding, padding)

                            page.draw_rect(
                                padded_rect,
                                color=(1, 0, 0),
                                width=1
                            )

                            found_fields.add(widget.field_name)

                    missing = set(wanted_field_names) - found_fields
                    if missing:
                        print(f"Could not find these AcroForm fields on page: {missing}")

                # Save this ONE student's temp PDF
                tmp_pdf_path = os.path.join(tmp_dir, f"progress_form_tmp_{i}.pdf")
                doc_tmp.save(tmp_pdf_path)
                doc_tmp.close()

                tmp_pdf_files.append(tmp_pdf_path)

            # Combine the completed one-student temp PDFs at the last moment
            doc2 = fitz.open()

            for tmp_pdf_path in tmp_pdf_files:
                tmp_doc = fitz.open(tmp_pdf_path)
                doc2.insert_pdf(tmp_doc)
                tmp_doc.close()

            # Save final PDF
            save_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf")],
                initialdir=PROJECT_ROOT / "output" / "Progress Forms",
                initialfile=f'''Progress Forms {day_str} @ {site} ({datetime.now().strftime("%Y-%m-%d_%H-%M-%S")})'''
            )

            if save_path:
                pdf_title = os.path.splitext(os.path.basename(save_path))[0]

                # Clear existing metadata first
                doc2.set_metadata({})

                # Set fresh metadata
                doc2.set_metadata({
                    "title": pdf_title,
                    "author": "",
                    "subject": "",
                    "keywords": "",
                    "creator": "",
                    "producer": "",
                })

                # Remove XMP metadata if present
                try:
                    doc2.del_xml_metadata()
                except Exception as e:
                    print(f"Could not delete XML metadata: {e}")

                doc2.save(save_path)

                doc2.close()

                if os.name == "nt":
                    os.startfile(save_path)
                elif os.name == "posix":
                    os.system(f"open {save_path}" if sys.platform == "darwin" else f"xdg-open {save_path}")
            else:
                doc2.close()
    def clear_right_tree(self):
        # Remove all rows from self.right_tree
        for row in self.right_tree.get_children():
            self.right_tree.delete(row)

    def limit_left_frame_size(self, event=None):
        """ Ensure left frame does not exceed half the window width. """
        max_width = self.winfo_width() // 2
        current_width = self.left_frame.winfo_width()
        
        # Set the minimum width if smaller, otherwise limit max width
        self.left_frame.config(width=min(max_width, current_width))
    
    def load_students_right(self):
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        checked_test_ids = []

        for item in self.checked_items:
            # Get the values of the item
            values = self.left_tree.item(item, "values")

            # Get the last column value
            if values:  # Ensure values exist
                checked_test_pk = values[-1]
                checked_test_ids.append(checked_test_pk)

        # Ensure self.test_ids_left is not empty
        if not checked_test_ids:
            return []  # Return empty list if there are no test IDs

        # Create placeholders for the IN clause dynamically
        placeholders = ", ".join(["?"] * len(checked_test_ids))

        # query = f'''
        #     SELECT
        #         s.last_name, 
        #         s.first_name,
        #         t.date,
        #         IIF(pre_or_post = 'Post', t.site, st.site_name),
        #         t.pre_or_post,
        #         t.level || t.form,
        #         t.reading_scaled,
        #         t.writing_scaled,
        #         t.combined,
        #         IFNULL(t.laces, '')
        #     FROM
        #         Testing t
        #     LEFT JOIN
        #         Student s
        #     ON
        #         t.student_id = s.student_id
        #     LEFT JOIN
        #         StudentSites ss
        #     ON
        #         t.student_id = ss.student_id
        #     LEFT JOIN
        #         Sites st
        #     ON
        #         ss.site_pk = st.site_pk
        #     WHERE
        #         t.testing_pk IN ({placeholders})
        #         AND t.form <> 'Locator'
        #     ORDER BY
        #         s.last_name ASCf
        # '''
        query = f"""
            SELECT
                s.last_name,
                s.first_name,
                t.date,
                CASE
                    WHEN t.pre_or_post IN ('Pre', 'Post') THEN COALESCE(
                        (
                            SELECT GROUP_CONCAT(x.site_name, CHAR(10))
                            FROM (
                                SELECT DISTINCT st2.site_name AS site_name
                                FROM StudentSites ss2
                                JOIN Sites st2 ON ss2.site_pk = st2.site_pk
                                WHERE ss2.student_id = s.student_id
                            ) AS x
                        ),
                        t.site
                    )
                    ELSE t.site
                END AS grouped_sites,
                t.pre_or_post,
                t.level || t.form AS test_level,
                t.reading_scaled,
                t.writing_scaled,
                t.combined,
                IFNULL(t.laces, '') AS laces,
                s.student_id,
                s.dob,
                s.gender,
                (
                    SELECT GROUP_CONCAT(y.teacher, '; ')
                    FROM (
                        SELECT DISTINCT st2.teacher AS teacher
                        FROM StudentSites ss2
                        JOIN Sites st2 ON ss2.site_pk = st2.site_pk
                        WHERE ss2.student_id = s.student_id
                    ) AS y
                ) AS teacher,
                s.residency_document
            FROM Testing t
            JOIN Student s ON t.student_id = s.student_id
            WHERE
                t.testing_pk IN ({placeholders})
                AND t.form <> 'Locator'
            ORDER BY s.last_name ASC;
        """



        cursor.execute(query, tuple(checked_test_ids))
        students = cursor.fetchall()
        conn.close()

        # Clear existing rows
        # for row in self.left_tree.get_children():
        #     self.left_tree.delete(row)

        # Populate students and remove leading zeros
        # for student in students:
        #     last_name, first_name, pre_or_post, reading_form_level, reading_scaled, writing_form_level, writing_scaled, combined = student
        #     item = self.left_tree.insert('', 'end', values=("", last_name, first_name, pre_or_post, reading_form_level, reading_scaled, writing_form_level, writing_scaled, combined), image=self.unchecked_img)
        #     self.checked_items.discard(item)

        for student in students:
            last_name, first_name, date, site, pre_or_post, form_level, reading_scaled, writing_scaled, combined, laces, student_id, dob, gender, teacher, residency_document = student
            self.right_tree.insert('', 'end', values=(last_name, first_name, format_date(date, display_format='%#m/%#d/%y'), site, pre_or_post, form_level, reading_scaled, writing_scaled, combined, laces, student_id, dob, gender, teacher, residency_document))
        
        # Uncheck students from left tree after adding them to right tree
        for row in self.left_tree.get_children():
            self.left_tree.item(row, image=self.unchecked_img, values=("",) + self.left_tree.item(row, 'values')[1:])
            self.checked_items.discard(row)


    def sort_students_left(self, col):
        # Toggle the sorting direction
        self.sort_directions_left[col] = not self.sort_directions_left[col]
        data = [(self.left_tree.set(child, col), child) for child in self.left_tree.get_children('')]
        if col == 'Last Class':
            # Handle sorting for date column
            data.sort(key=lambda x: datetime.strptime(x[0], '%Y-%m-%d') if x[0] else datetime.min, reverse=self.sort_directions_left[col])
        else:
            # Handle sorting for text columns
            data.sort(reverse=self.sort_directions_left[col])

        for index, (val, child) in enumerate(data):
            self.left_tree.move(child, '', index)

        # Update heading with new sort direction
        self.left_tree.heading(col, command=lambda: self.sort_students(col))

    def sort_students_right(self, col):
        # Toggle the sorting direction
        self.sort_directions_right[col] = not self.sort_directions_right[col]
        data = [(self.right_tree.set(child, col), child) for child in self.right_tree.get_children('')]
        if col == 'Last Class':
            # Handle sorting for date column
            data.sort(key=lambda x: datetime.strptime(x[0], '%Y-%m-%d') if x[0] else datetime.min, reverse=self.sort_directions_left[col])
        else:
            # Handle sorting for text columns
            data.sort(reverse=self.sort_directions_right[col])

        for index, (val, child) in enumerate(data):
            self.right_tree.move(child, '', index)

        # Update heading with new sort direction
        self.right_tree.heading(col, command=lambda: self.sort_students(col))

    def toggle_checkbox(self, event):
        item = self.left_tree.identify_row(event.y)
        if item:
            if item in self.checked_items:
                self.left_tree.item(item, image=self.unchecked_img, values=("",) + self.left_tree.item(item, 'values')[1:])
                self.checked_items.discard(item)
                # c = self.total_attendance_records_counter.get()
                # if c != 0:
                #     c -= 1
                #     self.total_attendance_records_counter.set(c)
            else:
                self.left_tree.item(item, image=self.checked_img, values=("✓",) + self.left_tree.item(item, 'values')[1:])
                self.checked_items.add(item)
                # c = self.total_attendance_records_counter.get()
                # c += 1
                # self.total_attendance_records_counter.set(c)


    def enable_toggle_checkbox(self, event):
        region = self.left_tree.identify('region', event.x, event.y)
        if region != 'heading':
            self.toggle_checkbox(event)

        
    def load_students_left(self, event=None):
        test_date = self.test_date.get()
        site = self.registration_site.get()
        if None in (site, test_date):
            return
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        query = '''
            SELECT
                s.last_name, 
                s.first_name,
                t.pre_or_post,
                t.form || t.level,
                t.reading_scaled,
                t.form || t.level,
                t.writing_scaled,
                t.combined,
                t.testing_pk
            FROM
                Testing t
            LEFT JOIN
                Student s
            ON
                t.student_id = s.student_id
            WHERE
                t.date = ?
                AND t.site = ?
                AND t.form <> 'Locator'
            ORDER BY
                s.last_name ASC
        '''

        cursor.execute(query, (test_date, site,))
        students = cursor.fetchall()
        conn.close()

        # Clear existing rows
        for row in self.left_tree.get_children():
            self.left_tree.delete(row)

        # Populate students and remove leading zeros
        # for student in students:
        #     last_name, first_name, pre_or_post, reading_form_level, reading_scaled, writing_form_level, writing_scaled, combined = student
        #     item = self.left_tree.insert('', 'end', values=("", last_name, first_name, pre_or_post, reading_form_level, reading_scaled, writing_form_level, writing_scaled, combined), image=self.unchecked_img)
        #     self.checked_items.discard(item)

        for student in students:
            last_name, first_name, pre_or_post, reading_form_level, reading_scaled, writing_form_level, writing_scaled, combined, testing_pk = student
            item = self.left_tree.insert('', 'end', values=("", last_name, first_name, testing_pk), image=self.unchecked_img)
            self.checked_items.discard(item)
            
class TestingTab(ttk.Frame):
    def __init__(self, parent, default_font):
        super().__init__(parent)
        self.default_font = default_font
        self.total_testing_records_counter = StringVar()
        self.test_progress_form_tab = None

        self.create_widgets()
        self.populate_treeview()       

    def create_widgets(self):

        treeview_columns = [
            "Test ID",
            "Student ID",
            "Last Name",
            "First Name",
            "LACES ID",
            "Site",
            "Staff",
            "Date",
            "Pre/Post",
            "Level",
            "Form",
            "Reading Raw",
            "Reading Scaled",
            "Reading NRS",
            "Writing Raw",
            "Writing Scaled",
            "Writing NRS",
            "Combined",
            "MSG",
            "MSG Notes",
            "Test Notes"
        ]
        
        # Create Treeview widget
        self.tree = ttk.Treeview(self, columns=tuple(treeview_columns), show='headings')
        for col in treeview_columns:
            self.tree.heading(col, text=str(col), command=lambda c=col: self.sort_by_column(c, False))      
            # Add some padding to make the heading width a little larger
            self.tree.column(col, anchor='center', stretch=NO)

        self.update()

        # Bind double click to treeview items (to view student's file)
        self.tree.bind("<Double-1>", self.on_treeview_click)

        # Pack Treeview widget
        self.tree.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky='nsew')

        # Add vertical and horizontal scrollbars
        self.vscrollbar = ttk.Scrollbar(self, orient=VERTICAL, command=self.tree.yview) 
        self.vscrollbar.grid(column=2, row=1, sticky=(N,S)) 
        self.tree['yscrollcommand'] = self.vscrollbar.set

        self.xscrollbar = ttk.Scrollbar(self, orient=HORIZONTAL, command=self.tree.xview) 
        self.xscrollbar.grid(column=0, row=2, columnspan=2, sticky=(E,W))
        self.tree['xscrollcommand'] = self.xscrollbar.set

        # Button frame
        button_frame = ttk.Frame(self)
        button_frame.grid(row=4, column=0, columnspan=4, padx=5, pady=(2, 0), sticky='ew')

        # Configure the button frame's grid columns
        button_frame.grid_columnconfigure(0, weight=1)  # Column for centering
        button_frame.grid_columnconfigure(1, weight=0)  # Column for right-justified buttons
        button_frame.grid_columnconfigure(2, weight=0)  # Column for right-justified buttons

        # Add Test button (centered)
        self.add_test_button = Button(button_frame, text="Add Test", command=self.add_test, font=self.default_font)
        self.add_test_button.grid(row=0, column=0, padx=2, pady=1, sticky='')

        # Filter button (right-justified)
        self.filter_button = Button(button_frame, text="Filter", command=self.open_filter_window, font=self.default_font)
        self.filter_button.grid(row=0, column=1, padx=2, pady=1, sticky='e')

        # Clear Filter button (right-justified)
        self.clear_filter_button = Button(button_frame, text="Clear Filter", command=self.reset_filters, font=self.default_font)
        self.clear_filter_button.grid(row=0, column=2, padx=2, pady=1, sticky='e')

        # Create Registration Form
        self.create_registration_form_button = Button(button_frame, text="Create Registration Form", command=self.create_registration_form, font=self.default_font)
        self.create_registration_form_button.grid(row=0, column=3, padx=2, pady=1, sticky='e')

        
        # Total Attendance Records label frame
        total_record_frame = ttk.Frame(self)
        total_record_frame.grid(row=3, column=0, columnspan=4, padx=5, pady=(2, 0), sticky='w')

        # Total Testing Records label
        self.total_testing_records_prefix = Label(total_record_frame, text="Total Testing Records:", font=self.default_font)
        self.total_testing_records_prefix.grid(row=0, column=0, padx=2, pady=1, sticky='w')

        # Total Testing Records label
        self.total_testing_records = Label(total_record_frame, textvariable=self.total_testing_records_counter, font=self.default_font)
        self.total_testing_records.grid(row=0, column=1, padx=2, pady=1, sticky='w')

        # Configure grid to expand properly
        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(0, weight=1)

    def populate_treeview(self, filtered_rows=None):
        # Clear existing rows
        for row in self.tree.get_children():
            self.tree.delete(row)

        # Connect to the SQLite database
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        if filtered_rows is None:
            # Fetch student data along with the most recent class_date
            query = '''
                SELECT 
                    t.testing_pk,
                    t.student_id,
                    s.last_name AS last_name,
                    s.first_name AS first_name,
                    t.laces,
                    t.site,
                    t.staff,
                    t.date,
                    t.pre_or_post,
                    t.level,
                    t.form,
                    t.reading_raw,
                    t.reading_scaled,
                    t.reading_nrs,
                    t.writing_raw,
                    t.writing_scaled,
                    t.writing_nrs,
                    t.combined,
                    t.got_msg,
                    t.msg_area,
                    t.test_notes
                FROM
                    Testing t
                JOIN
                    Student s
                ON
                    s.student_id = t.student_id
                ORDER BY
                    t.testing_pk DESC   
            '''
            cursor.execute(query)
            rows = cursor.fetchall()
        else:
            rows = filtered_rows

        # Update student counter intvar
        c = len(rows)
        self.total_testing_records_counter.set(c)

        # Insert data into the Treeview
        for row in rows:
            formatted_row = list(row)
            formatted_row = [f if f != None else "" for f in formatted_row]
            formatted_row[7] = format_date(formatted_row[7]) if formatted_row[5] else ''  # Format last_class date
            self.tree.insert('', 'end', values=formatted_row)

        # Adjust column widths based on content
        for col in self.tree['columns']:
            # Start with the header width
            max_width = tkFont.Font().measure(col)
            
            # Create list of possible cell widths - use the max to determine the column width 
            widths = []
            for item in self.tree.get_children():
                cell_value = str(self.tree.item(item, 'values')[self.tree['columns'].index(col)])
                cell_width = tkFont.Font().measure(cell_value)

                # Update max_width if the cell content is wider
                if cell_width > max_width:
                    max_width = cell_width

                cell_padding = 50
                widths.append(max_width + cell_padding)

            self.tree.column(col, width=max(widths), stretch=NO)

            self.update()

        # Close the database connection
        conn.close()
    
    def on_treeview_click(self, event):
        region = self.tree.identify('region', event.x, event.y)
        if region != 'heading':
            self.on_double_click(event)

    def on_double_click(self, event):
        curItem = self.tree.focus()
        if not curItem:
            return
        item_values = self.tree.item(curItem)['values']
        (
            testing_pk,
            student_id,
            last_name,
            first_name,
            laces,
            site,
            staff,
            date,
            pre_or_post,
            level,
            form,
            reading_raw,
            reading_scaled,
            reading_nrs,
            writing_raw,
            writing_scaled,
            writing_nrs,
            combined,
            got_msg,
            msg_area,
            test_notes
        ) = item_values

        self.item_dict = {
            "testing_pk": testing_pk,
            "student_id": student_id,
            "last_name": last_name,
            "first_name": first_name,
            "laces": laces,
            "site": site,
            "staff": staff,
            "date": date,
            "pre_or_post": pre_or_post,
            "level": level,
            "form": form,
            "reading_raw": reading_raw,
            "reading_scaled": reading_scaled,
            "reading_nrs": reading_nrs,
            "writing_raw": writing_raw,
            "writing_scaled": writing_scaled,
            "writing_nrs": writing_nrs,
            "combined": combined,
            "got_msg": got_msg,
            "msg_area": msg_area,
            "test_notes": test_notes
        }

        title_string = f'{last_name}, {first_name} | Level {level}, Form {form} | {date}'

        self.testing_record_window = Toplevel(self)
        self.testing_record_window.title(title_string)
        center_window(self.testing_record_window, width=700, height=700)

        # Add notebook widget 
        self.testing_record_notebook = ttk.Notebook(self.testing_record_window)

        self.testing_record_notebook.grid_rowconfigure(0, weight=1)
        self.testing_record_notebook.grid_columnconfigure(0, weight=1)

        # Create and add tabs
        self.test_results_tab = TestResultsTab(self.testing_record_notebook, self.default_font, self.item_dict['testing_pk'], self)
        self.test_display_tab = TestDisplayTab(self.testing_record_notebook, self.default_font, self.item_dict)
        self.test_progress_form_tab = TestProgressFormTab(self.testing_record_notebook, self.default_font, self.testing_record_window, self.item_dict)

        self.testing_record_notebook.add(self.test_results_tab, text='Results')
        self.testing_record_notebook.add(self.test_display_tab, text='Responses')
        self.testing_record_notebook.add(self.test_progress_form_tab, text='Progress Form')

        # Pack the Notebook widget to make it visible
        self.testing_record_notebook.grid(row=0, column=0, padx=10, pady=10, sticky='nsew')

        # Bind the event to handle tab changes
        self.testing_record_notebook.bind("<<NotebookTabChanged>>", self.on_tab_changed)
    
    def cleanup_temp_files(self):
            if hasattr(self, 'test_progress_form_tab') and self.test_progress_form_tab:
                self.test_progress_form_tab.cleanup_temp_files()
    
    def sort_by_column(self, col, reverse):
        # Fetch data from Treeview
        data = [(self.tree.set(child, col), child) for child in self.tree.get_children('')]

        # Convert data to appropriate types for sorting
        if col == 'Date':
            data = [(datetime.strptime(d, '%m-%d-%Y') if d else datetime.min, child) for d, child in data]
        elif col in ['Test ID', 'Student ID', 'LACES ID', 'Reading Raw', 'Reading Scaled', 'Writing Raw', 'Writing Scaled', 'Combined', 'MSG']:
            data = [(int(d) if d else 0, child) for d, child in data]
        else:
            data = [(d, child) for d, child in data]

        # Sort data
        data.sort(reverse=reverse)

        # Rearrange items in sorted positions
        for index, (val, child) in enumerate(data):
            self.tree.move(child, '', index)

        # Reverse sort next time
        self._sort_column = col
        self._sort_reverse = not reverse

        # Update heading with new sort direction
        self.tree.heading(col, command=lambda: self.sort_by_column(col, not reverse))

    def open_filter_window(self):
        self.filter_window = Toplevel(self)
        self.filter_window.title("Filter Records")

        Label(self.filter_window, text="Select Column:", font=self.default_font).grid(row=0, column=0, padx=10, pady=5, sticky='w')
        self.filter_column = StringVar()
        column_menu = AutocompleteCombobox(self.filter_window, textvariable=self.filter_column, font=self.default_font)
        column_menu['values'] = self.tree["columns"]
        column_menu.grid(row=0, column=1, padx=10, pady=5, sticky='w')

        Label(self.filter_window, text="Enter Filter Text:", font=self.default_font).grid(row=1, column=0, padx=10, pady=5, sticky='w')
        self.filter_text = StringVar()
        filter_entry = Entry(self.filter_window, textvariable=self.filter_text, font=self.default_font)
        filter_entry.grid(row=1, column=1, padx=10, pady=5, sticky='w')

        Button(self.filter_window, text="Apply Filter", command=self.apply_filter, font=self.default_font).grid(row=2, column=0, columnspan=2, pady=5)
        Button(self.filter_window, text="Reset", command=self.reset_filters, font=self.default_font).grid(row=3, column=0, columnspan=2, pady=5)

    def apply_filter(self):
        column = self.filter_column.get()
        filter_text = self.filter_text.get().lower().strip()
        filtered_rows = []

        if column and filter_text:
            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()

            query = """
                SELECT 
                    t.testing_pk,
                    t.student_id,
                    s.last_name AS last_name,
                    s.first_name AS first_name,
                    t.laces,
                    t.site,
                    t.staff,
                    t.date,
                    t.pre_or_post,
                    t.level,
                    t.form,
                    t.reading_raw,
                    t.reading_scaled,
                    t.reading_nrs,
                    t.writing_raw,
                    t.writing_scaled,
                    t.writing_nrs,
                    t.combined,
                    t.got_msg,
                    t.msg_area,
                    t.test_notes
                FROM
                    Testing t
                JOIN
                    Student s
                ON
                    s.student_id = t.student_id
                ORDER BY
                    t.date DESC
            """
            cursor.execute(query)
            rows = cursor.fetchall()
            conn.close()

            # Remove leading zeros from the filter text
            filter_text = re.sub(r'\b0+(\d)', r'\1', filter_text)

            # Apply the filter to the rows
            for row in rows:
                formatted_row = list(row)

                # Format dates for filtering
                date = formatted_row[7]

                # Convert date to different searchable formats
                def date_matches(date_str, search_text):
                    try:
                        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                        formatted_date = [
                            date_obj.strftime('%Y-%m-%d'),  # Full date
                            date_obj.strftime('%m-%d-%Y'),  # Alternative full date format
                            date_obj.strftime('%y-%m-%d'),  # Short year with full date
                            date_obj.strftime('%m-%d'),     # Month-day only
                            date_obj.strftime('%Y-%m'),     # Month
                            date_obj.strftime('%m'),        # Year and month
                            date_obj.strftime('%Y'),        # Year only
                            date_obj.strftime('%y'),        # Short year only
                            date_obj.strftime('%B'),        # full month name
                            date_obj.strftime('%b'),        # short month name
                            date_obj.strftime('%B %Y'),     # full month name with year
                            date_obj.strftime('%b %Y')      # short month name with yar                        
                        ]

                        # Normalize month and day by removing leading zeros
                        normalized_date = [
                            re.sub(r'\b0+(\d)', r'\1', d) for d in formatted_date
                        ]

                        # Check if any normalized version matches the search text
                        return any(search_text in nd.lower() for nd in normalized_date)
                    except ValueError:
                        return False

                # Check if the filter text matches in any format
                if column == 'Date':
                    if date_matches(date, filter_text):
                        filtered_rows.append(formatted_row)
                else:
                    col_index = self.tree['columns'].index(column)
                    if filter_text in str(formatted_row[col_index]).lower():
                        filtered_rows.append(formatted_row)

        self.filter_window.destroy()
        self.populate_treeview(filtered_rows if filtered_rows else None)

    def reset_filters(self):
        self.filter_text.set('')
        self.filter_column.set('')
        self.populate_treeview()

    def create_registration_form(self):
        self.create_registration_form_window = Toplevel(self)
        self.create_registration_form_window.state("zoomed")
        # center_window(self.create_registration_form_window, width=screen_width-10, height=screen_height-10)
        self.create_registration_form_window.title("Create Registration Form")
        RegistrationFormWindow(self.create_registration_form_window, self.default_font)

    def add_test(self):
        run_test_grading_gui(self.default_font, on_save=self.populate_treeview)

    def on_tab_changed(self, *args):
            selected_tab = self.testing_record_notebook.tab(self.testing_record_notebook.select(), "text")
            if selected_tab == "Progress Form":
                # center_window(self.add_student_window, width=600, height=800)
                self.testing_record_window.state("zoomed")
            else:
                self.testing_record_window.state("normal")
                center_window(self.testing_record_window, width=650, height=950)

class TestResultsTab(ttk.Frame):
    def __init__(self, parent, default_font, test_id, testing_tab):
        super().__init__(parent)
        self.default_font = default_font
        self.test_id = test_id  # The ID of the selected test result
        self.testing_tab = testing_tab
        self.fields = [
            ("Laces:", StringVar()),
            ("Site:", StringVar()),
            ("Staff:", StringVar()),
            ("Primary Class:", StringVar()),
            ("Date:", StringVar()),
            ("Pre/Post Test:", StringVar()),
            ("Level:", IntVar()),
            ("Form:", StringVar()),
            ("Reading Raw:", StringVar()),
            ("Reading Scaled:", IntVar()),
            ("Reading NRS:", StringVar()),
            ("Writing Raw:", StringVar()),
            ("Writing Scaled:", IntVar()),
            ("Writing NRS:", StringVar()),
            ("Combined Score:", IntVar()),
            ("Got Message:", IntVar()),
            ("Message Area:", StringVar()),
            ("Locator Raw:", StringVar()),
            ("Locator Result:", StringVar()),
            ("Test Notes:", StringVar())
        ]
        self.edit_mode = False
        self.create_widgets()
        self.populate_data()  # Populate the fields with the data for the selected test

    def create_widgets(self):
        # Create a bold font
        bold_font = tkFont.Font(weight="bold")

        # Create a canvas for scrolling
        self.canvas = Canvas(self, height=800)  # Adjust the height to make the scrollable area larger
        self.canvas.grid(row=0, column=0, columnspan=3, sticky='nsew')

        # Add a vertical scrollbar linked to the canvas
        scrollbar_y = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        scrollbar_y.grid(row=0, column=3, sticky='ns')
        self.canvas.configure(yscrollcommand=scrollbar_y.set)

        # Create a frame inside the canvas to hold the labels and entries
        self.scrollable_frame = ttk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw", width=400)  # Set width for consistent layout

        # Update scroll region when the frame changes size
        self.scrollable_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))

        # Bind mouse wheel scrolling to the canvas
        self.canvas.bind_all("<MouseWheel>", lambda e: self.on_mouse_wheel(e, self.canvas))

        # Create labels, data labels, and entries for each field inside the scrollable frame
        self.labels = {}
        self.data_labels = {}
        self.entries = {}

        for i, (label_text, string_var) in enumerate(self.fields):
                label = ttk.Label(self.scrollable_frame, text=label_text, font=self.default_font)
                data_label = ttk.Label(self.scrollable_frame, textvariable=string_var, font=self.default_font)
                entry = Entry(self.scrollable_frame, textvariable=string_var, font=self.default_font)

                label.grid(row=i, column=0, padx=5, pady=5, sticky='w')
                data_label.grid(row=i, column=1, padx=5, pady=5, sticky='w')
                entry.grid(row=i, column=2, padx=5, pady=5, sticky='w')
                entry.grid_remove()  # Hide the entry widget initially

                self.labels[label_text] = label
                self.data_labels[label_text] = data_label
                self.entries[label_text] = entry

        # Add Edit, Save, and Delete buttons below the scrollable area
        self.edit_button = Button(self, text="Edit", command=self.enable_editing, font=self.default_font)
        self.edit_button.grid(row=1, column=0, padx=10, pady=10, sticky='w')

        self.save_button = Button(self, text="Save", command=self.save_changes, font=self.default_font)
        self.save_button.grid(row=1, column=1, padx=10, pady=10, sticky='w')
        self.save_button.grid_remove()  # Hide the Save button initially

        self.delete_button = Button(self, text="Delete", command=self.delete_test, font=self.default_font)
        self.delete_button.grid(row=1, column=2, padx=10, pady=10, sticky='e')

    def populate_data(self):
        """Populate the labels and entries with data from the database for the selected test result."""
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()

        cursor.execute('''
            SELECT laces, site, staff, primary_class, date, pre_or_post, level, form,
                   reading_raw, reading_scaled, reading_nrs, writing_raw, writing_scaled,
                   writing_nrs, combined, got_msg, msg_area, locator_raw, locator_result, test_notes
            FROM Testing WHERE testing_pk = ?
        ''', (self.test_id,))
        test_data = cursor.fetchone()

        conn.close()

        # Set the data into the StringVars/IntVars
        if test_data:
            for (label_text, string_var), value in zip(self.fields, test_data):
                if value is None:
                    # Remove both the label and the entry if value is None
                    self.labels[label_text].grid_remove()
                    self.data_labels[label_text].grid_remove()
                    self.entries[label_text].grid_remove()
                else:
                    if label_text == "Date:":
                        string_var.set(format_date(value))  # Assuming format_date is a custom function to format dates
                    else:
                        string_var.set(value)

    def enable_editing(self):
        """Enable editing mode for the current test."""
        for label_text in self.labels.keys():
            self.data_labels[label_text].grid_remove()
            self.entries[label_text].grid()
        self.edit_button.config(state='disabled')
        self.save_button.grid()  # Show the Save button

    def save_changes(self):
        """Save changes made to the test result."""
        updated_data = {label: var.get() for label, var in self.fields}

        if self.test_id is not None:
            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()

            # Update the Testing table
            cursor.execute('''
                UPDATE Testing
                SET laces = ?, site = ?, staff = ?, primary_class = ?, date = ?, pre_or_post = ?, level = ?, form = ?,
                    reading_raw = ?, reading_scaled = ?, reading_nrs = ?, writing_raw = ?, writing_scaled = ?,
                    writing_nrs = ?, combined = ?, got_msg = ?, msg_area = ?, locator_raw = ?, locator_result = ?, test_notes = ?
                WHERE testing_pk = ?
            ''', (
                updated_data['Laces:'], updated_data['Site:'], updated_data['Staff:'], updated_data['Primary Class:'],
                format_date(updated_data['Date:'], out=True), updated_data['Pre/Post Test:'], updated_data['Level:'], updated_data['Form:'],
                updated_data['Reading Raw:'], updated_data['Reading Scaled:'], updated_data['Reading NRS:'],
                updated_data['Writing Raw:'], updated_data['Writing Scaled:'], updated_data['Writing NRS:'],
                updated_data['Combined Score:'], updated_data['Got Message:'], updated_data['Message Area:'],
                updated_data['Locator Raw:'], updated_data['Locator Result:'],
                updated_data['Test Notes:'], self.test_id
            ))

            conn.commit()
            conn.close()

            self.disable_editing()

    def disable_editing(self):
        """Disable editing mode."""
        for label_text in self.labels.keys():
            self.entries[label_text].grid_remove()
            self.data_labels[label_text].grid()
        self.edit_button.config(state='normal')
        self.save_button.grid_remove()

    def delete_test(self):
        """Delete the currently selected test result."""
        confirm_delete = messagebox.askyesno("Delete Test", "Are you sure you want to delete this test result?")
        if confirm_delete and self.test_id is not None:
            conn = sqlite3.connect(DB_FILE_PATH)
            cursor = conn.cursor()
            cursor.execute('DELETE FROM Testing WHERE testing_pk = ?', (self.test_id,))
            conn.commit()
            conn.close()

            messagebox.showinfo("Success", "Test result deleted successfully.")

            self.testing_tab.populate_treeview()

    def on_mouse_wheel(self, event, canvas):
        """Handle mouse wheel events for scrolling the canvas."""
        if event.delta:
            scroll_speed = int(event.delta / 120)
        else:
            scroll_speed = event.num if event.num != 0 else 0
        scroll_speed = -scroll_speed
        if scroll_speed == 0:
            scroll_speed = -1 if event.delta < 0 else 1
        canvas.yview_scroll(scroll_speed, "units")

class TestDisplayTab(ttk.Frame):
    def __init__(self, parent, default_font, item_dict):
        super().__init__(parent)
        self.default_font = default_font
        self.item_dict = item_dict

        # Add student's answers to item_dict
        conn = sqlite3.connect(DB_FILE_PATH)
        cursor = conn.cursor()
        query = '''
            SELECT
                reading_answers,
                writing_answers,
                writing_folio_answers,
                locator_answers,
                locator_result
            FROM
                Testing
            WHERE
                testing_pk = ?
        '''
        cursor.execute(query, (item_dict['testing_pk'],))
        answer_data = cursor.fetchone()

        new_fields = [                
            'reading_answers',
            'writing_answers',
            'writing_folio_answers',
            'locator_answers',
            'locator_result'
        ]
        for i, new_field in enumerate(new_fields):
            item_dict.update({new_field:answer_data[i]})


class TestProgressFormTab(ttk.Frame):
    def __init__(self, parent, default_font, testing_record_window, item_dict):
        super().__init__(parent)
        self.default_font = default_font
        self.testing_record_window = testing_record_window
        self.item_dict = item_dict
        self.temp_files = []  # To track temporary files
        ORIGINAL_WIDTH = 1200  # Replace with the actual width of the original image
        ORIGINAL_HEIGHT = 1600  # Replace with the actual height of the original image
        self.original_image_size = (ORIGINAL_WIDTH, ORIGINAL_HEIGHT)  # Set these values according to your original image size
        
        self.create_widgets()
        self.populate_treeview()

    def create_widgets(self):
        # self.image_label = tk.Label(self)
        # self.image_label.grid(row=0, column=0, padx=5, pady=5)

        # Load Progress Form image
        def load_and_resize_image(self, new_size=None):
            IMAGE_PATH = os.path.join(get_test_png_path(), 'Student Progress Form ESOL_page_1.png')
            image = Image.open(IMAGE_PATH)
            MAX_SIZE = (800, 1000)  # Max width and height in pixels
            if os.name != 'nt':
                new_size = (1000, 1200)
            else:
                new_size = (1800, 1200)

            # If a new size is provided, use it; otherwise, resize based on original dimensions
            if new_size:
                ratio = min(new_size[0] / image.width, new_size[1] / image.height)
            else:
                ratio = min(MAX_SIZE[0] / image.width, MAX_SIZE[1] / image.height)
            
            new_size = (int(image.width * ratio), int(image.height * ratio))
            resized_image = image.resize(new_size, Image.LANCZOS)
            return ImageTk.PhotoImage(resized_image), ratio, new_size

        def update_test_display(self, scaling_factor=1.0):
            # Resize the image based on the scaling factor
            new_image_size = (int(self.original_image_size[0] * scaling_factor), int(self.original_image_size[1] * scaling_factor))
            test_image, ratio, new_size = load_and_resize_image(self, new_size=new_image_size)
            self.image_label.config(image=test_image)
            self.image_label.image = test_image

        # Setup canvas and scrollbars
        self.canvas = tk.Canvas(self, width=1850, height=950)
        v_scrollbar = tk.Scrollbar(self, orient=tk.VERTICAL, command=self.canvas.yview)
        h_scrollbar = tk.Scrollbar(self, orient=tk.HORIZONTAL, command=self.canvas.xview)
        
        self.canvas.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        # Place the vertical scrollbar on the right side
        v_scrollbar.grid(row=0, column=1, sticky='ns')

        # Place the horizontal scrollbar on the bottom
        h_scrollbar.grid(row=1, column=0, sticky='ew')

        # Place the canvas in the remaining space, filling the area
        self.canvas.grid(row=0, column=0, sticky='nsew')

        # Configure the grid to allow the canvas to expand
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        # Create a frame inside the canvas to hold all content
        self.content_frame = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.content_frame, anchor="nw")
        
        # Ensure the scroll region matches the content size
        self.content_frame.bind("<Configure>", lambda event: self.canvas.configure(scrollregion=self.canvas.bbox("all")))

        # Bind the right-click event (Button-3 in Tkinter) to the function
        self.testing_record_window.bind("<Button-3>", right_click)
        self.canvas.bind_all("<MouseWheel>", lambda e: on_mouse_wheel(e, self.canvas))

        # center_window(self.answer_sheet_window, 1600, 850)

        # Label to display the test image
        self.image_label = tk.Label(self.content_frame)
        self.image_label.grid(row=0, column=0, padx=5, pady=5)

        center_window(self.testing_record_window, 1903, 964)

        # Update the display with the constant image and size
        update_test_display(self)

        # Plot labels over progess form
        self.site_label = tk.Label(self.content_frame, text=self.item_dict["site"], bg="white", fg="black", padx=0, pady=0, borderwidth=0)
        if os.name == 'nt':
            self.site_label.place(x=183, y=59)
            self.site_label.config(font=("Calibri", 14))

        self.staff_label = tk.Label(self.content_frame, text=self.item_dict["staff"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, )
        if os.name == 'nt':
            self.staff_label.place(x=183, y=109)
            self.staff_label.config(font=("Calibri", 14))

        # Eventually, we'll want to check the primary class
        self.class_label = tk.Label(self.content_frame, text='ESOL', bg="white", fg="black", padx=2, pady=2, borderwidth=3, relief='solid')
        if os.name == 'nt':
            self.class_label.place(x=192, y=159)
            self.class_label.config(font=("Calibri", 14, "bold"))

        self.last_name_label = tk.Label(self.content_frame, text=self.item_dict["last_name"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, )
        if os.name == 'nt':
            self.last_name_label.place(x=171, y=243)
            self.last_name_label.config(font=("Calibri", 14))

        self.first_name_label = tk.Label(self.content_frame, text=self.item_dict["first_name"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, )
        if os.name == 'nt':
            self.first_name_label.place(x=606, y=243)
            self.first_name_label.config(font=("Calibri", 14))

        self.laces_label = tk.Label(self.content_frame, text=self.item_dict["laces"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, )
        if os.name == 'nt':
            self.laces_label.place(x=1364, y=243)
            self.laces_label.config(font=("Calibri", 14))

        if self.item_dict['pre_or_post'] == 'Post':
            self.date_label = tk.Label(self.content_frame, text=self.item_dict["date"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.date_label.place(x=708, y=471)
                self.date_label.config(font=("Calibri", 11)) 

            self.reading_level_label = tk.Label(self.content_frame, text=f'{self.item_dict["level"]}{self.item_dict["form"]}', bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.reading_level_label.place(x=661, y=538)
                self.reading_level_label.config(font=("Calibri", 11)) 

            self.writing_level_label = tk.Label(self.content_frame, text=f'{self.item_dict["level"]}{self.item_dict["form"]}', bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.writing_level_label.place(x=661, y=572)
                self.writing_level_label.config(font=("Calibri", 11))  

            self.reading_scaled_label = tk.Label(self.content_frame, text=self.item_dict["reading_scaled"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.reading_scaled_label.place(x=814, y=538)
                self.reading_scaled_label.config(font=("Calibri", 11)) 

            self.writing_scaled_label = tk.Label(self.content_frame, text=self.item_dict["writing_scaled"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.writing_scaled_label.place(x=814, y=572)
                self.writing_scaled_label.config(font=("Calibri", 11)) 

            self.combined_label = tk.Label(self.content_frame, text=self.item_dict["combined"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.combined_label.place(x=814, y=606)
                self.combined_label.config(font=("Calibri", 11))
        else:
            self.date_label = tk.Label(self.content_frame, text=self.item_dict["date"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.date_label.place(x=708-308, y=471)
                self.date_label.config(font=("Calibri", 11)) 

            self.reading_level_label = tk.Label(self.content_frame, text=f'{self.item_dict["level"]}{self.item_dict["form"]}', bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.reading_level_label.place(x=661-300, y=538)
                self.reading_level_label.config(font=("Calibri", 11)) 

            self.writing_level_label = tk.Label(self.content_frame, text=f'{self.item_dict["level"]}{self.item_dict["form"]}', bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.writing_level_label.place(x=661-300, y=572)
                self.writing_level_label.config(font=("Calibri", 11))  

            self.reading_scaled_label = tk.Label(self.content_frame, text=self.item_dict["reading_scaled"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.reading_scaled_label.place(x=814-310, y=538)
                self.reading_scaled_label.config(font=("Calibri", 11)) 

            self.writing_scaled_label = tk.Label(self.content_frame, text=self.item_dict["writing_scaled"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.writing_scaled_label.place(x=814-310, y=572)
                self.writing_scaled_label.config(font=("Calibri", 11)) 

            self.combined_label = tk.Label(self.content_frame, text=self.item_dict["combined"], bg="white", fg="black", padx=0, pady=0, borderwidth=0, anchor='center')
            if os.name == 'nt':
                self.combined_label.place(x=814-310, y=606)
                self.combined_label.config(font=("Calibri", 11))

        self.print_button = tk.Button(self.content_frame, text="Print",command=self.print_progress_form, borderwidth=0, highlightthickness=0, bd=0, padx=0, pady=0)
        self.print_button.place(x=1600, y=58)
        self.print_button.config(font=("Calibri", 14)) 
        self.save_button = tk.Button(self.content_frame, text="Save",command=self.save_progress_form, borderwidth=0, highlightthickness=0, bd=0, padx=0, pady=0)
        self.save_button.place(x=1600, y=120)
        self.save_button.config(font=("Calibri", 14)) 
    
    def populate_treeview(self):
        pass

    def print_progress_form(self):
        pass

    def getter(self, widget):
        # Ensure the testing_record_window is focused and brought to the front
        self.testing_record_window.lift()
        self.testing_record_window.focus_force()

        # Temporarily hide the Print and Save buttons 
        self.print_button.place_forget() 
        self.save_button.place_forget()

        # Give the window some time to redraw and be in focus
        self.testing_record_window.update()

        # Get the visible region of the canvas
        x = widget.winfo_rootx()
        y = widget.winfo_rooty()
        x1 = x + widget.winfo_width()
        y1 = y + widget.winfo_height()

        # Capture the specific area of the screen corresponding to the widget's visible region
        return ImageGrab.grab(bbox=(x, y, x1, y1))

    def save_progress_form(self):
        # Prompt the user to select the save location for the PDF
        pdf_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")], title="Save as PDF")
        if not pdf_path:  # If the user cancels the save dialog
            return

        # Temporarily hide the Print and Save buttons
        self.print_button.place_forget()
        self.save_button.place_forget()

        # Scroll to the top and capture the visible area
        self.canvas.yview_moveto(0)
        top_image = self.getter(self.canvas)

        # Scroll down slightly to reveal the hidden section
        self.canvas.yview_scroll(1, "pages")
        self.testing_record_window.update()
        bottom_image = self.getter(self.canvas)

        # Combine the two images
        overlap_height = 745  # Adjust overlap if needed
        combined_height = top_image.height + (bottom_image.height - overlap_height)
        combined_image = Image.new('RGB', (top_image.width, combined_height))
        combined_image.paste(top_image, (0, 0))
        combined_image.paste(bottom_image, (0, top_image.height - overlap_height))

        # Define crop values (adjust as needed)
        crop_x = 10   # Crop 10 pixels from the left
        crop_y = 10   # Crop 10 pixels from the top
        crop_x1 = 50  # Crop 50 pixels from the right
        crop_y1 = 20  # Crop 20 pixels from the bottom

        # Apply cropping to the combined image
        final_image = combined_image.crop((crop_x, crop_y, combined_image.width - crop_x1, combined_image.height - crop_y1))

        # Resize to fit within 3300x2550 pixels (11x8.5 inches at 300 DPI) while maintaining aspect ratio
        target_width = 3200
        target_height = 2450

        # Calculate aspect ratio and determine the new size
        image_ratio = final_image.width / final_image.height
        target_ratio = target_width / target_height

        if image_ratio > target_ratio:
            # Fit to width
            new_width = target_width
            new_height = int(target_width / image_ratio)
        else:
            # Fit to height
            new_height = target_height
            new_width = int(target_height * image_ratio)

        # Resize the image proportionally
        resized_image = final_image.resize((new_width, new_height), Image.LANCZOS)

        # Create a temporary file to save the resized PNG
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_png:
            resized_image.save(temp_png.name)
            temp_png_path = temp_png.name

        # Convert the PNG to PDF
        resized_image.convert("RGB").save(pdf_path, "PDF", quality=95, resolution=300)

        # Delete the temporary PNG file after converting to PDF
        os.remove(temp_png_path)

        # Restore the Print and Save buttons
        self.print_button.place(x=1600, y=58)
        self.save_button.place(x=1600, y=120)

    def print_progress_form(self):
        # Temporarily hide the Print and Save buttons
        self.print_button.place_forget()
        self.save_button.place_forget()

        # Scroll to the top and capture the visible area
        self.canvas.yview_moveto(0)
        top_image = self.getter(self.canvas)

        # Scroll down slightly to reveal the hidden section
        self.canvas.yview_scroll(1, "pages")
        self.testing_record_window.update()
        bottom_image = self.getter(self.canvas)

        # Combine the two images
        overlap_height = 745
        combined_height = top_image.height + (bottom_image.height - overlap_height)
        combined_image = Image.new('RGB', (top_image.width, combined_height))
        combined_image.paste(top_image, (0, 0))
        combined_image.paste(bottom_image, (0, top_image.height - overlap_height))

        # Define crop values
        crop_x = 10
        crop_y = 10
        crop_x1 = 50
        crop_y1 = 20

        # Apply cropping to the combined image
        final_image = combined_image.crop((crop_x, crop_y, combined_image.width - crop_x1, combined_image.height - crop_y1))

        # Resize to fit within 3200x2450 pixels (slightly smaller than 11x8.5 inches at 300 DPI)
        target_width = 3200
        target_height = 2450

        # Calculate aspect ratio and determine the new size
        image_ratio = final_image.width / final_image.height
        target_ratio = target_width / target_height

        if image_ratio > target_ratio:
            # Fit to width
            new_width = target_width
            new_height = int(target_width / image_ratio)
        else:
            # Fit to height
            new_height = target_height
            new_width = int(target_height * image_ratio)

        # Resize the image proportionally
        resized_image = final_image.resize((new_width, new_height), Image.LANCZOS)

        # Create a temporary file to save the resized PNG
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_png:
            resized_image.save(temp_png.name)
            temp_png_path = temp_png.name
            self.temp_files.append(temp_png_path)  # Track the temp file

        # Create a temporary file to save the PDF
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
            temp_pdf_path = temp_pdf.name
            resized_image.convert("RGB").save(temp_pdf_path, "PDF", quality=95, resolution=300)
            self.temp_files.append(temp_pdf_path)  # Track the temp file

        # Open the print dialog to choose the printer and print
        self._print_pdf(temp_pdf_path)

        # Restore the Print and Save buttons
        self.print_button.place(x=1600, y=58)
        self.save_button.place(x=1600, y=120)

    def _print_pdf(self, pdf_path):
        if platform.system() == "Darwin":  # macOS
            subprocess.run(["open", "-a", "Preview", pdf_path])
        elif platform.system() == "Windows":  # Windows
            # os.startfile(pdf_path, "print")
            for proc in psutil.process_iter(['pid', 'name']):
                if proc.info['name'] in ['AcroRd32.exe', 'acrord32.exe']:
                    print(f"Terminating {proc.info['name']} with PID {proc.info['pid']}")
                    proc.kill()

            acrobat_path = r"C:\Program Files\Adobe\Acrobat DC\Acrobat\Acrobat.exe"
            subprocess.Popen([acrobat_path, "/p", pdf_path])
        elif platform.system() == "Linux":  # Linux
            subprocess.run(["xdg-open", pdf_path])

    def cleanup_temp_files(self):
        # Delete the temporary files when the application is closing
        for file_path in self.temp_files:
            try:
                os.remove(file_path)
            except Exception as e:
                print(f"Error deleting temporary file: {file_path}, {str(e)}")

class StatisticsTab(ttk.Frame):
    def __init__(self, parent, default_font, bold_font):
        super().__init__(parent)
        self.default_font = default_font
        self.bold_font = bold_font
        self.create_widgets()      

    def create_widgets(self):
        # Create a Notebook
        self.notebook = ttk.Notebook(self)

        # Create and add tabs
        self.attendance_tab = StatisticsAttendanceTab(self.notebook, self.default_font, self.bold_font)
        self.testing_tab = StatisticsTestingTab(self.notebook, self.default_font, self.bold_font)

        self.notebook.add(self.attendance_tab, text='Attendance')
        # self.notebook.add(self.site_tab, text='Site')
        self.notebook.add(self.testing_tab, text='Testing')

        # Pack the Notebook widget to make it visible
        self.notebook.grid(row=0, column=0, padx=10, pady=10, sticky='nsew')

        # Bind the tab change event
        self.notebook.bind("<<NotebookTabChanged>>", self.on_tab_changed)

        # Configure the master window to expand properly
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

    def on_tab_changed(self, event):
        # Check if the AttendanceAll tab is selected
        selected_tab = self.notebook.tab(self.notebook.select(), "text")
        if selected_tab == 'All':
            self.attendance_tab.populate_treeview()

class StatisticsAttendanceTab(ttk.Frame):
    def __init__(self, parent, default_font, bold_font):
        super().__init__(parent)
        self.default_font = default_font
        self.bold_font = bold_font
        self.start_date = None
        self.end_date = None
        # Set the default fiscal year range for FY24-25
        self.fiscal_start = datetime(2025, 7, 1)
        self.fiscal_end = datetime(2026, 6, 30)
        self.range_var = StringVar(value=f"FY Range:\n{self.fiscal_start.strftime('%#m/%#d/%y')} - {self.fiscal_end.strftime('%#m/%#d/%y')}")
        
        # Variables for each checkbutton
        self.data_labels_var = IntVar(value=0)
        self.regression_var = IntVar(value=0)  # Default to checked for regression line
        self.prediction_var = IntVar(value=0)  # Default to unchecked for prediction line
        self.outliers_var = IntVar(value=1)    # Default to checked for remove outliers

        # Configure grid layout for the entire frame
        self.grid_columnconfigure(0, weight=1, minsize=300)  # Ensure attendance_frame has space
        self.grid_columnconfigure(1, weight=3)               # Give plot_frame more space
        # self.grid_columnconfigure(2, weight=0, minsize=250) 

        self.create_widgets()

    def create_widgets(self):
        # Step 1: Create attendance_frame on the left side of the tab using grid
        self.attendance_frame = ttk.Frame(self)
        self.attendance_frame.grid(row=0, column=0, sticky="nsw", padx=2, pady=2)
        self.columnconfigure(0, weight=1, minsize=300)  # Make attendance_frame as small as possible

        # Plot frame in the middle
        self.plot_frame = ttk.Frame(self)
        self.plot_frame.grid(row=0, column=1, sticky="nsew", padx=10, pady=2)
        self.columnconfigure(1, weight=3)  # Let plot frame take most space

        # Summary frame on the far right
        self.summary_frame = ttk.Frame(self)
        self.summary_frame.grid(row=0, column=2, sticky="nsw", padx=10, pady=2)
        self.columnconfigure(2, weight=2, minsize=200)  # Ensure summary frame is visible but not stretched

        self.rowconfigure(0, weight=1)

        # Step 2: Create attendance_canvas within attendance_frame
        self.attendance_canvas = Canvas(self.attendance_frame, width=250)  # Reduce width to avoid stretching
        self.attendance_canvas.grid(row=0, column=0, sticky="nsw")

        row_counter = 0  # Keep track of row positions

        self.date_range_button = Button(self.attendance_canvas, text="Select Date Range", command=self.open_date_range_window, font=self.default_font)
        self.date_range_button.grid(row=row_counter, column=0, padx=2, pady=10, sticky='w')
        row_counter += 1

        self.date_label = ttk.Label(self.attendance_canvas, textvariable=self.range_var, font=self.default_font, wraplength=250)
        self.date_label.grid(row=row_counter, column=0, padx=2, pady=2, sticky='w')
        row_counter += 1

        frequency_label = ttk.Label(self.attendance_canvas, text="Frequency:", font=self.bold_font)
        frequency_label.grid(row=row_counter, column=0, padx=2, pady=2, sticky='w')
        row_counter += 1

        frequency_options = ["Daily", "Weekly", "Monthly", "FY23-24", "FY24-25", "FY25-26"]
        self.frequency_var = StringVar(value=frequency_options[0])  
        frequency_dropdown = AutocompleteCombobox(self.attendance_canvas, textvariable=self.frequency_var, values=frequency_options, state="readonly")
        frequency_dropdown.grid(row=row_counter, column=0, padx=2, pady=2, sticky='w')
        frequency_dropdown.bind("<<ComboboxSelected>>", lambda e: self.update_plot())
        row_counter += 1

        # Step 4: Class Sites Section (Scrollable Area)
        site_label = ttk.Label(self.attendance_canvas, text="Class Sites:", font=self.bold_font)
        site_label.grid(row=row_counter, column=0, padx=2, pady=2, sticky='w')
        row_counter += 1

        # Create a frame inside the canvas just for checkboxes
        self.checkbox_canvas_frame = ttk.Frame(self.attendance_canvas)
        self.checkbox_canvas_frame.grid(row=row_counter, column=0, sticky="nsew", padx=2, pady=2)

        # Create a Canvas to allow scrolling
        self.checkbox_canvas = Canvas(self.checkbox_canvas_frame, height=250, width=230)  # Set width to match attendance frame
        self.checkbox_canvas.grid(row=0, column=0, sticky="nsew")

        # Add a vertical scrollbar to the checkbox canvas
        vsb_checkbox_canvas = ttk.Scrollbar(self.checkbox_canvas_frame, orient="vertical", command=self.checkbox_canvas.yview)
        vsb_checkbox_canvas.grid(row=0, column=1, sticky='ns')

        self.checkbox_canvas.configure(yscrollcommand=vsb_checkbox_canvas.set)

        # Step 5: Create a frame inside the checkbox canvas for checkbuttons
        self.checkbox_frame = ttk.Frame(self.checkbox_canvas)
        self.checkbox_window = self.checkbox_canvas.create_window((0, 0), window=self.checkbox_frame, anchor="nw")

        # Step 6: Add checkbuttons inside the `checkbox_frame`
        site_names = get_site_names(ignore_empty=True)
        self.site_vars = {site: IntVar(value=1) for site in site_names}

        checkbutton_row = 0
        for site, var in self.site_vars.items():
            check_button = Checkbutton(self.checkbox_frame, text=site, variable=var, font=self.default_font, command=self.update_plot)
            check_button.grid(row=checkbutton_row, column=0, padx=2, pady=2, sticky='w')
            checkbutton_row += 1

        # Update scroll region when the content changes
        self.checkbox_frame.update_idletasks()
        self.checkbox_canvas.config(scrollregion=self.checkbox_canvas.bbox("all"))

        row_counter += 1  # Move past the checkboxes

        # Step 7: Regression and Prediction Line Checkbuttons
        line_type_label = ttk.Label(self.attendance_canvas, text="Line:", font=self.bold_font)
        line_type_label.grid(row=row_counter, column=0, padx=2, pady=2, sticky='w')
        row_counter += 1

        data_labels = Checkbutton(self.attendance_canvas, text="Data Labels", variable=self.data_labels_var, font=self.default_font, command=self.update_plot)
        data_labels.grid(row=row_counter, column=0, padx=2, pady=2, sticky='w')
        row_counter += 1

        regression_check = Checkbutton(self.attendance_canvas, text="Regression Line", variable=self.regression_var, font=self.default_font, command=self.update_plot)
        regression_check.grid(row=row_counter, column=0, padx=2, pady=2, sticky='w')
        row_counter += 1

        prediction_check = Checkbutton(self.attendance_canvas, text="Prediction Line", variable=self.prediction_var, font=self.default_font, command=self.update_plot)
        prediction_check.grid(row=row_counter, column=0, padx=2, pady=2, sticky='w')
        row_counter += 1

        line_type_label = ttk.Label(self.attendance_canvas, text="Data:", font=self.bold_font)
        line_type_label.grid(row=row_counter, column=0, padx=2, pady=2, sticky='w')
        row_counter += 1

        remove_outliers_check = Checkbutton(self.attendance_canvas, text="Remove Outliers", variable=self.outliers_var, font=self.default_font, command=self.update_plot)
        remove_outliers_check.grid(row=row_counter, column=0, padx=2, pady=2, sticky='w')
        row_counter += 1


        # Set up the plot area and initial plot
        self.update_plot()
    
    def open_date_range_window(self):
        # Open a new window with calendar for date range selection
        range_window = Toplevel(self)
        range_window.title("Select Date Range")
        center_window(range_window, 800, 450)

        # Configure the grid layout
        range_window.columnconfigure(0, weight=1)
        range_window.columnconfigure(1, weight=1)
        range_window.columnconfigure(2, weight=1)
        range_window.rowconfigure(1, weight=1)

        # Label for displaying selected date range
        large_font = font.Font(family=self.default_font.cget("family"), size=self.default_font.cget("size")+2, weight="bold")
        date_label = ttk.Label(range_window, textvariable=self.range_var, font=large_font)
        date_label.grid(row=0, column=0, columnspan=3, pady=5, sticky="ew")

        # Calendar widget
        self.cal = Calendar(range_window, selectmode='day',
                            year=datetime.today().year, month=datetime.today().month,
                            day=datetime.today().day, font=large_font,
                            headersbackground="lightblue",
                            background="lightgrey", foreground="black",
                            selectbackground="blue", selectforeground="white")
        self.cal.grid(row=1, column=0, columnspan=3, pady=10)

        # Methods for highlighting and clearing date range
        def highlight_range(start, end):
            """Highlight the range of dates between start and end."""
            clear_highlight()  # Clear previous highlights
            current_date = start
            while current_date <= end:
                self.cal.calevent_create(current_date, '', 'highlight')
                current_date += timedelta(days=1)
            self.cal.tag_config('highlight', background='blue', foreground='white')

        def clear_highlight():
            """Clear highlighted dates."""
            for event in self.cal.get_calevents(tag='highlight'):
                self.cal.calevent_remove(event)

        # Variables for tracking start and end dates within this function
        local_start_date = None
        local_end_date = None

        def on_date_click(event):
            """Handle date selection in calendar widget."""
            nonlocal local_start_date, local_end_date
            selected_date = self.cal.get_date()
            date_obj = datetime.strptime(selected_date, "%m/%d/%y")

            if not local_start_date:
                # Select start date
                local_start_date = date_obj
                self.range_var.set(f"Start Date Selected: {local_start_date.strftime('%m/%d/%Y')}")
            elif not local_end_date:
                # Select end date
                local_end_date = date_obj
                if local_end_date < local_start_date:
                    local_start_date, local_end_date = local_end_date, local_start_date
                highlight_range(local_start_date, local_end_date)
                self.range_var.set(f"Selected Date Range: {local_start_date.strftime('%m/%d/%Y')} - {local_end_date.strftime('%m/%d/%Y')}")
            else:
                # Reset if both dates are already selected
                local_start_date = date_obj
                local_end_date = None
                clear_highlight()
                self.range_var.set(f"Start Date Selected: {local_start_date.strftime('%m/%d/%Y')}")

        self.cal.bind("<<CalendarSelected>>", on_date_click)

        # Confirm and Cancel buttons
        def confirm_selection():
            """Save the selected date range and close the window."""
            self.start_date, self.end_date = local_start_date, local_end_date
            range_window.destroy()
            self.update_plot()

        def cancel_selection():
            """Close the window without saving changes."""
            range_window.destroy()

        def reset_selection():
            """Reset the date range to the default fiscal year."""
            self.start_date = self.fiscal_start
            self.end_date = self.fiscal_end
            self.range_var.set(f"FY Range: {self.fiscal_start.strftime('%m/%d/%Y')} - {self.fiscal_end.strftime('%m/%d/%Y')}")
            clear_highlight()
            highlight_range(self.fiscal_start, self.fiscal_end)
            self.update_plot()       

        # Add Confirm, Cancel, and Reset buttons
        confirm_button = Button(range_window, text="Confirm", command=confirm_selection, font=self.default_font)
        confirm_button.grid(row=2, column=0, padx=10, pady=10, sticky='e')

        reset_button = Button(range_window, text="Reset", command=reset_selection, font=self.default_font)
        reset_button.grid(row=2, column=1, padx=10, pady=10)

        cancel_button = Button(range_window, text="Cancel", command=cancel_selection, font=self.default_font)
        cancel_button.grid(row=2, column=2, padx=10, pady=10, sticky='w')
    
    def update_plot(self):
        if not hasattr(self, 'canvas') or not self.canvas:
            self.current_fig, ax = plt.subplots(figsize=(10, 6))
            self.canvas = FigureCanvasTkAgg(self.current_fig, master=self.plot_frame)
            self.canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
        else:
            self.current_fig.clf()
            ax = self.current_fig.add_subplot(111)

        selected_sites = [site for site, var in self.site_vars.items() if var.get() == 1]
        date_range = (self.start_date or self.fiscal_start, self.end_date or self.fiscal_end)
        date_range = format_date_range(date_range)
        frequency = self.frequency_var.get()
        data_labels = bool(self.data_labels_var.get())
        prediction = bool(self.prediction_var.get())
        add_regression = bool(self.regression_var.get())
        remove_outliers = bool(self.outliers_var.get())

        # Get figure and statistics from plot_attendance
        fig, stats = plot_attendance(
            site=selected_sites if selected_sites else 'all',
            frequency=frequency,
            date_range=date_range,
            data_labels=data_labels,
            prediction=prediction,
            add_regression=add_regression,
            remove_outliers=remove_outliers,
            fig=self.current_fig,
            ax=ax
        )

        # Update the canvas
        self.canvas.draw()

        # Display stats in summary_frame
        self.display_statistics(stats)

    def display_statistics(self, stats):
        for widget in self.summary_frame.winfo_children():
            widget.destroy()

        # Create labels for data summary on the right
        data_summary_label = tk.Label(self.summary_frame, text="Data Summary", font=self.bold_font)
        data_summary_label.grid(row=0, column=0, columnspan=2, padx=(0, 10), pady=5, sticky='w')

        idx = 1
        for stat, value in stats.items():
            prefix_label = tk.Label(self.summary_frame, text=f"{stat}", font=self.bold_font)
            prefix_label.grid(row=idx, column=0, padx=(0, 10), pady=5, sticky='w')  # Left padding only

            data_label = tk.Label(self.summary_frame, text=f"{value:.2f}", font=self.default_font)
            data_label.grid(row=idx, column=1, padx=(10, 10), pady=5, sticky='w')  # Right padding only

            idx += 1

    def get_selected_sites(self):
        # Collect selected sites from checkbuttons
        return [site for site, var in self.site_vars.items() if var.get() == 1]

    def get_date_range(self):
        # Return start and end date range for plotting
        return (self.start_date or self.fiscal_start, self.end_date or self.fiscal_end)


class StatisticsTestingTab(ttk.Frame):
    def __init__(self, parent, default_font, bold_font):
        self.default_font = default_font
        self.bold_font = bold_font
        super().__init__(parent)
        self.create_widgets() 

    def create_widgets(self):
        pass


# Set high resolution for Windows
if os.name == 'nt':
    windll.shcore.SetProcessDpiAwareness(1)
# Initialize the Tkinter root window
root = Tk()
# root.attributes('-fullscreen', True)

screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# Calculate the available vertical space on Windows
if os.name == 'nt':  # Check if the OS is Windows
    # Get the working area dimensions using tkinter's winfo method
    screen_height = screen_height - (root.winfo_vrooty() - root.winfo_rooty())
else:
    # For macOS and other OS, use the total screen height
    screen_height = screen_height

# root.geometry('%dx%d' % (screen_width, screen_height))
root.state("zoomed")

# Create an instance of the SRS_Master class
app = SRS_Master(root)

# Register the cleanup function to run when the app is closing
atexit.register(app.cleanup_temp_files)

# Run the application
root.mainloop()
